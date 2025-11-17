"""
Utility functions for the AI Tester framework
Extracted from legacy code
"""

import re
import json
import base64
from typing import Optional, Dict
from html import escape as _html_escape


def slugify(value: str) -> str:
    """Convert text to URL-friendly slug."""
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9\-\s]", "", value)
    value = re.sub(r"\s+", "-", value)
    value = re.sub(r"-+", "-", value)
    return value or "ticket"


def safe_json_extract(text: str) -> Optional[dict]:
    """Safely extract JSON from text that may contain markdown code blocks."""
    if not text:
        print("DEBUG safe_json_extract: text is empty")
        return None
    s = text.strip()
    
    # Remove markdown code blocks
    s = re.sub(r"^\s*```(?:json)?\s*|\s*```\s*$", "", s, flags=re.I)
    
    # Try direct parse
    try:
        result = json.loads(s)
        print(f"DEBUG safe_json_extract: Successfully parsed JSON directly")
        return result
    except Exception as e:
        print(f"DEBUG safe_json_extract: Direct parse failed: {e}")
    
    # Try to find JSON object in text
    m = re.search(r"(\{(?:.|\n)*\})", s, re.DOTALL)
    if m:
        try:
            result = json.loads(m.group(1))
            print(f"DEBUG safe_json_extract: Successfully extracted JSON via regex")
            return result
        except Exception as e:
            print(f"DEBUG safe_json_extract: Regex extraction parse failed: {e}")
            print(f"DEBUG safe_json_extract: Extracted text (first 200 chars): {m.group(1)[:200]}")
    
    print("DEBUG safe_json_extract: All parsing attempts failed")
    return None


def clean_jira_text_for_llm(text: str) -> str:
    """
    Clean Jira text to remove out-of-scope content before sending to LLM.
    Removes:
    - Strikethrough text (~~text~~)
    - Parenthetical scope removal notes
    - Lines marked as removed from scope
    """
    if not text:
        return ""
    
    lines = text.split("\n")
    cleaned_lines = []
    
    for line in lines:
        # Skip lines that are entirely strikethrough or contain removal markers
        if line.strip().startswith("~~") or "(removed from scope)" in line.lower():
            continue
        
        # Remove inline strikethrough text
        line = re.sub(r"~~[^~]+~~", "", line)
        
        # Remove parenthetical scope notes with preceding word
        line = re.sub(r"\b\w+\s*\([^)]*removed from scope[^)]*\)", "", line, flags=re.IGNORECASE)
        
        if line.strip():
            cleaned_lines.append(line)
    
    return "\n".join(cleaned_lines)


# ---------- Jira-like HTML rendering ----------
JIRA_HTML_CSS = """
<style>
.jira { font-family: 'Segoe UI','Inter',Arial; color:#E6E9EF; font-weight:600; line-height:1.6; }
.jira h1{font-weight:800;font-size:22px;margin:16px 0 8px}
.jira h2{font-weight:800;font-size:20px;margin:14px 0 6px}
.jira h3{font-weight:800;font-size:18px;margin:12px 0 6px}
.jira h4{font-weight:800;font-size:16px;margin:10px 0 6px}
.jira p{margin:8px 0}
.jira ul, .jira ol{margin:8px 0 8px 24px}
.jira li{margin:4px 0}
.jira blockquote{border-left:3px solid #3a4761; padding:4px 10px; color:#cfd6e6; margin:8px 0}
.jira pre{background:#0b1220;border:1px solid #1a2740;border-radius:8px;padding:10px; overflow:auto}
.jira code{background:#0b1220;border:1px solid #1a2740;border-radius:6px;padding:1px 4px;}
.jira hr{border:0;height:1px;background:#1a2740;margin:12px 0}
.jira table{border-collapse:collapse;margin:8px 0}
.jira th,.jira td{border:1px solid #2a3650;padding:6px 8px}
.jira .panel{border:1px solid #2a3650;border-radius:8px;padding:8px 10px;background:#0e1524;margin:8px 0}
</style>
"""


def _inline_html(nodes):
    """Convert inline nodes to HTML."""
    out = []
    for n in nodes or []:
        if n.get("type") == "text":
            txt = _html_escape(n.get("text", "") or "")
            for m in (n.get("marks") or []):
                t = m.get("type")
                if t == "strong":
                    txt = f"<strong>{txt}</strong>"
                elif t == "em":
                    txt = f"<em>{txt}</em>"
                elif t == "code":
                    txt = f"<code>{txt}</code>"
                elif t == "underline":
                    txt = f"<u>{txt}</u>"
                elif t == "strike":
                    txt = f"<s>{txt}</s>"
            out.append(txt)
        elif n.get("type") == "hardBreak":
            out.append("<br/>")
    return "".join(out)


def adf_to_html(adf: Dict) -> str:
    """Convert Atlassian Document Format to HTML."""
    def walk(node):
        t = node.get("type")
        if t == "doc":
            return "".join(walk(c) for c in (node.get("content") or []))
        if t == "paragraph":
            html = _inline_html(node.get("content", []))
            return f"<p>{html if html.strip() else '&nbsp;'}</p>"
        if t == "heading":
            level = node.get("attrs", {}).get("level", 2)
            return f"<h{level}>{_inline_html(node.get('content', []))}</h{level}>"
        if t == "bulletList":
            return "<ul>" + "".join(walk(li) for li in (node.get("content") or [])) + "</ul>"
        if t == "orderedList":
            start = node.get("attrs", {}).get("order", 1)
            return f"<ol start='{start}'>" + "".join(walk(li) for li in (node.get("content") or [])) + "</ol>"
        if t == "listItem":
            return f"<li>{''.join(walk(c) for c in (node.get('content') or []))}</li>"
        if t == "blockquote":
            return f"<blockquote>{''.join(walk(c) for c in (node.get('content') or []))}</blockquote>"
        if t == "codeBlock":
            return f"<pre><code>{_inline_html(node.get('content', []))}</code></pre>"
        if t == "rule":
            return "<hr/>"
        if t == "panel":
            return f"<div class='panel'>{''.join(walk(c) for c in (node.get('content') or []))}</div>"
        if t == "table":
            return "<table>" + "".join(walk(r) for r in (node.get("content") or [])) + "</table>"
        if t == "tableRow":
            return "<tr>" + "".join(walk(c) for c in (node.get("content") or [])) + "</tr>"
        if t in ("tableHeader", "tableCell"):
            tag = "th" if t == "tableHeader" else "td"
            return f"<{tag}>" + "".join(walk(c) for c in (node.get("content") or [])) + f"</{tag}>"
        if t == "text":
            return _html_escape(node.get("text", ""))
        return ""
    
    try:
        return JIRA_HTML_CSS + f"<div class='jira'>{walk(adf or {})}</div>"
    except Exception:
        return JIRA_HTML_CSS + "<div class='jira'><p>(unable to render)</p></div>"


def plain_to_html(text: str) -> str:
    """Convert plain text to HTML."""
    safe = _html_escape(text or "").replace("\n\n", "</p><p>").replace("\n", "<br/>")
    return JIRA_HTML_CSS + f"<div class='jira'><p>{safe}</p></div>"


def adf_to_plaintext(adf: Dict) -> str:
    """Convert Atlassian Document Format to readable plain text."""
    out = []
    
    def walk(node, indent_level=0):
        node_type = node.get("type", "")
        content = node.get("content", [])
        
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            heading_text = "".join([c.get("text", "") for c in content if c.get("type") == "text"])
            out.append("\n" + heading_text + "\n")
        
        elif node_type == "paragraph":
            para_text = ""
            for c in content:
                if c.get("type") == "text":
                    para_text += c.get("text", "")
                elif c.get("type") == "hardBreak":
                    para_text += "\n"
            if para_text.strip():
                out.append("  " * indent_level + para_text + "\n")
        
        elif node_type == "bulletList":
            for item in content:
                if item.get("type") == "listItem":
                    item_text = ""
                    for item_content in item.get("content", []):
                        if item_content.get("type") == "paragraph":
                            for text_node in item_content.get("content", []):
                                if text_node.get("type") == "text":
                                    item_text += text_node.get("text", "")
                    if item_text.strip():
                        out.append("  " * indent_level + "• " + item_text + "\n")
                    
                    for nested in item.get("content", []):
                        if nested.get("type") in ["bulletList", "orderedList"]:
                            walk(nested, indent_level + 1)
        
        elif node_type == "orderedList":
            for idx, item in enumerate(content, 1):
                if item.get("type") == "listItem":
                    item_text = ""
                    for item_content in item.get("content", []):
                        if item_content.get("type") == "paragraph":
                            for text_node in item_content.get("content", []):
                                if text_node.get("type") == "text":
                                    item_text += text_node.get("text", "")
                    if item_text.strip():
                        out.append("  " * indent_level + f"{idx}. " + item_text + "\n")
                    
                    for nested in item.get("content", []):
                        if nested.get("type") in ["bulletList", "orderedList"]:
                            walk(nested, indent_level + 1)
        
        elif node_type == "codeBlock":
            code_text = "".join([c.get("text", "") for c in content if c.get("type") == "text"])
            out.append("\n" + code_text + "\n")
        
        elif node_type == "text":
            out.append(node.get("text", ""))
        
        else:
            for child in content:
                walk(child, indent_level)
    
    if adf:
        walk(adf)
    
    result = "".join(out)
    result = re.sub(r"\n{4,}", "\n\n", result)
    return result.strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        import PyPDF2
        import io
        pdf_file = io.BytesIO(pdf_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


def extract_text_from_word(docx_bytes: bytes) -> str:
    """Extract text from Word document bytes."""
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"Error extracting Word text: {e}")
        return ""


def extract_images_from_word(docx_bytes: bytes) -> list:
    """
    Extract embedded images from Word document bytes.

    Args:
        docx_bytes: Word document as bytes

    Returns:
        List of dicts with 'data' (base64) and 'mime_type' keys
    """
    try:
        import docx
        import io
        from docx.oxml import parse_xml

        doc = docx.Document(io.BytesIO(docx_bytes))
        images = []

        # Get all image relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob

                    # Determine mime type from content type
                    mime_type = image_part.content_type

                    # Encode to base64
                    image_base64 = encode_image_to_base64(image_bytes, mime_type)

                    images.append({
                        'data': image_base64,
                        'mime_type': mime_type,
                        'data_url': f"data:{mime_type};base64,{image_base64}"
                    })
                    print(f"DEBUG: Extracted image from Word doc: {mime_type}, {len(image_bytes)} bytes")
                except Exception as e:
                    print(f"DEBUG: Failed to extract image from Word doc: {e}")
                    continue

        print(f"DEBUG: Extracted {len(images)} images from Word document")
        return images

    except Exception as e:
        print(f"Error extracting images from Word: {e}")
        return []


def encode_image_to_base64(image_bytes: bytes, mime_type: str) -> str:
    """Encode image bytes to base64 for AI vision."""
    return base64.b64encode(image_bytes).decode('utf-8')

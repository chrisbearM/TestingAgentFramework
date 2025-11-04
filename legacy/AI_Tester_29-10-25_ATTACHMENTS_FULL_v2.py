#!/usr/bin/env python3
"""
AI Tester GUI – Jira → Azure DevOps

New:
- **Feature to analyze** tab (first): enter Epic key, fetch Epic + all child issues; keep a
  Feature Context for this session (epic + children).
- Ticket generation now includes Feature Context → more robust, cross-story test cases.
- **Comprehensive Attachment Processing**: Fetches and processes images, PDFs, and Word docs
  * Works across ALL workflows: Feature/Epic/Initiative loading, Overview generation, 
    Readiness assessment, Test case generation
  * Images analyzed with GPT-4o vision API
  * Text extracted from PDFs and Word documents
  * Attachment content incorporated into all AI analyses

Existing:
- Sign In first (centered); Jira base hard-coded to https://powerfleet.atlassian.net/
- Top-centered tabs after login; bold typography (Nebula theme)
- Jira-like HTML rendering on Feature & Ticket pages
- Test Cases page: floating/wrapping rounded tiles with expand/collapse + select in corner
- Export CSV/TSV + Azure DevOps CSV bundle

Requirements for attachment processing:
- PyPDF2: pip install PyPDF2
- python-docx: pip install python-docx
"""
from __future__ import annotations

import os, sys, re, csv, json, time, zipfile, base64, io
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from html import escape as _html_escape

import requests
from PySide6.QtCore import Qt, QThread, Signal, QPoint, QRect, QSize, QTimer, QPropertyAnimation, QEasingCurve, QParallelAnimationGroup, QSequentialAnimationGroup
from PySide6.QtGui import QFontDatabase, QFont, QColor
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QFormLayout,
    QLineEdit, QPushButton, QLabel, QStackedWidget, QMessageBox,
    QTableWidget, QTableWidgetItem, QFileDialog, QSpinBox, QDialog, QDialogButtonBox,
    QCheckBox, QFrame, QProgressDialog, QPlainTextEdit, QGraphicsDropShadowEffect,
    QTextBrowser, QScrollArea, QSizePolicy, QLayout, QComboBox, QProgressBar,
    QGraphicsOpacityEffect, QGraphicsBlurEffect
)

# ---------- Constants ----------
JIRA_BASE_URL = "https://powerfleet.atlassian.net/"

# ---------- Utilities ----------
def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9\-\s]", "", value)
    value = re.sub(r"\s+", "-", value)
    value = re.sub(r"-+", "-", value)
    return value or "ticket"

def safe_json_extract(text: str) -> Optional[dict]:
    if not text:
        print("DEBUG safe_json_extract: text is empty")
        return None
    s = text.strip()
    
    # Remove markdown code blocks
    s = re.sub(r"^\s*```(?:json)?\s*|\s*```\s*$", "", s, flags=re.I)
    
    # Try direct parse
    try:
        result = json.loads(s)
        print(f"DEBUG safe_json_extract: Successfully parsed JSON directly")
        return result
    except Exception as e:
        print(f"DEBUG safe_json_extract: Direct parse failed: {e}")
    
    # Try to find JSON object in text
    m = re.search(r"(\{(?:.|\n)*\})", s, re.DOTALL)
    if m:
        try:
            result = json.loads(m.group(1))
            print(f"DEBUG safe_json_extract: Successfully extracted JSON via regex")
            return result
        except Exception as e:
            print(f"DEBUG safe_json_extract: Regex extraction parse failed: {e}")
            # Print the extracted text for debugging
            print(f"DEBUG safe_json_extract: Extracted text (first 200 chars): {m.group(1)[:200]}")
    
    print("DEBUG safe_json_extract: All parsing attempts failed")
    return None

# ---------- Fonts ----------
from PySide6.QtGui import QFont
def load_brand_fonts():
    candidates = [
        str(Path(__file__).with_name("Inter-Variable.ttf")),
        str(Path(__file__).with_name("PlusJakartaSans-Variable.ttf")),
        str(Path(__file__).parent.joinpath("fonts", "Inter-Variable.ttf")),
        str(Path(__file__).parent.joinpath("fonts", "PlusJakartaSans-Variable.ttf")),
    ]
    for p in candidates:
        if Path(p).exists():
            QFontDatabase.addApplicationFont(p)
    app_font = QFont("Inter"); app_font.setPointSizeF(10.5); QApplication.setFont(app_font)

# ---------- Visual helpers ----------
def apply_soft_shadow(w: QWidget, radius=32, opacity=0.25, dx=0, dy=10, color="#000000"):
    eff = QGraphicsDropShadowEffect(w)
    eff.setBlurRadius(radius)
    eff.setOffset(dx, dy)
    c = QColor(color)
    c.setAlphaF(opacity)
    eff.setColor(c)
    w.setGraphicsEffect(eff)

def apply_glow_effect(w: QWidget, color="#3B82F6", radius=20, opacity=0.4):
    """Apply a colored glow effect to widgets."""
    eff = QGraphicsDropShadowEffect(w)
    eff.setBlurRadius(radius)
    eff.setOffset(0, 0)
    c = QColor(color)
    c.setAlphaF(opacity)
    eff.setColor(c)
    w.setGraphicsEffect(eff)

# ---------- Jira-like HTML rendering ----------
JIRA_HTML_CSS = """
<style>
.jira { font-family: 'Segoe UI','Inter',Arial; color:#E6E9EF; font-weight:600; line-height:1.6; }
.jira h1{font-weight:800;font-size:22px;margin:16px 0 8px}
.jira h2{font-weight:800;font-size:20px;margin:14px 0 6px}
.jira h3{font-weight:800;font-size:18px;margin:12px 0 6px}
.jira h4{font-weight:800;font-size:16px;margin:10px 0 6px}
.jira p{margin:8px 0}
.jira ul, .jira ol{margin:8px 0 8px 24px}
.jira li{margin:4px 0}
.jira blockquote{border-left:3px solid #3a4761; padding:4px 10px; color:#cfd6e6; margin:8px 0}
.jira pre{background:#0b1220;border:1px solid #1a2740;border-radius:8px;padding:10px; overflow:auto}
.jira code{background:#0b1220;border:1px solid #1a2740;border-radius:6px;padding:1px 4px;}
.jira hr{border:0;height:1px;background:#1a2740;margin:12px 0}
.jira table{border-collapse:collapse;margin:8px 0}
.jira th,.jira td{border:1px solid #2a3650;padding:6px 8px}
.jira .panel{border:1px solid #2a3650;border-radius:8px;padding:8px 10px;background:#0e1524;margin:8px 0}
</style>
"""
def _inline_html(nodes):
    out=[]
    for n in nodes or []:
        if n.get("type")=="text":
            txt=_html_escape(n.get("text","") or "")
            for m in (n.get("marks") or []):
                t=m.get("type")
                if t=="strong": txt=f"<strong>{txt}</strong>"
                elif t=="em": txt=f"<em>{txt}</em>"
                elif t=="code": txt=f"<code>{txt}</code>"
                elif t=="underline": txt=f"<u>{txt}</u>"
                elif t=="strike": txt=f"<s>{txt}</s>"
            out.append(txt)
        elif n.get("type")=="hardBreak":
            out.append("<br/>")
    return "".join(out)

def adf_to_html(adf: Dict) -> str:
    def walk(node):
        t=node.get("type")
        if t=="doc": return "".join(walk(c) for c in (node.get("content") or []))
        if t=="paragraph":
            html=_inline_html(node.get("content",[]))
            return f"<p>{html if html.strip() else '&nbsp;'}</p>"
        if t=="heading":
            level=node.get("attrs",{}).get("level",2)
            return f"<h{level}>{_inline_html(node.get('content',[]))}</h{level}>"
        if t=="bulletList": return "<ul>"+"".join(walk(li) for li in (node.get("content") or []))+"</ul>"
        if t=="orderedList":
            start=node.get("attrs",{}).get("order",1)
            return f"<ol start='{start}'>"+"".join(walk(li) for li in (node.get("content") or []))+"</ol>"
        if t=="listItem": return f"<li>{''.join(walk(c) for c in (node.get('content') or []))}</li>"
        if t=="blockquote": return f"<blockquote>{''.join(walk(c) for c in (node.get('content') or []))}</blockquote>"
        if t=="codeBlock": return f"<pre><code>{_inline_html(node.get('content',[]))}</code></pre>"
        if t=="rule": return "<hr/>"
        if t=="panel": return f"<div class='panel'>{''.join(walk(c) for c in (node.get('content') or []))}</div>"
        if t=="table": return "<table>"+"".join(walk(r) for r in (node.get("content") or []))+"</table>"
        if t=="tableRow": return "<tr>"+"".join(walk(c) for c in (node.get("content") or []))+"</tr>"
        if t in ("tableHeader","tableCell"):
            tag="th" if t=="tableHeader" else "td"
            return f"<{tag}>"+"".join(walk(c) for c in (node.get("content") or []))+f"</{tag}>"
        if t=="text": return _html_escape(node.get("text",""))
        return ""
    try:
        return JIRA_HTML_CSS+f"<div class='jira'>{walk(adf or {})}</div>"
    except Exception:
        return JIRA_HTML_CSS+"<div class='jira'><p>(unable to render)</p></div>"

def plain_to_html(text: str) -> str:
    safe=_html_escape(text or "").replace("\n\n","</p><p>").replace("\n","<br/>")
    return JIRA_HTML_CSS+f"<div class='jira'><p>{safe}</p></div>"

def adf_to_plaintext(adf: Dict) -> str:
    """Convert Atlassian Document Format to readable plain text with preserved structure."""
    out = []
    
    def walk(node, indent_level=0):
        node_type = node.get("type", "")
        content = node.get("content", [])
        
        # Headings
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            heading_text = "".join([c.get("text", "") for c in content if c.get("type") == "text"])
            out.append("\n" + heading_text + "\n")
        
        # Paragraphs
        elif node_type == "paragraph":
            para_text = ""
            for c in content:
                if c.get("type") == "text":
                    text = c.get("text", "")
                    # Check for bold/italic marks
                    marks = c.get("marks", [])
                    if marks:
                        # Just add the text, don't add markdown symbols
                        para_text += text
                    else:
                        para_text += text
                elif c.get("type") == "hardBreak":
                    para_text += "\n"
            if para_text.strip():
                out.append("  " * indent_level + para_text + "\n")
        
        # Bullet lists
        elif node_type == "bulletList":
            for item in content:
                if item.get("type") == "listItem":
                    # Get text from the list item
                    item_text = ""
                    for item_content in item.get("content", []):
                        if item_content.get("type") == "paragraph":
                            for text_node in item_content.get("content", []):
                                if text_node.get("type") == "text":
                                    item_text += text_node.get("text", "")
                    if item_text.strip():
                        out.append("  " * indent_level + "• " + item_text + "\n")
                    
                    # Handle nested lists
                    for nested in item.get("content", []):
                        if nested.get("type") in ["bulletList", "orderedList"]:
                            walk(nested, indent_level + 1)
        
        # Ordered lists
        elif node_type == "orderedList":
            for idx, item in enumerate(content, 1):
                if item.get("type") == "listItem":
                    item_text = ""
                    for item_content in item.get("content", []):
                        if item_content.get("type") == "paragraph":
                            for text_node in item_content.get("content", []):
                                if text_node.get("type") == "text":
                                    item_text += text_node.get("text", "")
                    if item_text.strip():
                        out.append("  " * indent_level + f"{idx}. " + item_text + "\n")
                    
                    # Handle nested lists
                    for nested in item.get("content", []):
                        if nested.get("type") in ["bulletList", "orderedList"]:
                            walk(nested, indent_level + 1)
        
        # Code blocks
        elif node_type == "codeBlock":
            code_text = "".join([c.get("text", "") for c in content if c.get("type") == "text"])
            out.append("\n" + code_text + "\n")
        
        # Plain text nodes (shouldn't happen at top level, but handle it)
        elif node_type == "text":
            out.append(node.get("text", ""))
        
        # Recursively process other content
        else:
            for child in content:
                walk(child, indent_level)
    
    if adf:
        walk(adf)
    
    # Clean up excessive newlines while preserving structure
    result = "".join(out)
    result = re.sub(r"\n{4,}", "\n\n", result)  # Max 2 consecutive newlines
    return result.strip()


def clean_jira_text_for_llm(text: str) -> str:
    """
    Clean Jira text to remove out-of-scope content before sending to LLM.
    Removes:
    - Strikethrough text (~~text~~)
    - Parenthetical scope removal notes AND the word/phrase before them
    - Lines that are marked as removed from scope
    - ANY operations/terms that are marked as removed ANYWHERE in the text
    """
    if not text:
        return text
    
    # Detect ANY words marked as removed (strikethrough or with removal notes)
    removed_terms = set()
    
    # Find strikethrough terms: ~~word~~
    strikethrough_matches = re.findall(r'~~(\w+)~~', text, re.IGNORECASE)
    removed_terms.update([term.lower() for term in strikethrough_matches])
    
    # Find words followed by removal notes: "word (removed from scope)"
    removal_note_matches = re.findall(
        r'\b(\w+)\s*\([^)]*(?:removed from scope|out of scope|not in scope|scope removed|deleted from scope)[^)]*\)',
        text,
        re.IGNORECASE
    )
    removed_terms.update([term.lower() for term in removal_note_matches])
    
    # Remove strikethrough text (Jira markdown: ~~text~~)
    text = re.sub(r'~{2,}[^~]+~{2,}', '', text)
    
    # Remove words/phrases followed by removal notes
    removal_phrases = [
        r'\b\w+\s*\([^)]*(?:removed from scope|out of scope|not in scope|scope removed|deleted from scope)[^)]*\)',
        r',?\s*and\s+\w+\s*\([^)]*(?:removed from scope|out of scope)[^)]*\)',
        r'\band\s+~~[^~]+~~',
    ]
    for pattern in removal_phrases:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # Remove standalone parenthetical notes about removed scope
    removal_patterns = [
        r'\([^)]*removed from scope[^)]*\)',
        r'\([^)]*out of scope[^)]*\)',
        r'\([^)]*not in scope[^)]*\)',
        r'\([^)]*scope removed[^)]*\)',
        r'\([^)]*deleted from scope[^)]*\)',
        r'\([^)]*operation removed from scope[^)]*\)',
    ]
    for pattern in removal_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    # For each removed term, remove ALL occurrences throughout the text
    for term in removed_terms:
        if not term:
            continue
        
        # Remove from comma-separated lists: "create, update, and delete" -> "create and update"
        text = re.sub(rf',?\s+and\s+{term}\b', '', text, flags=re.IGNORECASE)
        text = re.sub(rf'\b{term},?\s+and\s+', '', text, flags=re.IGNORECASE)
        
        # Remove from slash-separated lists: "Create/Update/Delete" -> "Create/Update"
        text = re.sub(rf'/{term}(?=/|\)|\s|,|$)', '', text, flags=re.IGNORECASE)
        text = re.sub(rf'{term}/(?=\w)', '', text, flags=re.IGNORECASE)
        
        # Remove standalone occurrences with "operations": "delete operations" -> "operations"
        text = re.sub(rf'\b{term}\s+operations?\b', 'operations', text, flags=re.IGNORECASE)
    
    # Remove lines that contain removal indicators
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line_lower = line.lower()
        # Skip lines that are mostly strikethrough or contain removal notes
        if (line.strip().startswith('~~') or 
            'removed from scope' in line_lower or
            'out of scope' in line_lower or
            'not in scope' in line_lower):
            continue
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # Clean up artifacts
    text = re.sub(r'\s*,\s*,\s*', ', ', text)  # Double commas
    text = re.sub(r'\s+and\s+and\s+', ' and ', text)  # Double "and"
    text = re.sub(r',\s*and\s+', ' and ', text)  # ", and" -> " and"
    text = re.sub(r'\(\s*\)', '', text)  # Empty parentheses
    text = re.sub(r'\s+operations', ' operations', text)  # Extra space before operations
    text = re.sub(r'/\s*/', '/', text)  # Clean up double slashes
    text = re.sub(r'(^|[^/])/$', r'\1', text)  # Remove trailing slash
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple blank lines
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces
    text = re.sub(r'\.\s*\.', '.', text)  # Double periods
    text = text.strip()
    
    return text


# ---------- Attachment Processing Helpers ----------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        text_content = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_content.append(f"[Page {page_num}]\n{page_text}")
        
        full_text = "\n\n".join(text_content)
        print(f"DEBUG: Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
        return full_text
    
    except ImportError:
        print("DEBUG: PyPDF2 not installed - PDF text extraction unavailable")
        return "[PDF content - PyPDF2 not installed. Install with: pip install PyPDF2]"
    except Exception as e:
        print(f"DEBUG: Error extracting text from PDF: {e}")
        return f"[PDF content - Error extracting text: {str(e)}]"


def extract_text_from_word(file_bytes: bytes) -> str:
    """Extract text content from a Word document (.docx)."""
    try:
        from docx import Document
        
        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            table_text = [f"\n[Table {table_num}]"]
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            text_content.extend(table_text)
        
        full_text = "\n\n".join(text_content)
        print(f"DEBUG: Extracted {len(full_text)} characters from Word document")
        return full_text
    
    except ImportError:
        print("DEBUG: python-docx not installed - Word document text extraction unavailable")
        return "[Word document - python-docx not installed. Install with: pip install python-docx]"
    except Exception as e:
        print(f"DEBUG: Error extracting text from Word document: {e}")
        return f"[Word document - Error extracting text: {str(e)}]"


def encode_image_to_base64(file_bytes: bytes, mime_type: str) -> str:
    """Encode image bytes to base64 for GPT-4o vision."""
    return base64.b64encode(file_bytes).decode('utf-8')


def build_ticket_html(summary: str, desc_html: str, ac_blocks: List[str]) -> str:
    parts=[JIRA_HTML_CSS,"<div class='jira'>"]
    if summary: parts.append(f"<h2>{_html_escape(summary)}</h2>")
    parts.append(desc_html.replace(JIRA_HTML_CSS,""))
    if ac_blocks:
        parts.append("<h3>Acceptance Criteria</h3>")
        for block in ac_blocks:
            lines=[ln.strip() for ln in (block or "").splitlines() if ln.strip()]
            if any(re.match(r"^[-*•]\s+", ln) for ln in lines):
                items="".join(f"<li>{_html_escape(re.sub(r'^[-*•]\\s+','',ln))}</li>" for ln in lines)
                parts.append(f"<ul>{items}</ul>")
            else:
                parts.append(f"<p>{_html_escape(block).replace('\\n','<br/>')}</p>")
    parts.append("</div>"); return "".join(parts)

def build_feature_html(epic_key: str, epic_title: str, epic_html: str, children: List[Dict[str,str]]) -> str:
    parts=[JIRA_HTML_CSS,"<div class='jira'>"]
    parts.append(f"<h2>Epic {epic_key}: {_html_escape(epic_title)}</h2>")
    parts.append(epic_html.replace(JIRA_HTML_CSS,""))
    parts.append("<h3>Child issues</h3>")
    if children:
        parts.append("<ul>")
        for ch in children:
            parts.append(f"<li><strong>{_html_escape(ch['key'])}</strong>: {_html_escape(ch['summary'])}</li>")
        parts.append("</ul>")
    else:
        parts.append("<p>(no child issues found)</p>")
    parts.append("</div>")
    return "".join(parts)

def build_initiative_html(init_data: Dict) -> str:
    """Build HTML view for Initiative with its Epics and children."""
    parts = [JIRA_HTML_CSS, "<div class='jira'>"]
    
    # Initiative header
    init = init_data["initiative"]
    parts.append(f"<h1>Initiative {init['key']}: {_html_escape(init['summary'])}</h1>")
    
    if init.get('desc'):
        # Show full description with preserved structure
        # Convert plain text to HTML with proper formatting
        desc_html = _html_escape(init['desc'])
        # Preserve line breaks and structure
        desc_html = desc_html.replace('\n', '<br/>')
        parts.append(f"<div style='white-space: pre-wrap;'>{desc_html}</div>")
    
    # Epics section
    epics = init_data.get("epics", [])
    parts.append(f"<h2>Epics ({len(epics)})</h2>")
    
    if epics:
        for epic in epics:
            parts.append(f"<h3>Epic {_html_escape(epic['key'])}: {_html_escape(epic['summary'])}</h3>")
            if epic.get('desc'):
                # Show first 300 chars with preserved structure
                desc_preview = epic['desc'][:300]
                desc_html = _html_escape(desc_preview).replace('\n', '<br/>')
                parts.append(f"<div style='white-space: pre-wrap;'>{desc_html}...</div>")
            
            # Children
            children = epic.get('children', [])
            if children:
                parts.append(f"<h4>Child Issues ({len(children)})</h4>")
                parts.append("<ul>")
                for child in children[:10]:  # Show first 10
                    parts.append(f"<li><strong>{_html_escape(child['key'])}</strong>: {_html_escape(child['summary'])}</li>")
                if len(children) > 10:
                    parts.append(f"<li><em>... and {len(children) - 10} more</em></li>")
                parts.append("</ul>")
            else:
                parts.append("<p><em>(no child issues)</em></p>")
    else:
        parts.append("<p>(no epics found)</p>")
    
    parts.append("</div>")
    return "".join(parts)

def build_overview_html(overview_data: Dict) -> str:
    parts = [JIRA_HTML_CSS, "<div class='jira'>"]
    
    if overview_data.get("overview"):
        parts.append(f"<h3>Overview</h3><p>{_html_escape(overview_data['overview'])}</p>")
    
    if overview_data.get("problem_statement"):
        parts.append(f"<h3>Problem Statement</h3><p>{_html_escape(overview_data['problem_statement'])}</p>")
    
    if overview_data.get("solution_approach"):
        parts.append(f"<h3>Solution Approach</h3><p>{_html_escape(overview_data['solution_approach'])}</p>")
    
    if overview_data.get("key_capabilities"):
        parts.append("<h3>Key Capabilities</h3><ul>")
        for cap in overview_data["key_capabilities"]:
            parts.append(f"<li>{_html_escape(cap)}</li>")
        parts.append("</ul>")
    
    if overview_data.get("user_impact"):
        parts.append(f"<h3>User Impact</h3><p>{_html_escape(overview_data['user_impact'])}</p>")
    
    if overview_data.get("technical_considerations"):
        parts.append("<h3>Technical Considerations</h3><ul>")
        for tech in overview_data["technical_considerations"]:
            parts.append(f"<li>{_html_escape(tech)}</li>")
        parts.append("</ul>")
    
    if overview_data.get("business_value"):
        parts.append(f"<h3>Business Value</h3><p>{_html_escape(overview_data['business_value'])}</p>")
    
    parts.append("</div>")
    return "".join(parts)

def build_readiness_html(readiness_data: Dict) -> str:
    parts = [JIRA_HTML_CSS, "<div class='jira'>"]
    
    if readiness_data.get("summary"):
        parts.append(f"<h3>Assessment Summary</h3><p>{_html_escape(readiness_data['summary'])}</p>")
    
    if readiness_data.get("strengths"):
        parts.append("<h3>✓ Strengths</h3><ul>")
        for strength in readiness_data["strengths"]:
            parts.append(f"<li style='color:#6EE7B7;'>{_html_escape(strength)}</li>")
        parts.append("</ul>")
    
    if readiness_data.get("missing_elements"):
        parts.append("<h3>⚠ Missing or Unclear Elements</h3><ul>")
        for missing in readiness_data["missing_elements"]:
            parts.append(f"<li style='color:#FCA5A5;'>{_html_escape(missing)}</li>")
        parts.append("</ul>")
    
    if readiness_data.get("recommendations"):
        parts.append("<h3>💡 Recommendations</h3><ul>")
        for rec in readiness_data["recommendations"]:
            parts.append(f"<li style='color:#FDE047;'>{_html_escape(rec)}</li>")
        parts.append("</ul>")
    
    if readiness_data.get("quality_concerns"):
        concerns = readiness_data["quality_concerns"]
        # Handle if it's a string instead of a list
        if isinstance(concerns, str):
            concerns = [concerns]
        if concerns:  # Only show if there are actual concerns
            parts.append("<h3>Quality Concerns</h3><ul>")
            for concern in concerns:
                if concern:  # Skip empty strings
                    parts.append(f"<li>{_html_escape(str(concern))}</li>")
            parts.append("</ul>")
    
    if readiness_data.get("questions_for_author"):
        parts.append("<h3>❓ Questions for Ticket Author</h3>")
        parts.append("<p style='color:#A7B3C6; font-size:13px; margin-bottom:12px;'>Ask the ticket author these questions to clarify requirements and improve testability:</p>")
        parts.append("<ul>")
        for question in readiness_data["questions_for_author"]:
            parts.append(f"<li style='color:#93C5FD;'>{_html_escape(question)}</li>")
        parts.append("</ul>")
    
    if readiness_data.get("ideal_ticket_example"):
        parts.append("<h3>✨ Ideal Ticket Example (100% Ready)</h3>")
        parts.append("<p style='color:#A7B3C6; font-size:13px; margin-bottom:12px;'>Here's how this ticket would look if it scored 'Excellent' with all recommended improvements:</p>")
        
        # Parse the ideal ticket example and format it like a real Jira ticket
        ideal_text = readiness_data["ideal_ticket_example"]
        
        # Try to extract Summary, Description, and Acceptance Criteria sections
        import re
        
        # More robust regex patterns that handle various formats
        summary_match = re.search(r'(?:Summary|Title):\s*\n?(.+?)(?=\n\n|\nDescription|\nAcceptance|$)', ideal_text, re.IGNORECASE | re.DOTALL)
        desc_match = re.search(r'(?:Description|Details?):\s*\n?(.+?)(?=\n(?:Acceptance Criteria|Context provided|$))', ideal_text, re.IGNORECASE | re.DOTALL)
        ac_match = re.search(r'Acceptance Criteria:\s*\n?(.+?)(?=\n(?:Context provided|$))', ideal_text, re.IGNORECASE | re.DOTALL)
        
        ideal_summary = summary_match.group(1).strip() if summary_match else ""
        ideal_desc = desc_match.group(1).strip() if desc_match else ""
        ideal_ac = ac_match.group(1).strip() if ac_match else ""
        
        # If we couldn't parse sections, just show the whole thing as description
        if not ideal_summary and not ideal_desc and not ideal_ac:
            ideal_desc = ideal_text
        
        # Build Jira-style HTML for the ideal ticket
        parts.append("<div style='background:#0b1220; border:2px solid #1e3a5f; padding:20px; margin:12px 0; border-radius:12px;'>")
        parts.append("<div class='jira'>")
        
        if ideal_summary:
            parts.append(f"<h2>{_html_escape(ideal_summary)}</h2>")
        
        if ideal_desc:
            # Format description with proper line breaks
            desc_formatted = _html_escape(ideal_desc).replace('\n\n', '</p><p>').replace('\n', '<br/>')
            parts.append(f"<p>{desc_formatted}</p>")
        
        if ideal_ac:
            parts.append("<h3>Acceptance Criteria</h3>")
            # Split AC into lines and format as list if appropriate
            ac_lines = [ln.strip() for ln in ideal_ac.splitlines() if ln.strip()]
            if any(re.match(r"^[-*•]\s+", ln) for ln in ac_lines):
                items = "".join(f"<li>{_html_escape(re.sub(r'^[-*•]\\s+','',ln))}</li>" for ln in ac_lines)
                parts.append(f"<ul>{items}</ul>")
            else:
                # Format as bullet list even if not explicitly bulleted
                items = "".join(f"<li>{_html_escape(ln)}</li>" for ln in ac_lines)
                parts.append(f"<ul>{items}</ul>")
        
        parts.append("</div>")
        parts.append("</div>")
        
        # Add a clear marker at the end
        parts.append("<p style='color:#4ADE80; font-size:12px; font-style:italic; margin-top:8px; text-align:center;'>✓ Ideal ticket example complete</p>")
    
    parts.append("</div>")
    return "".join(parts)

# ---------- Jira Client (Agile epic issues + robust search fallbacks) ----------
class JiraClient:
    def __init__(self, base_url: str, email: str, api_token: str):
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        self.session.auth = (email, api_token)
        self.session.headers.update({
            "Accept": "application/json",
            "Content-Type": "application/json",
        })

    def get_issue(self, key: str) -> Dict:
        url = f"{self.base_url}/rest/api/3/issue/{key}"
        r = self.session.get(url, timeout=30)
        if r.status_code == 404:
            raise ValueError("Issue not found or you lack permission.")
        r.raise_for_status()
        return r.json()
    
    def get_attachments(self, issue_key: str) -> List[Dict]:
        """
        Fetch all attachments for a given issue.
        Returns list of attachment metadata dicts with keys: id, filename, mimeType, size, content (download URL)
        """
        try:
            issue = self.get_issue(issue_key)
            attachments = issue.get("fields", {}).get("attachment", [])
            print(f"DEBUG: Found {len(attachments)} attachments for {issue_key}")
            return attachments
        except Exception as e:
            print(f"DEBUG: Error fetching attachments for {issue_key}: {e}")
            return []
    
    def download_attachment(self, attachment_url: str) -> Optional[bytes]:
        """Download attachment content from URL."""
        try:
            print(f"DEBUG: Downloading attachment from {attachment_url}")
            r = self.session.get(attachment_url, timeout=60)
            r.raise_for_status()
            print(f"DEBUG: Downloaded {len(r.content)} bytes")
            return r.content
        except Exception as e:
            print(f"DEBUG: Error downloading attachment: {e}")
            return None
    
    def process_attachment(self, attachment: Dict) -> Optional[Dict]:
        """
        Process an attachment and extract relevant information.
        Returns dict with: filename, type (text/image), content
        """
        filename = attachment.get("filename", "unknown")
        mime_type = attachment.get("mimeType", "")
        size = attachment.get("size", 0)
        content_url = attachment.get("content", "")
        
        print(f"DEBUG: Processing attachment: {filename} ({mime_type}, {size} bytes)")
        
        # Skip if too large (> 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if size > max_size:
            print(f"DEBUG: Skipping {filename} - too large ({size} bytes)")
            return None
        
        # Download the file
        file_bytes = self.download_attachment(content_url)
        if not file_bytes:
            return None
        
        # Determine how to process based on mime type
        result = {
            "filename": filename,
            "mime_type": mime_type,
            "size": size
        }
        
        # PDF files
        if mime_type == "application/pdf" or filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
            result["type"] = "document"
            result["content"] = text
            return result
        
        # Word documents
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] or filename.lower().endswith(('.docx', '.doc')):
            text = extract_text_from_word(file_bytes)
            result["type"] = "document"
            result["content"] = text
            return result
        
        # Images (for GPT-4o vision)
        elif mime_type.startswith("image/"):
            # Supported image types
            if mime_type in ["image/jpeg", "image/png", "image/gif", "image/webp"]:
                base64_data = encode_image_to_base64(file_bytes, mime_type)
                result["type"] = "image"
                result["content"] = base64_data
                result["data_url"] = f"data:{mime_type};base64,{base64_data}"
                return result
            else:
                print(f"DEBUG: Unsupported image format: {mime_type}")
                return None
        
        # Plain text files
        elif mime_type.startswith("text/") or filename.lower().endswith(('.txt', '.md', '.csv')):
            try:
                text = file_bytes.decode('utf-8')
                result["type"] = "document"
                result["content"] = text
                return result
            except Exception as e:
                print(f"DEBUG: Error decoding text file: {e}")
                return None
        
        else:
            print(f"DEBUG: Unsupported file type: {mime_type}")
            return None

    # -- Issue fetch (unchanged)
    def _agile_epic_issues(self, epic_key: str, fields: List[str]) -> Optional[List[Dict]]:
        """
        Try Jira Software endpoint first:
          GET /rest/agile/1.0/epic/{epicKey}/issue?fields=...&maxResults=...&startAt=...
        Not all tenants have Agile; if 404/403/etc, return None so caller can try JQL.
        """
        base = f"{self.base_url}/rest/agile/1.0/epic/{epic_key}/issue"
        params = {
            "fields": ",".join(fields),
            "maxResults": "200",
            "startAt": "0",
        }
        issues: List[Dict] = []
        while True:
            r = self.session.get(base, params=params, timeout=30)
            if r.status_code in (401, 403, 404, 405):
                # Agile not available / not permitted / not found -> fallback to JQL
                return None
            if r.status_code >= 400:
                # Treat other errors as a hard stop for Agile; let caller try JQL
                return None
            data = r.json()
            batch = data.get("issues", []) or []
            issues.extend(batch)
            if len(batch) < int(params["maxResults"]):
                break
            # paginate
            params["startAt"] = str(int(params["startAt"]) + int(params["maxResults"]))
        return issues

    # -- low-level search variant runner
    def _search_once(self, api_ver: str, method: str, jql: str, fields: List[str], max_results: int):
        url = f"{self.base_url}/rest/api/{api_ver}/search"
        if method == "POST":
            payload = {
                "jql": jql,
                "maxResults": max_results,
                "fields": fields
            }
            print(f"DEBUG _search_once: POST to {url}")
            print(f"DEBUG _search_once: Payload: {payload}")
            r = self.session.post(url, json=payload, timeout=30)
        else:
            params = {"jql": jql, "maxResults": str(max_results), "fields": ",".join(fields)}
            print(f"DEBUG _search_once: GET to {url}")
            print(f"DEBUG _search_once: Params: {params}")
            r = self.session.get(url, params=params, timeout=30)
        
        print(f"DEBUG _search_once: Status code: {r.status_code}")
        
        if r.status_code == 410:
            # This shape is disabled on the tenant; allow fallback to another shape
            return None
        if r.status_code >= 400:
            print(f"DEBUG _search_once: Error response: {r.text[:500]}")
            raise requests.HTTPError(f"{r.status_code} {r.reason}: {r.text}")
        return r.json().get("issues", [])

    # -- robust search that tries POST/GET across v3 and v2
    def search_jql(self, jql: str, fields: List[str], max_results: int = 200) -> List[Dict]:
        variants = [("3", "POST"), ("3", "GET"), ("2", "POST"), ("2", "GET")]
        last_err = None
        all_410 = True
        
        for ver, method in variants:
            try:
                issues = self._search_once(ver, method, jql, fields, max_results)
                if issues is not None:
                    all_410 = False
                    return issues
                # issues is None means we got a 410, continue to next variant
            except Exception as e:
                all_410 = False  # Got a different error, not all 410s
                last_err = e
                continue
        
        # If all variants returned 410, return empty list instead of raising error
        if all_410:
            print(f"DEBUG: All search variants returned HTTP 410 (Gone) for JQL: {jql}")
            print(f"DEBUG: This usually means the Jira search endpoints are unavailable.")
            return []
        
        # If we got other errors, raise
        raise ValueError(f"Jira search failed for all variants. Last error: {last_err}")

    # -- public: get all children in an epic with multi-strategy fallback
    def get_children_of_epic(self, epic_key: str) -> List[Dict]:
        fields = ["summary", "description", "issuetype", "status"]

        # 1) Try the Agile epic issues endpoint first (best coverage, no JQL)
        agile = self._agile_epic_issues(epic_key, fields)
        if isinstance(agile, list) and agile:
            return agile

        # 2) Fall back to robust JQL search (cover classic & team-managed)
        jql_variants = [
            f'parent = {epic_key}',                        # team-managed
            f'"Epic Link" = {epic_key}',                   # company-managed classic
            f'issue in childIssuesOf("{epic_key}")',       # Advanced Roadmaps function
            f'parentEpic = {epic_key}',                    # some Cloud sites expose this
        ]
        seen = set()
        results: List[Dict] = []
        last_err = None
        for jql in jql_variants:
            try:
                issues = self.search_jql(jql, fields, max_results=500)
                for it in issues:
                    key = it.get("key")
                    if key and key not in seen:
                        seen.add(key); results.append(it)
            except Exception as e:
                last_err = e
                continue

        if not results and last_err:
            raise ValueError(
                f"Could not fetch child issues for epic {epic_key}. "
                f"Tried Agile endpoint and JQL variants: {', '.join(jql_variants)}. "
                f"Last error: {last_err}"
            )
        return results
    
    def get_initiative_details(self, initiative_key: str) -> Dict:
        """
        Fetch an Initiative and all its related Epics and their children.
        Returns: {
            "initiative": {key, summary, description},
            "epics": [{key, summary, description, children: [...]}]
        }
        """
        # 1. Get the Initiative itself
        initiative = self.get_issue(initiative_key)
        init_fields = initiative.get("fields", {})
        init_summary = init_fields.get("summary", initiative_key)
        init_desc = init_fields.get("description")
        
        # Parse description
        if isinstance(init_desc, dict) and init_desc.get("type") == "doc":
            init_desc_plain = adf_to_plaintext(init_desc)
        elif isinstance(init_desc, str):
            init_desc_plain = init_desc
        else:
            init_desc_plain = ""
        
        # Clean initiative description for LLM
        init_desc_plain = clean_jira_text_for_llm(init_desc_plain)
        
        # DEBUG: Print all available fields
        print("=" * 80)
        print(f"DEBUG: Initiative {initiative_key} fields available:")
        for field_key in init_fields.keys():
            field_value = init_fields.get(field_key)
            # Print field name and type, but not full content
            if field_value is not None:
                print(f"  - {field_key}: {type(field_value).__name__}")
        print("=" * 80)

        # DEBUG: Check custom fields that might contain Epic references
        print("DEBUG: Checking custom fields for Epic references...")
        for field_key, field_value in init_fields.items():
            if field_key.startswith("customfield_"):
                # Check if it's a list or dict that might contain issue keys
                if isinstance(field_value, list) and field_value:
                    # Print first item to see structure
                    sample = field_value[0] if field_value else None
                    if sample and isinstance(sample, dict):
                        if 'key' in sample or 'id' in sample:
                            print(f"  {field_key} contains: {field_value}")
                elif isinstance(field_value, dict) and field_value:
                    if 'key' in field_value or 'id' in field_value:
                        print(f"  {field_key} contains: {field_value}")
                elif isinstance(field_value, str) and ('UEX' in field_value or 'PFI' in field_value or 'Epic' in field_value):
                    print(f"  {field_key} = {field_value}")
        print("=" * 80)
        
        epic_issues = []
        seen_epics = set()
        
        # 2a. Check for issue links
        issue_links = init_fields.get("issuelinks", [])
        print(f"DEBUG: Found {len(issue_links)} issue links")
        
        for link in issue_links:
            # Links can be inward or outward
            linked_issue = link.get("outwardIssue") or link.get("inwardIssue")
            if linked_issue:
                linked_key = linked_issue.get("key")
                linked_type = linked_issue.get("fields", {}).get("issuetype", {}).get("name", "")
                link_type = link.get("type", {}).get("name", "")
                
                print(f"DEBUG: Issue link: {linked_key} ({linked_type}) via '{link_type}'")
                
                if "epic" in linked_type.lower() and linked_key not in seen_epics:
                    print(f"DEBUG: ✓ Found Epic via issue link: {linked_key}")
                    seen_epics.add(linked_key)
                    # Fetch full Epic details
                    try:
                        epic_full = self.get_issue(linked_key)
                        epic_issues.append(epic_full)
                    except Exception as e:
                        print(f"DEBUG: Could not fetch Epic {linked_key}: {e}")
        
        # 2b. Check subtasks field (some orgs structure it this way)
        subtasks = init_fields.get("subtasks", [])
        print(f"DEBUG: Found {len(subtasks)} subtasks")
        for subtask in subtasks:
            subtask_key = subtask.get("key")
            subtask_type = subtask.get("fields", {}).get("issuetype", {}).get("name", "")
            print(f"DEBUG: Subtask: {subtask_key} ({subtask_type})")
            if "epic" in subtask_type.lower() and subtask_key not in seen_epics:
                print(f"DEBUG: ✓ Found Epic as subtask: {subtask_key}")
                seen_epics.add(subtask_key)
                try:
                    epic_full = self.get_issue(subtask_key)
                    epic_issues.append(epic_full)
                except Exception as e:
                    print(f"DEBUG: Could not fetch Epic {subtask_key}: {e}")
        
        # 2c. Try Agile API
        try:
            url = f"{self.base_url}/rest/agile/1.0/issue/{initiative_key}/child"
            r = self.session.get(url, timeout=30)
            print(f"DEBUG: Agile API response status: {r.status_code}")
            if r.status_code == 200:
                agile_data = r.json()
                for issue in agile_data.get("issues", []):
                    epic_key = issue.get("key")
                    if epic_key and epic_key not in seen_epics:
                        print(f"DEBUG: ✓ Found Epic via Agile API: {epic_key}")
                        seen_epics.add(epic_key)
                        epic_issues.append(issue)
        except Exception as e:
            print(f"DEBUG: Agile API not available: {e}")
        
        # 2d. Try the Portfolio/Advanced Roadmaps hierarchy API
        try:
            url = f"{self.base_url}/rest/api/3/issue/{initiative_key}/hierarchy"
            r = self.session.get(url, timeout=30)
            print(f"DEBUG: Hierarchy API response status: {r.status_code}")
            if r.status_code == 200:
                hierarchy_data = r.json()
                print(f"DEBUG: Hierarchy data: {hierarchy_data}")
                
                # Navigate the hierarchy to find children
                for child in hierarchy_data.get("children", []):
                    child_key = child.get("key")
                    child_type = child.get("issueType", {}).get("name", "")
                    print(f"DEBUG: Hierarchy child: {child_key} ({child_type})")
                    
                    if child_key and child_key not in seen_epics:
                        print(f"DEBUG: ✓ Found Epic via Hierarchy API: {child_key}")
                        seen_epics.add(child_key)
                        try:
                            epic_full = self.get_issue(child_key)
                            epic_issues.append(epic_full)
                        except Exception as e:
                            print(f"DEBUG: Could not fetch Epic {child_key}: {e}")
        except Exception as e:
            print(f"DEBUG: Hierarchy API error: {e}")

        # 2e. Last resort: Find all Epics in related projects and check parent relationship
        try:
            # Try both the Initiative's project and common Epic projects
            project_keys = [
                initiative_key.split('-')[0],  # PFI
                'UEX',  # The Epics in the screenshot are UEX-*
                'PF',   # Common abbreviation
            ]
            
            for proj_key in project_keys:
                try:
                    jql = f'project = {proj_key} AND type = Epic ORDER BY created DESC'
                    print(f"DEBUG: Searching for Epics in project {proj_key}")
                    
                    all_epics = self.search_jql(jql, ["summary", "description", "parent", "customfield_10019"], max_results=200)
                    print(f"DEBUG: Found {len(all_epics)} Epics in project {proj_key}")
                    
                    for epic in all_epics:
                        epic_key = epic.get("key")
                        epic_fields = epic.get("fields", {})
                        
                        # Check if this Epic has a parent field pointing to our Initiative
                        parent = epic_fields.get("parent")
                        if parent:
                            if isinstance(parent, dict):
                                parent_key = parent.get("key", "")
                            elif isinstance(parent, str):
                                parent_key = parent
                            else:
                                parent_key = ""
                            
                            if parent_key == initiative_key and epic_key not in seen_epics:
                                print(f"DEBUG: ✓ Found Epic {epic_key} via parent field = {parent_key}")
                                seen_epics.add(epic_key)
                                epic_issues.append(epic)
                                continue
                        
                        # Check customfield_10019
                        cf_10019 = epic_fields.get("customfield_10019")
                        if cf_10019:
                            if cf_10019 == initiative_key and epic_key not in seen_epics:
                                print(f"DEBUG: ✓ Found Epic {epic_key} via customfield_10019 = {cf_10019}")
                                seen_epics.add(epic_key)
                                epic_issues.append(epic)
                    
                    # If we found epics in this project, no need to search other projects
                    if epic_issues:
                        break
                        
                except Exception as e:
                    print(f"DEBUG: Search in project {proj_key} failed: {e}")
                    continue
                    
        except Exception as e:
            print(f"DEBUG: Project search error: {e}")

        jql_variants = [
            f'parent = {initiative_key}',
            f'"Parent Link" = {initiative_key}',
            f'issue in childIssuesOf("{initiative_key}")',
            f'"Epic Link" = {initiative_key}',
            f'parentEpic = {initiative_key}',
            f'cf[10000] = {initiative_key}',
            f'"Portfolio Parent" = {initiative_key}',
            f'"Initiative Link" = {initiative_key}',
        ]
        
        print(f"DEBUG: Trying {len(jql_variants)} JQL variants...")
        for jql in jql_variants:
            try:
                print(f"DEBUG: JQL: {jql}")
                results = self.search_jql(jql, ["summary", "description", "issuetype"], max_results=100)
                print(f"DEBUG: -> Found {len(results)} results")
                
                for epic in results:
                    epic_key = epic.get("key")
                    epic_type = epic.get("fields", {}).get("issuetype", {}).get("name", "")
                    
                    if epic_key and epic_key not in seen_epics:
                        print(f"DEBUG: ✓ Found via JQL: {epic_key} ({epic_type})")
                        seen_epics.add(epic_key)
                        epic_issues.append(epic)
            except Exception as e:
                print(f"DEBUG: JQL failed: {e}")
                continue
        
        print("=" * 80)
        print(f"DEBUG: TOTAL Epics found: {len(epic_issues)}")
        print("=" * 80)
        
        # 3. For each Epic, get its children
        epics_with_children = []
        for epic in epic_issues:
            epic_key = epic.get("key", "")
            epic_fields = epic.get("fields", {})
            epic_summary = epic_fields.get("summary", "")
            epic_desc = epic_fields.get("description")
            
            # Parse Epic description
            if isinstance(epic_desc, dict) and epic_desc.get("type") == "doc":
                epic_desc_plain = adf_to_plaintext(epic_desc)
            elif isinstance(epic_desc, str):
                epic_desc_plain = epic_desc
            else:
                epic_desc_plain = ""
            
            # Clean for LLM (remove strikethrough and out-of-scope content)
            epic_desc_plain = clean_jira_text_for_llm(epic_desc_plain)
            
            # Get Epic's children
            try:
                children_issues = self.get_children_of_epic(epic_key)
                children = []
                for child in children_issues:
                    child_key = child.get("key", "")
                    child_fields = child.get("fields", {})
                    child_summary = child_fields.get("summary", "")
                    child_desc = child_fields.get("description")
                    
                    if isinstance(child_desc, dict) and child_desc.get("type") == "doc":
                        child_desc_plain = adf_to_plaintext(child_desc)
                    elif isinstance(child_desc, str):
                        child_desc_plain = child_desc
                    else:
                        child_desc_plain = ""
                    
                    # Clean for LLM
                    child_desc_plain = clean_jira_text_for_llm(child_desc_plain)
                    
                    children.append({
                        "key": child_key,
                        "summary": child_summary,
                        "desc": child_desc_plain
                    })
                    
                print(f"DEBUG: Epic {epic_key} has {len(children)} children")
            except Exception as e:
                print(f"DEBUG: Failed to get children for {epic_key}: {e}")
                children = []
            
            epics_with_children.append({
                "key": epic_key,
                "summary": epic_summary,
                "desc": epic_desc_plain,
                "children": children
            })
        
        return {
            "initiative": {
                "key": initiative_key,
                "summary": init_summary,
                "desc": init_desc_plain
            },
            "epics": epics_with_children
        }
    
    def load_manual_epics(self, epic_keys: List[str]) -> List[Dict]:
        """Load epics manually by their keys."""
        epics_with_children = []
        
        for epic_key in epic_keys:
            epic_key = epic_key.strip().upper()
            if not epic_key:
                continue
            
            try:
                print(f"DEBUG: Loading Epic {epic_key}")
                epic = self.get_issue(epic_key)
                epic_fields = epic.get("fields", {})
                epic_summary = epic_fields.get("summary", "")
                epic_desc = epic_fields.get("description")
                
                # Parse Epic description
                if isinstance(epic_desc, dict) and epic_desc.get("type") == "doc":
                    epic_desc_plain = adf_to_plaintext(epic_desc)
                elif isinstance(epic_desc, str):
                    epic_desc_plain = epic_desc
                else:
                    epic_desc_plain = ""
                
                # Get Epic's children
                try:
                    children_issues = self.get_children_of_epic(epic_key)
                    children = []
                    for child in children_issues:
                        child_key = child.get("key", "")
                        child_fields = child.get("fields", {})
                        child_summary = child_fields.get("summary", "")
                        child_desc = child_fields.get("description")
                        
                        if isinstance(child_desc, dict) and child_desc.get("type") == "doc":
                            child_desc_plain = adf_to_plaintext(child_desc)
                        elif isinstance(child_desc, str):
                            child_desc_plain = child_desc
                        else:
                            child_desc_plain = ""
                        
                        children.append({
                            "key": child_key,
                            "summary": child_summary,
                            "desc": child_desc_plain
                        })
                    
                    print(f"DEBUG: Epic {epic_key} has {len(children)} children")
                except Exception as e:
                    print(f"DEBUG: Failed to get children for {epic_key}: {e}")
                    children = []
                
                epics_with_children.append({
                    "key": epic_key,
                    "summary": epic_summary,
                    "desc": epic_desc_plain,
                    "children": children
                })
                
            except Exception as e:
                print(f"DEBUG: Failed to load Epic {epic_key}: {e}")
                continue
        
        return epics_with_children


# ---------- OpenAI LLM ----------
class LLM:
    def __init__(self, enabled: bool=True, model: Optional[str]=None):
        self.api_key = os.getenv("OPENAI_API_KEY")
        # Use gpt-4o-2024-08-06 for Structured Outputs support
        self.model = model or os.getenv("OPENAI_MODEL","gpt-4o-2024-08-06")
        self.enabled = enabled and bool(self.api_key)
        self.import_ok = True
        self.supports_structured_outputs = False
        if self.enabled:
            try:
                from openai import OpenAI  # noqa
                # Check if model supports structured outputs
                if "gpt-4o" in self.model or "gpt-4o-mini" in self.model:
                    self.supports_structured_outputs = True
            except Exception:
                self.import_ok = False; self.enabled = False
    
    def status_label(self) -> str:
        if not self.api_key: return "AI: OFF (no key)"
        if not self.import_ok: return "AI: OFF (install openai)"
        suffix = " [Structured]" if self.supports_structured_outputs else ""
        return f"AI: ON ({self.model}{suffix})" if self.enabled else "AI: OFF"
    
    def complete_json(self, sys_prompt: str, user_prompt: str, max_tokens=2000, retries=2, pydantic_model=None) -> Tuple[str, Optional[str]]:
        if not self.enabled: return ("","AI disabled or missing key/openai")
        try:
            from openai import OpenAI  # type: ignore
            client = OpenAI(api_key=self.api_key)
        except Exception as e:
            return ("", f"OpenAI import failed: {e}")
        
        last_err=None
        for attempt in range(retries+1):
            try:
                # Build base kwargs
                kwargs = dict(
                    model=self.model,
                    messages=[
                        {"role":"system","content":sys_prompt},
                        {"role":"user","content":user_prompt}
                    ]
                )
                
                # Use Structured Outputs if pydantic model provided and supported
                if pydantic_model and self.supports_structured_outputs and PYDANTIC_AVAILABLE:
                    print(f"DEBUG: Using Structured Outputs with {pydantic_model.__name__}")
                    # Use beta.chat.completions.parse for structured outputs
                    kwargs["response_format"] = pydantic_model
                    
                    # Use correct parameters based on model
                    if self.model.startswith(("o1",)):
                        kwargs["max_completion_tokens"] = max_tokens
                    else:
                        kwargs["max_tokens"] = max_tokens
                        kwargs["temperature"] = 0.0  # Maximum consistency
                        kwargs["seed"] = 12345  # Reproducibility
                    
                    resp = client.beta.chat.completions.parse(**kwargs)
                    
                    # Check for refusal
                    if resp.choices[0].message.refusal:
                        return ("", f"Model refused: {resp.choices[0].message.refusal}")
                    
                    # Get parsed response
                    parsed = resp.choices[0].message.parsed
                    if parsed:
                        # Convert Pydantic model to JSON string
                        return (parsed.model_dump_json(indent=2), None)
                    else:
                        return ("", "No parsed response from structured output")
                
                else:
                    # Fallback to regular JSON mode
                    print("DEBUG: Using regular JSON mode (structured outputs not available)")
                    kwargs["response_format"] = {"type":"json_object"}
                    
                    # Use correct parameters based on model
                    if self.model.startswith(("o1",)):
                        kwargs["max_completion_tokens"] = max_tokens
                    else:
                        kwargs["max_tokens"] = max_tokens
                        kwargs["temperature"] = 0.0
                        kwargs["seed"] = 12345
                        kwargs["top_p"] = 1.0
                    
                    resp = client.chat.completions.create(**kwargs)
                    return ((resp.choices[0].message.content or "").strip(), None)
                    
            except Exception as e:
                msg=str(e); last_err=msg
                print(f"DEBUG: API call failed (attempt {attempt+1}): {msg}")
                if "unsupported" in msg.lower() or "parse" in msg.lower():
                    try:
                        # Fallback: remove structured output and try basic JSON
                        kwargs["response_format"] = {"type":"json_object"}
                        if "max_completion_tokens" in kwargs:
                            kwargs["max_tokens"] = kwargs.pop("max_completion_tokens")
                        resp=client.chat.completions.create(**kwargs)
                        return ((resp.choices[0].message.content or "").strip(), None)
                    except Exception as e2:
                        last_err=str(e2)
                time.sleep(0.8*(attempt+1))
        return ("", last_err or "Unknown OpenAI error")

# ---------- Data Models ----------
# Add Pydantic for structured outputs
try:
    from pydantic import BaseModel, Field
    from typing import Literal
    PYDANTIC_AVAILABLE = True
except ImportError:
    PYDANTIC_AVAILABLE = False
    BaseModel = object  # Fallback

# Pydantic models for OpenAI Structured Outputs
if PYDANTIC_AVAILABLE:
    class TestStepSchema(BaseModel):
        action: str = Field(description="Clear, specific action to perform")
        expected: str = Field(description="Observable, verifiable expected result")
    
    class RequirementSchema(BaseModel):
        id: str = Field(description="Unique ID like REQ-001, REQ-002")
        description: str = Field(description="Clear, testable requirement description")
        source: str = Field(description="Source: 'Acceptance Criteria', 'Description', or 'User Story'")
    
    class TestCaseSchema(BaseModel):
        requirement_id: str = Field(description="Links to requirement ID (e.g., REQ-001)")
        requirement_desc: str = Field(description="Brief summary of the requirement")
        title: str = Field(description="Clear, descriptive test case title")
        priority: Literal[1, 2, 3, 4] = Field(description="1=Critical, 2=High, 3=Medium, 4=Low")
        test_type: Literal["Positive", "Negative", "Edge Case"] = Field(description="Test type classification")
        tags: List[str] = Field(default_factory=list, description="Relevant tags for categorization")
        steps: List[TestStepSchema] = Field(description="3-8 detailed test steps covering complete user journey")
    
    class TestGenerationResponse(BaseModel):
        requirements: List[RequirementSchema] = Field(description="All identified requirements from ticket")
        test_cases: List[TestCaseSchema] = Field(description="Exactly 3 test cases per requirement (1 Positive, 1 Negative, 1 Edge)")

    class TestCaseIssue(BaseModel):
        test_case_title: str = Field(description="Title of the test case with the issue")
        requirement_id: str = Field(description="Associated requirement ID")
        issue_type: Literal["missing_steps", "incomplete_flow", "unclear_expected", "wrong_test_type", "missing_requirement", "insufficient_requirements", "over_consolidated_requirements", "duplicate", "priority_mismatch", "other"]
        severity: Literal["critical", "major", "minor"] = Field(description="Severity of the issue")
        description: str = Field(description="Clear description of what's wrong")
        suggestion: str = Field(description="Specific suggestion for how to fix it")
    
    class CriticReviewResponse(BaseModel):
        overall_quality: Literal["excellent", "good", "needs_improvement", "poor"] = Field(description="Overall quality assessment")
        approved: bool = Field(description="True if test cases meet quality standards, False if refinement needed")
        confidence_score: int = Field(description="Confidence in assessment (0-100)", ge=0, le=100)
        
        # Validation checks
        requirement_count_correct: bool = Field(description="True if all requirements are properly identified")
        test_count_correct: bool = Field(description="True if exactly 3 test cases per requirement (formula: N×3)")
        test_type_distribution_correct: bool = Field(description="True if each requirement has 1 Positive, 1 Negative, 1 Edge")
        steps_complete: bool = Field(description="True if all test cases have 3-8 detailed steps")
        traceability_correct: bool = Field(description="True if all test cases properly link to requirements")
        
        # Detailed findings
        strengths: List[str] = Field(description="What was done well (2-4 points)")
        issues_found: List[TestCaseIssue] = Field(description="Specific issues that need correction")
        missing_test_scenarios: List[str] = Field(description="Important scenarios that should be tested but aren't covered")
        
        # Summary
        summary: str = Field(description="2-3 sentence summary of the review")
        recommendation: str = Field(description="Clear recommendation: 'Approve' or specific actions needed")
    
    # New models for test ticket creation
    class TestTicketSplit(BaseModel):
        functional_area: str = Field(description="Name of the functional area (e.g., 'Fleet Data Migration')")
        child_tickets: List[str] = Field(description="Child ticket keys that relate to this area")
        estimated_ac_count: int = Field(description="Estimated number of acceptance criteria (5-8 recommended)")
        rationale: str = Field(description="Why these tickets are grouped together")
    
    class TestTicketSplitResponse(BaseModel):
        recommended_splits: List[TestTicketSplit] = Field(description="Recommended test ticket structure")
        total_tickets: int = Field(description="Total number of test tickets to create")
        reasoning: str = Field(description="Overall strategy for splitting")
    
    class TestTicketContent(BaseModel):
        summary: str = Field(description="Ticket summary following Epic naming convention")
        description: str = Field(description="Detailed description matching Epic's writing style")
        acceptance_criteria: List[str] = Field(description="5-8 black-box acceptance criteria for manual testing")
        quality_estimate: int = Field(description="Self-estimated quality score (0-100)", ge=0, le=100)
    
    class TestTicketReview(BaseModel):
        approved: bool = Field(description="True if ticket is ready, False if needs revision")
        quality_score: int = Field(description="Quality assessment score (0-100)", ge=0, le=100)
        strengths: List[str] = Field(description="What was done well")
        issues: List[str] = Field(description="Problems that need fixing")
        recommendations: List[str] = Field(description="Specific improvements needed")
        revised_content: Optional[TestTicketContent] = Field(default=None, description="Revised ticket if approved=False")

# Dataclasses for internal use
@dataclass
class TestStep:
    action: str
    expected: str

@dataclass
class Requirement:
    id: str  # REQ-001, AC-001, etc.
    description: str
    source: str  # "Acceptance Criteria", "Description", "User Story", etc.

@dataclass
class TestCase:
    title: str
    objective: str = ""
    priority: int = 2
    test_type: str = "Positive"  # Positive, Negative, or Edge Case
    tags: List[str] = field(default_factory=list)
    steps: List[TestStep] = field(default_factory=list)
    requirement_id: str = ""  # Links to Requirement.id
    requirement_desc: str = ""  # For display purposes

@dataclass
class GeneratedTestTicket:
    """Stores a generated test ticket with its state."""
    id: int
    title: str
    summary: str
    description: str
    acceptance_criteria: List[str]
    quality_score: int
    ac_count: int
    analyzed: bool = False
    test_cases: Optional[List[TestCase]] = None
    requirements: Optional[List[Dict]] = None
    export_timestamp: Optional[str] = None
    child_tickets: List[Dict] = field(default_factory=list)  # Source child tickets with key and summary

# ---------- Analyzer ----------
class Analyzer:
    def __init__(self, llm: LLM):
        self.llm = llm

    def extract_sections(self, description: str) -> Tuple[str, List[str]]:
        parts = re.split(r"\n\s*\n", description or "")
        ac_texts=[]; body=[]
        for block in parts:
            if "acceptance" in block.lower() or re.search(r"^(given|when|then|and|but)\b", block.strip(), re.I|re.M):
                ac_texts.append(block.strip())
            else:
                body.append(block.strip())
        return ("\n\n".join(body).strip(), ac_texts)

    @staticmethod
    def _fallback_cases(summary: str, body: str, ac_blocks: List[str]) -> List[TestCase]:
        text=" ".join([summary, body, "\n".join(ac_blocks)]).strip()
        sentences=[s.strip() for s in re.split(r"[.\n]", text) if s.strip()]
        tc=[]
        
        # Generate positive, negative, and edge cases for first few sentences
        test_types = ["Positive", "Negative", "Edge Case"]
        for i,s in enumerate(sentences[:3],1):  # Take first 3 requirements
            for test_type in test_types:
                if test_type == "Positive":
                    title = f"Verify: {s[:70]}"
                    action = f"Test valid scenario: {s[:60]}"
                elif test_type == "Negative":
                    title = f"Verify invalid input handling: {s[:50]}"
                    action = f"Test with invalid/missing data for: {s[:50]}"
                else:  # Edge Case
                    title = f"Verify boundary condition: {s[:50]}"
                    action = f"Test edge case for: {s[:50]}"
                
                tc.append(TestCase(
                    title=title.rstrip("."), 
                    priority=2, 
                    test_type=test_type,
                    tags=["fallback", test_type],
                    steps=[TestStep(action, "Behavior matches requirement.")]
                ))
        
        if not tc: 
            tc=[TestCase(
                title=summary or "Review ticket", 
                test_type="Positive",
                steps=[TestStep("Walk through flow","Meets acceptance criteria")]
            )]
        return tc

    @staticmethod
    def _system_prompt() -> str:
        return (
            "You are a senior QA engineer creating comprehensive manual black-box test cases.\n\n"
            "TESTING PHILOSOPHY:\n"
            "For EACH requirement identified, create exactly THREE test cases:\n"
            "1. One POSITIVE test (happy path)\n"
            "2. One NEGATIVE test (error handling)\n"
            "3. One EDGE CASE test (boundary conditions)\n\n"
            "This ensures complete coverage with clear traceability.\n\n"
            
            "CRITICAL: You MUST follow this two-phase structured approach:\n\n"
            
            "PHASE 1 - REQUIREMENT EXTRACTION:\n"
            "First, identify and extract ALL distinct requirements from the ticket:\n"
            "- Extract from Acceptance Criteria (highest priority)\n"
            "- Extract from Description (look for behaviors, validations, UI elements)\n"
            "- Extract from Epic and child ticket context\n"
            "- Each requirement must be clear, atomic, and testable\n"
            "- Assign each a unique ID (REQ-001, REQ-002, etc.)\n"
            "- BE THOROUGH: Identify all testable behaviors\n\n"
            
            "EXAMPLES OF REQUIREMENTS:\n"
            "✓ Form field validation rules (each field = 1 requirement)\n"
            "✓ Button states and behaviors\n"
            "✓ Navigation and UI elements\n"
            "✓ Data persistence/session management\n"
            "✓ Integration points\n"
            "✓ Business logic rules\n\n"
            
            "PHASE 2 - TEST CASE GENERATION:\n"
            "After identifying ALL requirements, create test cases using the 1:3 FORMULA:\n\n"
            
            "CRITICAL FORMULA:\n"
            "For N requirements identified → Generate exactly N × 3 test cases\n"
            "Example: 10 requirements → 30 test cases (10 Positive + 10 Negative + 10 Edge Cases)\n"
            "Example: 13 requirements → 39 test cases (13 Positive + 13 Negative + 13 Edge Cases)\n\n"
            
            "FOR EACH REQUIREMENT, CREATE THESE THREE TEST CASES:\n\n"
            
            "1. **POSITIVE TEST** (Happy Path):\n"
            "   - Test the requirement with valid inputs and expected behavior\n"
            "   - Verify successful workflow\n"
            "   - Priority: Usually High (1) or Medium (2)\n"
            "   - test_type: 'Positive'\n"
            "   - Steps: 1-6 steps depending on complexity\n"
            "   Example: 'REQ-001 Positive: Valid email format accepted'\n\n"
            
            "2. **NEGATIVE TEST** (Error Handling):\n"
            "   - Test the requirement with invalid inputs\n"
            "   - Verify proper error handling and messages\n"
            "   - Priority: Usually Medium (2) or High (1) for critical validations\n"
            "   - test_type: 'Negative'\n"
            "   - Steps: 1-6 steps depending on complexity\n"
            "   Example: 'REQ-001 Negative: Invalid email format rejected'\n\n"
            
            "3. **EDGE CASE TEST** (Boundary Conditions):\n"
            "   - Test the requirement at boundaries and limits\n"
            "   - Verify behavior with special characters, empty values, max length, etc.\n"
            "   - Priority: Usually Medium (2) or Low (3)\n"
            "   - test_type: 'Edge Case'\n"
            "   - Steps: 1-6 steps depending on complexity\n"
            "   Example: 'REQ-001 Edge Case: Email with special characters'\n\n"
            
            "TEST CASE NAMING CONVENTION:\n"
            "Use this format: '[REQ-ID] [Type]: [Description]'\n"
            "✓ 'REQ-001 Positive: Complete form with valid personal data'\n"
            "✓ 'REQ-001 Negative: Submit form with missing first name'\n"
            "✓ 'REQ-001 Edge Case: First name with maximum length (50 chars)'\n\n"
            
            "STEP COUNT GUIDELINES:\n"
            "- Simple validation: 1-3 steps\n"
            "- Moderate workflow: 3-4 steps\n"
            "- Complex scenario: 4-6 steps\n"
            "- Keep steps focused and specific to the requirement\n"
            "- Use context from Epic and other tickets to inform realistic test data\n\n"
            
            "STEP COMPLETENESS:\n"
            "Each step MUST have:\n"
            "- action: Clear, specific action to perform\n"
            "- expected: Observable, verifiable expected result\n\n"
            
            "EXAMPLE - Testing a Form Field:\n"
            "REQ-005: First name field is required and accepts 1-50 characters\n\n"
            
            "TC-001: REQ-005 Positive: Valid first name accepted\n"
            "  Steps:\n"
            "  1. Navigate to registration form → Form displays\n"
            "  2. Enter first name: 'John' → Field accepts input\n"
            "  3. Click Submit → Form validates successfully\n\n"
            
            "TC-002: REQ-005 Negative: Missing first name shows error\n"
            "  Steps:\n"
            "  1. Navigate to registration form → Form displays\n"
            "  2. Leave first name field empty → Field is blank\n"
            "  3. Click Submit → Error message: 'First name is required'\n\n"
            
            "TC-003: REQ-005 Edge Case: First name at maximum length\n"
            "  Steps:\n"
            "  1. Navigate to registration form → Form displays\n"
            "  2. Enter first name with 50 characters → Field accepts exactly 50 chars\n"
            "  3. Attempt to enter 51st character → Character is rejected\n"
            "  4. Click Submit → Form validates successfully with 50 char name\n\n"
            
            "REQUIRED JSON OUTPUT STRUCTURE:\n"
            "{\n"
            '  "requirements": [\n'
            '    {"id": "REQ-001", "description": "requirement text", "source": "Acceptance Criteria"}\n'
            "  ],\n"
            '  "test_cases": [\n'
            "    {\n"
            '      "requirement_id": "REQ-001",  // REQUIRED: Single requirement ID this test covers\n'
            '      "requirement_desc": "Tests first name field validation",  // REQUIRED: What requirement is being tested\n'
            '      "title": "REQ-001 Positive: Valid first name accepted",  // REQUIRED: Use naming convention\n'
            '      "priority": 1,  // REQUIRED: 1=High, 2=Medium, 3=Low\n'
            '      "test_type": "Positive",  // REQUIRED: "Positive", "Negative", or "Edge Case"\n'
            '      "tags": ["form", "validation", "first-name"],  // REQUIRED: Array of tags\n'
            '      "steps": [  // REQUIRED: 1-6 steps with action and expected\n'
            '        {"action": "Navigate to registration form", "expected": "Form displays"},\n'
            '        {"action": "Enter first name: John", "expected": "Field accepts input"},\n'
            '        {"action": "Click Submit", "expected": "Form validates successfully"}\n'
            "      ]\n"
            "    }\n"
            "  ]\n"
            "}\n\n"
            
            "MANDATORY RULES:\n"
            "1. Identify ALL requirements first (PHASE 1)\n"
            "2. For EACH requirement, create EXACTLY 3 test cases (1 Positive, 1 Negative, 1 Edge Case)\n"
            "3. Formula: N requirements → N × 3 test cases (this will be verified!)\n"
            "4. Each test case MUST include ALL required fields:\n"
            "   - requirement_id (single REQ-ID)\n"
            "   - requirement_desc\n"
            "   - title (using naming convention)\n"
            "   - priority\n"
            "   - test_type (exactly: 'Positive', 'Negative', or 'Edge Case')\n"
            "   - tags\n"
            "   - steps (1-6 steps with action and expected)\n"
            "5. Use Epic and child ticket context to inform test scenarios\n"
            "6. Keep steps focused - simple tests may only need 1-3 steps\n"
            "7. Test cases missing ANY required field will be REJECTED\n\n"
            
            "VALIDATION BEFORE OUTPUT:\n"
            "Before generating output, verify:\n"
            "✓ All requirements are identified\n"
            "✓ Each requirement has exactly 3 test cases (Positive, Negative, Edge Case)\n"
            "✓ Total test cases = Requirements count × 3\n"
            "✓ Every test case has ALL required fields\n"
            "✓ Step count is appropriate (1-6 per test case)\n"
            "✓ Test case titles follow the naming convention\n"
        )

    @staticmethod
    def _summarize_feature_context(feature_ctx: Optional[Dict]) -> str:
        if not feature_ctx: 
            print("DEBUG: No feature context available for test generation")
            return "(none)"
        
        print("=" * 80)
        print("DEBUG: Feature Context Summary Being Sent to AI:")
        print(f"  Type: {feature_ctx.get('type', 'unknown')}")
        
        lines = []
        
        if feature_ctx.get("type") == "initiative":
            # Initiative context
            print(f"  Initiative: {feature_ctx.get('initiative_key')}")
            print(f"  Epic count: {feature_ctx.get('epic_count', 0)}")
            print(f"  Total child issues: {feature_ctx.get('total_children', 0)}")
            print(f"  Sample epics: {feature_ctx.get('epics', [])[:3]}")
            
            lines.append(f"Initiative: {feature_ctx.get('initiative_key', '')} – {feature_ctx.get('initiative_summary', '')}")
            if feature_ctx.get('initiative_desc'):
                lines.append("Initiative overview: " + feature_ctx['initiative_desc'][:800])
            
            lines.append(f"\nEpics ({feature_ctx.get('epic_count', 0)}):")
            for epic_summary in (feature_ctx.get('epics', [])[:20]):  # Cap at 20 epics
                lines.append(f"  - {epic_summary}")
            
            lines.append(f"\nTotal child issues across all epics: {feature_ctx.get('total_children', 0)}")
            lines.append("\nSample child issues:")
            children = feature_ctx.get("children", [])[:30]  # Cap at 30
            for ch in children:
                s = f"- {ch.get('key')}: {ch.get('summary','')}"
                if ch.get("desc"):
                    s += " | " + ch["desc"][:300]
                lines.append(s)
        else:
            # Epic context (original logic)
            print(f"  Epic: {feature_ctx.get('epic_key')}")
            print(f"  Child count: {len(feature_ctx.get('children', []))}")
            
            lines.append(f"Epic: {feature_ctx.get('epic_key','')} – {feature_ctx.get('epic_summary','')}")
            epic_desc = (feature_ctx.get('epic_desc','') or "")
            if epic_desc:
                lines.append("Epic overview: " + epic_desc[:1200])
            children = feature_ctx.get("children", [])[:30]
            for ch in children:
                s = f"- {ch.get('key')}: {ch.get('summary','')}"
                if ch.get("desc"):
                    s += " | " + ch["desc"][:400]
                lines.append(s)
        
        summary = "\n".join(lines)
        print(f"  Total context length: {len(summary)} characters")
        print("=" * 80)
        
        return summary

    def analyze_images_with_vision(self, attachments: List[Dict]) -> str:
        """
        Analyze image attachments using GPT-4o vision API.
        Returns a text description of what's in the images.
        """
        if not self.llm.enabled:
            return ""
        
        try:
            from openai import OpenAI
            client = OpenAI(api_key=self.llm.api_key)
        except Exception as e:
            print(f"DEBUG: OpenAI import failed: {e}")
            return ""
        
        # Filter for image attachments
        images = [att for att in attachments if att.get("type") == "image"]
        if not images:
            return ""
        
        print(f"DEBUG: Analyzing {len(images)} images with GPT-4o vision...")
        
        # Build the vision request
        content = [
            {
                "type": "text",
                "text": (
                    "You are analyzing images attached to a Jira ticket for software testing purposes. "
                    "For each image, describe:\n"
                    "1. What UI elements, screens, or diagrams are shown\n"
                    "2. Any text, labels, buttons, or form fields visible\n"
                    "3. Workflows, user interactions, or processes illustrated\n"
                    "4. Error messages, validation rules, or business logic\n"
                    "5. Any technical details relevant for test case creation\n\n"
                    "Be specific and detailed. This information will be used to create test cases."
                )
            }
        ]
        
        # Add all images
        for i, img in enumerate(images, 1):
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": img.get("data_url"),
                    "detail": "high"  # Use high detail for better analysis
                }
            })
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o",  # Use gpt-4o for vision capabilities
                messages=[
                    {
                        "role": "user",
                        "content": content
                    }
                ],
                max_tokens=2000
            )
            
            analysis = response.choices[0].message.content
            print(f"DEBUG: Image analysis complete ({len(analysis)} chars)")
            return analysis
        
        except Exception as e:
            print(f"DEBUG: Error analyzing images: {e}")
            return ""
    
    @staticmethod
    def _format_attachments(attachments: Optional[List[Dict]]) -> str:
        """Format attachments for inclusion in the prompt."""
        if not attachments:
            return ""
        
        lines = []
        
        # Separate documents and images
        documents = [att for att in attachments if att.get("type") == "document"]
        images = [att for att in attachments if att.get("type") == "image"]
        
        # Add documents
        if documents:
            lines.append("📄 Documents:")
            for doc in documents:
                filename = doc.get("filename", "unknown")
                content = doc.get("content", "")
                # Truncate long documents
                if len(content) > 3000:
                    content = content[:3000] + "\n\n...[truncated for length]"
                lines.append(f"\n--- {filename} ---\n{content}\n")
        
        # Add image analysis
        if images:
            lines.append("\n🖼️ Images:")
            for img in images:
                filename = img.get("filename", "unknown")
                lines.append(f"- {filename}")
            lines.append("\n(Image content has been analyzed and will inform test case generation)")
        
        return "\n".join(lines)

    @staticmethod
    def _user_prompt(summary: str, body: str, ac_blocks: List[str], feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None) -> str:
        return (
            "🚨 CRITICAL INSTRUCTION 🚨\n"
            "You MUST generate ALL requirements AND ALL test cases in this SINGLE response.\n"
            "Do NOT stop after generating just 1 test case. Generate EVERY test case for EVERY requirement.\n"
            "The 'test_cases' array in your JSON response must contain ALL test cases, not just the first one.\n\n"
            "==================================================================================\n\n"
            f"Story Summary:\n{summary}\n\n"
            f"Story Description:\n{body}\n\n"
            f"Acceptance Criteria Blocks:\n{chr(10).join(ac_blocks) if ac_blocks else '(none)'}\n\n"
            + (f"Attachments Analysis:\n{Analyzer._format_attachments(attachments)}\n\n" if attachments else "")
            + "YOUR TASK:\n\n"
            "STEP 1 - REQUIREMENT IDENTIFICATION:\n"
            "Carefully analyze the ticket above and extract ALL distinct, testable requirements.\n"
            "Look in:\n"
            "- Acceptance Criteria (primary source)\n"
            "- Description (secondary source)\n"
            "- Summary (for high-level understanding)\n\n"
            
            "⚠️ SPECIAL NOTE FOR GIVEN-WHEN-THEN FORMAT:\n"
            "If acceptance criteria use 'Given-When-Then' format, treat EACH numbered item as a separate requirement.\n"
            "Do NOT consolidate similar-looking items - each deserves its own 3 test cases.\n"
            "Example: If you see 7 'Given X, when Y, then Z' statements, that's 7 requirements = 21 test cases.\n\n"
            
            "Each requirement should be:\n"
            "- Atomic (tests one specific functionality)\n"
            "- Clear and unambiguous\n"
            "- Testable (has observable outcomes)\n"
            "Assign each a unique ID starting from REQ-001.\n\n"
            
            "STEP 2 - TEST CASE GENERATION:\n"
            "🚨 CRITICAL: For EACH requirement you identified, you MUST create EXACTLY THREE test cases 🚨\n"
            "Generate ALL test cases for ALL requirements in this single response.\n"
            "Do NOT stop after 1 test case - continue until you have generated test cases for ALL requirements.\n\n"
            
            "For each requirement, create these 3 test cases:\n\n"
            
            "1. POSITIVE TEST (Happy Path):\n"
            "   Title: 'REQ-XXX Positive: [description]'\n"
            "   - Valid inputs and expected successful behavior\n"
            "   - 1-6 steps depending on complexity\n"
            "   - Priority: High (1) or Medium (2)\n\n"
            
            "2. NEGATIVE TEST (Error Handling):\n"
            "   Title: 'REQ-XXX Negative: [description]'\n"
            "   - Invalid inputs or missing data\n"
            "   - Verify proper error handling\n"
            "   - 1-6 steps depending on complexity\n"
            "   - Priority: Medium (2) or High (1) for critical validations\n\n"
            
            "3. EDGE CASE TEST (Boundary Conditions):\n"
            "   Title: 'REQ-XXX Edge Case: [description]'\n"
            "   - Boundary values, special characters, limits\n"
            "   - 1-6 steps depending on complexity\n"
            "   - Priority: Medium (2) or Low (3)\n\n"
            
            "🚨 CRITICAL FORMULA CHECK 🚨\n"
            "If you identify N requirements, you MUST generate N × 3 test cases IN THIS RESPONSE.\n"
            "Example: 7 requirements → 21 test cases (7 Positive + 7 Negative + 7 Edge)\n"
            "Example: 10 requirements → 30 test cases (10 Positive + 10 Negative + 10 Edge)\n"
            "Example: 13 requirements → 39 test cases (13 Positive + 13 Negative + 13 Edge)\n\n"
            
            "🚨 IMPORTANT: Even if requirements have similar Given-When-Then patterns, you must generate ALL test cases 🚨\n"
            "Do NOT assume similarity means you can skip test cases. Generate each one explicitly.\n\n"
            
            "CONCRETE EXAMPLE - If you have 3 requirements, your JSON must look like:\n"
            "{\n"
            '  "requirements": [\n'
            '    {"id": "REQ-001", "description": "...", "source": "..."},\n'
            '    {"id": "REQ-002", "description": "...", "source": "..."},\n'
            '    {"id": "REQ-003", "description": "...", "source": "..."}\n'
            "  ],\n"
            '  "test_cases": [\n'
            '    {"requirement_id": "REQ-001", "title": "REQ-001 Positive: ...", "test_type": "Positive", ...},\n'
            '    {"requirement_id": "REQ-001", "title": "REQ-001 Negative: ...", "test_type": "Negative", ...},\n'
            '    {"requirement_id": "REQ-001", "title": "REQ-001 Edge Case: ...", "test_type": "Edge Case", ...},\n'
            '    {"requirement_id": "REQ-002", "title": "REQ-002 Positive: ...", "test_type": "Positive", ...},\n'
            '    {"requirement_id": "REQ-002", "title": "REQ-002 Negative: ...", "test_type": "Negative", ...},\n'
            '    {"requirement_id": "REQ-002", "title": "REQ-002 Edge Case: ...", "test_type": "Edge Case", ...},\n'
            '    {"requirement_id": "REQ-003", "title": "REQ-003 Positive: ...", "test_type": "Positive", ...},\n'
            '    {"requirement_id": "REQ-003", "title": "REQ-003 Negative: ...", "test_type": "Negative", ...},\n'
            '    {"requirement_id": "REQ-003", "title": "REQ-003 Edge Case: ...", "test_type": "Edge Case", ...}\n'
            "  ]\n"
            "}\n"
            "Notice: 3 requirements = 9 test cases in the array (3 × 3 = 9).\n\n"
            
            "The 'test_cases' array in your JSON response MUST contain ALL test cases.\n"
            "Do not generate just 1 test case and stop - generate ALL of them!\n\n"
            
            "STEP GUIDELINES:\n"
            "- Simple field validation: 1-3 steps\n"
            "- Moderate workflows: 3-4 steps\n"
            "- Complex scenarios: 4-6 steps\n"
            "- Use realistic test data based on the ticket description\n"
            "- Each step needs clear action + expected result\n\n"
            
            "MANDATORY FIELDS FOR EACH TEST CASE:\n"
            "Every test case MUST include:\n"
            "- requirement_id: Single requirement ID (e.g., 'REQ-001')\n"
            "- requirement_desc: What you're testing\n"
            "- title: Following naming convention 'REQ-XXX [Type]: Description'\n"
            "- priority: 1, 2, or 3\n"
            "- test_type: Exactly 'Positive', 'Negative', or 'Edge Case'\n"
            "- tags: Array of relevant tags\n"
            "- steps: 1-6 steps with action and expected\n\n"
            
            "OUTPUT FORMAT:\n"
            "Return valid JSON with 'requirements' array and 'test_cases' array.\n"
            "Each test case must have ALL the fields listed above.\n"
            "The 'test_cases' array must contain ALL test cases for ALL requirements.\n\n"
            
            "BEFORE YOU OUTPUT - SELF-CHECK:\n"
            "Count your requirements and test cases:\n"
            "- Requirements count: ___ (e.g., 7)\n"
            "- Test cases count: ___ (must be requirements × 3, e.g., 21)\n"
            "- Verify each requirement has exactly 3 test cases (1 Positive, 1 Negative, 1 Edge Case)\n"
            "- Verify your 'test_cases' array has ALL test cases, not just 1\n"
        )
    
    
    
    @staticmethod
    def _refinement_prompt(
        summary: str, 
        body: str, 
        ac_blocks: List[str], 
        feature_ctx: Optional[Dict],
        previous_attempt: str,
        critic_feedback: Dict,
        iteration: int
    ) -> str:
        """Generate a refinement prompt based on critic feedback."""
        
        prompt = (
            f"===== REFINEMENT ATTEMPT #{iteration} =====\n\n"
            f"Your PREVIOUS ATTEMPT was REJECTED by the QA Manager.\n"
            f"Study the feedback below carefully and generate CORRECTED test cases.\n\n"
            f"=== CRITIC REVIEW FEEDBACK ===\n"
            f"Overall Quality: {critic_feedback.get('overall_quality', 'unknown')}\n"
            f"Approved: {critic_feedback.get('approved', False)}\n"
            f"Confidence: {critic_feedback.get('confidence_score', 0)}%\n\n"
        )
        
        # Add validation failures
        prompt += "=== VALIDATION FAILURES ===\n"
        if not critic_feedback.get('requirement_count_correct'):
            prompt += "✗ CRITICAL: Requirement count is incorrect\n"
        if not critic_feedback.get('test_count_correct'):
            prompt += "✗ CRITICAL: Formula is WRONG - Must be exactly Requirements × 3\n"
        if not critic_feedback.get('test_type_distribution_correct'):
            prompt += "✗ CRITICAL: Test type distribution is wrong - each requirement needs exactly 1 Positive, 1 Negative, 1 Edge Case\n"
        if not critic_feedback.get('steps_complete'):
            prompt += "✗ Some test cases have step count issues (should be 1-6 steps)\n"
        if not critic_feedback.get('traceability_correct'):
            prompt += "✗ CRITICAL: Traceability is broken - requirements don't have their 3 test cases\n"
        
        prompt += "\n"
        
        # Add specific issues
        issues = critic_feedback.get('issues_found', [])
        if issues:
            prompt += f"=== SPECIFIC ISSUES ({len(issues)}) ===\n"
            for i, issue in enumerate(issues, 1):
                prompt += (
                    f"{i}. [{issue.get('severity', 'unknown').upper()}] {issue.get('test_case_title', 'N/A')}\n"
                    f"   Problem: {issue.get('description', '')}\n"
                    f"   Fix: {issue.get('suggestion', '')}\n\n"
                )
        
        # Add missing scenarios
        missing = critic_feedback.get('missing_test_scenarios', [])
        if missing:
            prompt += f"=== MISSING TEST SCENARIOS ({len(missing)}) ===\n"
            for scenario in missing:
                prompt += f"- {scenario}\n"
            prompt += "\n"
        
        # Add critic's recommendation and summary
        prompt += f"=== CRITIC'S RECOMMENDATION ===\n{critic_feedback.get('recommendation', 'Fix the issues above')}\n\n"
        prompt += f"=== CRITIC'S SUMMARY ===\n{critic_feedback.get('summary', 'Quality needs improvement')}\n\n"
        
        prompt += (
            "===== YOUR TASK =====\n\n"
            "🚨 CRITICAL: Generate ALL requirements AND ALL test cases in this SINGLE response 🚨\n"
            "Do NOT generate just 1 test case and stop. Generate EVERY test case for EVERY requirement.\n"
            "The 'test_cases' array must contain ALL test cases (Requirements × 3).\n\n"
            "Re-analyze the original ticket and generate CORRECTED test cases.\n\n"
            
            f"ORIGINAL TICKET:\n"
            f"Summary: {summary}\n"
            f"Description: {body[:2000]}\n"
            f"Acceptance Criteria: {chr(10).join(ac_blocks) if ac_blocks else '(none)'}\n\n"
            
            "===== MANDATORY APPROACH =====\n\n"
            
            "STEP 1: Extract ALL requirements\n"
            "- Review ticket thoroughly\n"
            "- Identify each distinct, testable requirement\n"
            "- ⚠️ If using Given-When-Then format: EACH item = separate requirement\n"
            "- Assign REQ-001, REQ-002, etc.\n"
            "- Count them: ___ requirements identified\n\n"
            
            "STEP 2: Apply the 1:3 FORMULA - Generate ALL test cases in this response\n"
            "🚨 CRITICAL: For EACH requirement, create EXACTLY 3 test cases in THIS RESPONSE:\n"
            "  1. REQ-XXX Positive: [happy path test]\n"
            "  2. REQ-XXX Negative: [error handling test]\n"
            "  3. REQ-XXX Edge Case: [boundary test]\n"
            "- Total test cases MUST equal: Requirements × 3\n"
            "- Example: 7 requirements → 21 test cases (ALL in this response)\n"
            "- Example: 10 requirements → 30 test cases (ALL in this response)\n"
            "- Example: 13 requirements → 39 test cases (ALL in this response)\n"
            "- Do NOT stop after generating 1 test case - continue until ALL are generated\n"
            "- Do NOT assume similarity means you can skip - generate each explicitly\n\n"
            
            "REMINDER: Your test_cases array must look like:\n"
            '"test_cases": [TC for REQ-001 Pos, TC for REQ-001 Neg, TC for REQ-001 Edge, TC for REQ-002 Pos, ...]\n'
            "Not just: [TC for REQ-001 Pos] ← This is WRONG\n\n"
            
            "STEP 3: Ensure EVERY test case has ALL fields:\n"
            "- requirement_id (single REQ-ID)\n"
            "- requirement_desc\n"
            "- title (using 'REQ-XXX [Type]: Description' format)\n"
            "- priority (1, 2, or 3)\n"
            "- test_type ('Positive', 'Negative', or 'Edge Case')\n"
            "- tags (array)\n"
            "- steps (1-6 steps with action and expected)\n\n"
            
            "STEP 4: Validate before output:\n"
            "- Count requirements: ___\n"
            "- Count test cases: ___ (must be requirements × 3)\n"
            "- Verify your 'test_cases' array contains ALL test cases, not just 1\n"
            "- Check each requirement has exactly 3 test cases\n"
            "- Check every test case has all required fields\n\n"
            
            "===== CRITICAL REMINDERS =====\n"
            "✓ The 1:3 formula is NON-NEGOTIABLE: N requirements → N × 3 test cases\n"
            "✓ Each requirement gets exactly 1 Positive, 1 Negative, 1 Edge Case\n"
            "✓ Test case titles MUST follow: 'REQ-XXX [Type]: Description'\n"
            "✓ Every test case MUST have all required fields and 1-6 steps\n"
            "✓ Step counts should be appropriate: 1-6 steps per test case\n"
            "✓ Generate ALL test cases in this single response - do not stop after 1\n"
            "✓ The 'test_cases' array MUST contain ALL test cases for ALL requirements\n\n"
            
            "Generate the corrected output now with ALL test cases. Double-check the formula before submitting!"
        )
        
        return prompt
    
    def generate_feature_overview(self, feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None) -> Tuple[str, Optional[str]]:

        """Generate an AI overview of the entire feature based on Epic + child issues."""
        if not feature_ctx:
            return ("No feature context available.", None)
        
        # Analyze images if present
        image_analysis = ""
        if attachments:
            images = [att for att in attachments if att.get("type") == "image"]
            if images:
                print(f"DEBUG: Analyzing {len(images)} images for feature overview...")
                image_analysis = self.analyze_images_with_vision(images)
        
        sys_prompt = (
            "You are a senior product analyst and technical writer. Analyze the provided Epic and its child issues "
            "to generate a comprehensive, detailed overview of what this feature does, how it works, and its value.\n"
            "Output ONLY JSON with this structure:\n"
            '{ "overview": string, "problem_statement": string, "solution_approach": string, '
            '"key_capabilities": [string], "user_impact": string, "technical_considerations": [string], '
            '"business_value": string }'
        )
        
        user_prompt = (
            f"Epic: {feature_ctx.get('epic_key', '')} - {feature_ctx.get('epic_summary', '')}\n\n"
            f"Epic Description:\n{feature_ctx.get('epic_desc', '')[:3000]}\n\n"
            f"Child Issues ({len(feature_ctx.get('children', []))}):\n"
        )
        
        for ch in feature_ctx.get("children", [])[:25]:  # Increased to 25 children
            user_prompt += f"\n- {ch.get('key')}: {ch.get('summary')}"
            if ch.get('desc'):
                user_prompt += f"\n  {ch['desc'][:500]}"  # More description per child
        
        # Add attachment context
        if attachments:
            user_prompt += f"\n\n{Analyzer._format_attachments(attachments)}"
        
        if image_analysis:
            user_prompt += f"\n\n{image_analysis}"
        
        user_prompt += (
            "\n\nProvide a detailed analysis:\n"
            "1. Overview: 3-4 sentences explaining what this feature does at a high level\n"
            "2. Problem Statement: 2-3 sentences describing the problem this feature solves\n"
            "3. Solution Approach: 3-4 sentences explaining how the feature addresses the problem\n"
            "4. Key Capabilities: 5-8 specific capabilities or functionalities this feature provides\n"
            "5. User Impact: 2-3 sentences on how this affects end users\n"
            "6. Technical Considerations: 3-5 important technical aspects, integrations, or dependencies\n"
            "7. Business Value: 2-3 sentences on the business impact and value delivered"
        )
        
        raw, ai_err = self.llm.complete_json(sys_prompt, user_prompt, max_tokens=2500)
        return (raw, ai_err)
    
    def assess_ticket_readiness(self, summary: str, body: str, ac_blocks: List[str], feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None) -> Tuple[str, Optional[str]]:
        """Assess ticket readiness for test case generation."""
        sys_prompt = (
            "You are a senior QA engineer assessing ticket quality and readiness for test case creation. "
            "Evaluate whether the ticket has sufficient information to generate high-quality, comprehensive test cases.\n"
            "IMPORTANT: Acceptance criteria and requirements may be embedded within the description, not explicitly labeled. "
            "Look for implicit requirements, expected behaviors, validation rules, and testable conditions throughout the ticket.\n"
            "Output ONLY JSON with this structure:\n"
            '{ "score": string (one of: "Excellent", "Good", "Poor"), '
            '"confidence": number (0-100), '
            '"summary": string (2-3 sentences overall assessment), '
            '"strengths": [string] (what is present and clear), '
            '"missing_elements": [string] (critical gaps in information), '
            '"recommendations": [string] (specific suggestions to improve), '
            '"quality_concerns": [string] (ambiguities or issues), '
            '"implicit_criteria_found": boolean (true if testable criteria exist even without AC section), '
            '"questions_for_author": [string] (specific questions to ask the ticket author), '
            '"ideal_ticket_example": string (a rewritten version of this ticket as it would look if it scored 100%, including all missing elements, clear AC, proper structure) }'
        )
        
        user_prompt = ""
        
        # Add attachment context
        if attachments:
            user_prompt += f"{Analyzer._format_attachments(attachments)}\n\n"
        
        user_prompt += (
            f"Ticket Summary:\n{summary}\n\n"
            f"Ticket Description:\n{body or '(no description provided)'}\n\n"
            f"Explicit Acceptance Criteria Section:\n{chr(10).join(ac_blocks) if ac_blocks else '(no explicit acceptance criteria section found)'}\n\n"
            f"Feature Context Available: {'Yes - Epic: ' + feature_ctx.get('epic_key', '') if feature_ctx else 'No'}\n\n"
            "Assess this ticket for test case generation readiness. **Look for testable requirements ANYWHERE in the ticket**, including:\n"
            "- Embedded acceptance criteria within the description (look for 'should', 'must', 'will', 'when', 'then')\n"
            "- Expected behaviors and outcomes described in the story\n"
            "- Validation rules and business logic\n"
            "- User flows and interaction patterns\n"
            "- Data requirements and constraints\n"
            "- Error handling and edge cases mentioned\n"
            "- Visual requirements from attached mockups/diagrams\n"
            "- Requirements from attached documents\n\n"
            "Evaluation Criteria:\n"
            "1. **Requirements Clarity**: Are user needs and system behaviors clear? (can be in description)\n"
            "2. **Testable Conditions**: Are there verifiable, measurable outcomes defined?\n"
            "3. **Behavioral Expectations**: Is expected functionality clearly described?\n"
            "4. **Edge Cases**: Are error states, boundaries, or special conditions mentioned?\n"
            "5. **Context Sufficiency**: Is there enough information to understand the feature?\n"
            "6. **Validation Points**: Can success/failure be objectively determined?\n\n"
            "Scoring Guidelines:\n"
            "- **Excellent (90-100%)**: Clear testable requirements (explicit or implicit), well-defined behaviors, "
            "edge cases considered, sufficient context. AC section is helpful but NOT required if requirements are clear in description.\n"
            "- **Good (70-89%)**: Core requirements and behaviors are clear enough for test creation, but missing some "
            "details like edge cases or validation rules. Testable but could benefit from more specificity.\n"
            "- **Poor (<70%)**: Vague or missing requirements, unclear expected behavior, insufficient detail for "
            "creating meaningful tests, lacks testable conditions.\n\n"
            "Remember: A ticket can score 'Excellent' or 'Good' WITHOUT an explicit AC section if the description "
            "contains clear, testable requirements and expected behaviors.\n\n"
            "Finally, generate:\n"
            "1. Questions for Author: 3-5 specific questions to ask the ticket author:\n"
            "   - Address the missing elements and gaps you identified\n"
            "   - Seek clarification on ambiguous or unclear points\n"
            "   - Request specifics for validation rules, edge cases, or error handling if missing\n"
            "   - Be actionable and help improve the ticket quality\n"
            "   - Be phrased professionally and constructively\n\n"
            "2. Ideal Ticket Example: Rewrite this ticket as it would appear if it scored 100% (Excellent):\n"
            "   - MATCH THE AUTHOR'S WRITING STYLE: Use similar tone, terminology, and structure as the original\n"
            "   - Keep the same voice and formatting preferences (bullets, paragraphs, etc.)\n"
            "   - Include the same core functionality but with all missing elements added\n"
            "   - Add clear, testable acceptance criteria (adapt format to author's style - Given/When/Then if they use it, or their preferred format)\n"
            "   - Include validation rules, edge cases, error handling in the author's voice\n"
            "   - Specify expected behaviors explicitly while maintaining their writing patterns\n"
            "   - Preserve any existing good sections verbatim - only enhance/add what's missing\n"
            "   - Keep it concise but comprehensive - this should feel like the author wrote it, just more complete"
        )
        
        
        raw, ai_err = self.llm.complete_json(sys_prompt, user_prompt, max_tokens=1500)
        return (raw, ai_err)

    def generate_test_cases(self, summary: str, body: str, ac_blocks: List[str], feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None, max_iterations: int = 3, progress_callback=None) -> Tuple[List[TestCase], List[Dict], bool, str, Optional[str], Optional[Dict]]:
        """
        Generate test cases with requirement traceability and critic review loop.
        
        Args:
            summary: Ticket summary
            body: Ticket description
            ac_blocks: Acceptance criteria blocks
            feature_ctx: Feature context (epic + children)
            max_iterations: Maximum refinement attempts (default: 3)
        
        Returns:
            Tuple of (test_cases, requirements, used_fallback, raw_response, error, critic_review)
        """
        
        # Process attachments if provided
        processed_attachments = attachments
        if attachments:
            print(f"DEBUG: Processing {len(attachments)} attachments...")
            # Analyze images with GPT-4o vision
            images = [att for att in attachments if att.get("type") == "image"]
            if images:
                print(f"DEBUG: Analyzing {len(images)} images with GPT-4o vision...")
                if progress_callback:
                    progress_callback(45, f"Analyzing {len(images)} attached images...")
                image_analysis = self.analyze_images_with_vision(images)
                if image_analysis:
                    # Add image analysis as a pseudo-document attachment
                    processed_attachments = list(attachments) + [{
                        "filename": "Image Analysis Results",
                        "type": "document",
                        "content": image_analysis
                    }]
                    print(f"DEBUG: Image analysis complete ({len(image_analysis)} chars)")
        
        iteration = 1
        critic_review = None
        previous_attempt = None
        
        while iteration <= max_iterations:
            print("=" * 80)
            print(f"DEBUG: Test Case Generation - ITERATION {iteration}/{max_iterations}")
            print("=" * 80)
            # Report progress
            if progress_callback:
                if iteration == 1:
                    progress_callback(55, f"Generating test cases (attempt {iteration}/{max_iterations})...")
                else:
                    progress_callback(55 + (iteration - 1) * 5, f"Refining test cases (attempt {iteration}/{max_iterations})...")
            print("DEBUG: Generating test cases with context:")
            print(f"  Ticket: {summary}")
            print("=" * 80)
            
            # Generate prompt based on iteration
            if iteration == 1:
                # First attempt: use standard prompt
                user_prompt = self._user_prompt(summary, body, ac_blocks, feature_ctx, processed_attachments)
            else:
                # Refinement attempt: use critic feedback
                user_prompt = self._refinement_prompt(
                    summary, body, ac_blocks, feature_ctx,
                    previous_attempt, critic_review, iteration
                )
            
            # Update progress before LLM call
            if progress_callback:
                progress_callback(60, "Sending request to AI model...")

            # Call LLM with structured outputs
            pydantic_model = TestGenerationResponse if PYDANTIC_AVAILABLE else None
            raw, ai_err = self.llm.complete_json(
                self._system_prompt(), 
                user_prompt,
                max_tokens=16000,  # Increased to ensure all test cases can be generated
                pydantic_model=pydantic_model
            )
            
            # Debug output
            print("=" * 80)
            print(f"DEBUG: Raw AI Response - Iteration {iteration}:")
            print(raw if raw else "None")
            print("=" * 80)
            print(f"DEBUG: Response length: {len(raw) if raw else 0} characters")
            print(f"DEBUG: AI Error: {ai_err}")
            
            # Parse response
            data = safe_json_extract(raw) if raw else None
            
            print("DEBUG: Extracted JSON data:")
            print(f"  Has 'requirements' key: {'requirements' in (data or {})}")
            print(f"  Has 'test_cases' key: {'test_cases' in (data or {})}")
            if data:
                print(f"  Requirements count: {len(data.get('requirements', []))}")
                print(f"  Test cases count: {len(data.get('test_cases', []))}")
            print("=" * 80)
            
            # Extract test cases and requirements
            cases = []
            requirements = []
            
            try:
                # Extract requirements first (for traceability)
                for req in (data or {}).get("requirements", []):
                    req_id = str(req.get("id", "")).strip()
                    req_desc = str(req.get("description", "")).strip()
                    req_source = str(req.get("source", "Unknown")).strip()
                    if req_id and req_desc:
                        requirements.append({"id": req_id, "description": req_desc, "source": req_source})
                
                # Parse test cases with requirement mapping
                for it in (data or {}).get("test_cases", []):
                    title = str(it.get("title","")).strip()
                    if not title: continue
                    prio = int(it.get("priority",2))
                    test_type = str(it.get("test_type","Positive")).strip()
                    tags = [str(t).strip() for t in (it.get("tags") or []) if str(t).strip()]
                    req_id = str(it.get("requirement_id","")).strip()
                    req_desc = str(it.get("requirement_desc","")).strip()
                    
                    # Add test_type as a tag automatically
                    if test_type and test_type not in tags:
                        tags.insert(0, test_type)
                    
                    # Add requirement_id as a tag for filtering
                    if req_id and req_id not in tags:
                        tags.append(req_id)
                    
                    steps = []
                    for s in (it.get("steps") or []):
                        a = str(s.get("action","")).strip()
                        e = str(s.get("expected","")).strip()
                        if a or e: steps.append(TestStep(a or "(action)", e or "(expected)"))
                    if not steps: steps = [TestStep("Execute scenario","Expected behavior as defined.")]
                    
                    cases.append(TestCase(
                        title=title[:160], 
                        priority=prio, 
                        test_type=test_type, 
                        tags=tags or ["ai"], 
                        steps=steps,
                        requirement_id=req_id,
                        requirement_desc=req_desc
                    ))
            except Exception as e:
                print(f"DEBUG: Exception during parsing: {e}")
                cases = []
            
            # Check if we got any cases
            if not cases:
                print("DEBUG: No cases generated, using fallback")
                fallback = Analyzer._fallback_cases(summary, body, ac_blocks)
                return (fallback, [], True, raw or "", ai_err, None)
            
            print(f"DEBUG: Successfully generated {len(cases)} cases with {len(requirements)} requirements")
            
            # === PRE-VALIDATION BEFORE CRITIC REVIEW ===
            print("=" * 80)
            print("DEBUG: Running structural pre-validation...")
            print("=" * 80)
            
            # Validate 1:3 formula
            expected_test_count = len(requirements) * 3
            actual_test_count = len(cases)
            formula_correct = (expected_test_count == actual_test_count)
            
            print(f"  Requirements: {len(requirements)}")
            print(f"  Expected test cases (Req × 3): {expected_test_count}")
            print(f"  Actual test cases: {actual_test_count}")
            print(f"  Formula correct: {formula_correct}")
            
            structural_issues = []
            
            # Check formula
            if not formula_correct:
                structural_issues.append(
                    f"Formula violation: {len(requirements)} requirements should yield "
                    f"{expected_test_count} test cases, but got {actual_test_count}"
                )
            
            # Validate structure - check if test cases have required fields from JSON
            for idx, tc in enumerate(cases, 1):
                # Check step count (should be 1-6)
                if len(tc.steps) < 1:
                    structural_issues.append(
                        f"TC-{idx:03d} '{tc.title}' has 0 steps (minimum 1 required)"
                    )
                elif len(tc.steps) > 6:
                    structural_issues.append(
                        f"TC-{idx:03d} '{tc.title}' has {len(tc.steps)} steps (maximum 6 recommended)"
                    )
            
            if structural_issues:
                print("⚠️  STRUCTURAL VALIDATION FAILED:")
                for issue in structural_issues:
                    print(f"  - {issue}")
                print("=" * 80)
                
                # If this is not the last iteration, regenerate with specific structural feedback
                if iteration < max_iterations:
                    print(f"⚠️  Structural issues detected. Attempting refinement (iteration {iteration + 1}/{max_iterations})...")
                    
                    # Create a minimal critic review for structural issues
                    critic_review = {
                        "approved": False,
                        "overall_quality": "poor",
                        "confidence_score": 100,
                        "requirement_count_correct": True,
                        "test_count_correct": formula_correct,
                        "test_type_distribution_correct": False,
                        "steps_complete": True,
                        "traceability_correct": False,
                        "issues_found": [
                            {
                                "severity": "critical",
                                "issue_type": "structural",
                                "test_case_title": "Multiple test cases",
                                "description": "Test cases are missing required fields or formula is incorrect. " + (
                                    f"Formula error: Expected {expected_test_count} test cases but got {actual_test_count}. " 
                                    if not formula_correct else ""
                                ) + "Some test cases have insufficient steps or missing fields.",
                                "suggestion": (
                                    f"CRITICAL: Apply 1:3 formula - {len(requirements)} requirements MUST yield exactly "
                                    f"{expected_test_count} test cases. Each requirement needs exactly 1 Positive, 1 Negative, "
                                    f"and 1 Edge Case test. Every test case MUST include all required fields and 1-6 steps."
                                )
                            }
                        ],
                        "missing_test_scenarios": [],
                        "strengths": [],
                        "summary": (
                            f"Structural validation failed. " +
                            (f"Formula is wrong: expected {expected_test_count} test cases but got {actual_test_count}. " 
                             if not formula_correct else "") +
                            "Some test cases are missing critical fields."
                        ),
                        "recommendation": (
                            f"Regenerate following 1:3 formula: {len(requirements)} requirements × 3 = "
                            f"{expected_test_count} test cases. Ensure each requirement has exactly 1 Positive, "
                            f"1 Negative, and 1 Edge Case test. Include all required fields in every test case."
                        )
                    }
                    
                    previous_attempt = raw
                    iteration += 1
                    continue
                else:
                    print("❌ FINAL ATTEMPT: Structural issues remain after maximum iterations")
                    print("   Returning with structural validation failure")
                    print("=" * 80)
            else:
                print("✅ Structural pre-validation passed")
                print("=" * 80)
            
            # Report progress before critic review
            if progress_callback:
                progress_callback(70, "Running quality review...")
            
            # === CRITIC REVIEW GUARDRAIL ===
            print("=" * 80)
            print(f"DEBUG: Starting QA Manager Critic Review - Iteration {iteration}...")
            print("=" * 80)
            
            critic_review, critic_err = self.critic_review_test_cases(
                summary=summary,
                body=body,
                ac_blocks=ac_blocks,
                requirements=requirements,
                test_cases_json=raw,
                generated_test_cases=cases
            )
            
            if critic_err:
                print(f"DEBUG: Critic review failed: {critic_err}")
                print("DEBUG: Proceeding with current test cases (critic unavailable)")
                return (cases, requirements, False, raw or "", ai_err, None)
            
            if not critic_review:
                print("DEBUG: No critic review received, proceeding with test cases")
                return (cases, requirements, False, raw or "", ai_err, None)
            
            # Display critic results
            print("DEBUG: Critic Review Results:")
            print(f"  Overall Quality: {critic_review.get('overall_quality', 'unknown')}")
            print(f"  Approved: {critic_review.get('approved', False)}")
            print(f"  Confidence: {critic_review.get('confidence_score', 0)}%")
            print(f"  Formula Correct: {critic_review.get('requirement_count_correct', False)} & {critic_review.get('test_count_correct', False)}")
            print(f"  Distribution Correct: {critic_review.get('test_type_distribution_correct', False)}")
            print(f"  Steps Complete: {critic_review.get('steps_complete', False)}")
            print(f"  Traceability Correct: {critic_review.get('traceability_correct', False)}")
            
            if critic_review.get('strengths'):
                print(f"\n  Strengths:")
                for strength in critic_review['strengths']:
                    print(f"    ✓ {strength}")
            
            if critic_review.get('issues_found'):
                print(f"\n  Issues Found ({len(critic_review['issues_found'])}):")
                for issue in critic_review['issues_found']:
                    severity_icon = "🔴" if issue.get('severity') == 'critical' else "🟡" if issue.get('severity') == 'major' else "🔵"
                    print(f"    {severity_icon} [{issue.get('issue_type', 'unknown')}] {issue.get('test_case_title', 'N/A')}")
                    print(f"       Problem: {issue.get('description', '')}")
                    print(f"       Fix: {issue.get('suggestion', '')}")
            
            if critic_review.get('missing_test_scenarios'):
                print(f"\n  Missing Scenarios:")
                for scenario in critic_review['missing_test_scenarios']:
                    print(f"    ⚠️  {scenario}")
            
            print(f"\n  Summary: {critic_review.get('summary', 'N/A')}")
            print(f"  Recommendation: {critic_review.get('recommendation', 'N/A')}")
            print("=" * 80)
            
            # Check if approved
            if critic_review.get('approved', False):
                print(f"✅ SUCCESS: Test cases APPROVED by critic on iteration {iteration}")
                print("=" * 80)
                # Report completion
                if progress_callback:
                    progress_callback(75, "Quality review passed!")
                return (cases, requirements, False, raw or "", ai_err, critic_review)
            
            # Not approved - check if we can retry
            if iteration < max_iterations:
                print(f"⚠️  Test cases REJECTED. Attempting refinement (iteration {iteration + 1}/{max_iterations})...")
                print("=" * 80)
                previous_attempt = raw
                iteration += 1
                continue
            else:
                print(f"❌ FINAL ATTEMPT: Test cases still not approved after {max_iterations} iterations")
                print("   Returning latest attempt with critic feedback for manual review")
                print("=" * 80)
                return (cases, requirements, False, raw or "", ai_err, critic_review)
        
        # Should never reach here, but just in case
        return (cases, requirements, False, raw or "", ai_err, critic_review)
    
    def critic_review_test_cases(
        self, 
        summary: str, 
        body: str, 
        ac_blocks: List[str],
        requirements: List[Dict],
        test_cases_json: str,
        generated_test_cases: List[TestCase]
    ) -> Tuple[Optional[Dict], Optional[str]]:
        """
        QA Manager Critic: Reviews generated test cases for quality and correctness.
        Returns (critic_response_dict, error_message)
        """
        
        sys_prompt = (
            "You are a Senior QA Manager and Test Architect with 15+ years of experience reviewing test documentation.\n"
            "Your role is to critically assess test case quality and enforce the 1:3 FORMULA.\n\n"
            
            "CRITICAL VALIDATION FORMULA:\n"
            "N Requirements → N × 3 Test Cases (EXACTLY)\n"
            "- For EACH requirement, there must be EXACTLY 3 test cases:\n"
            "  * 1 Positive test (happy path)\n"
            "  * 1 Negative test (error handling)\n"
            "  * 1 Edge Case test (boundary conditions)\n"
            "- Example: 10 requirements → 30 test cases ✓\n"
            "- Example: 13 requirements → 39 test cases ✓\n"
            "- Example: 10 requirements → 25 test cases ✗ REJECT (missing tests)\n"
            "- Example: 10 requirements → 35 test cases ✗ REJECT (too many tests)\n\n"
            
            "YOUR REVIEW CHECKLIST:\n\n"
            
            "1. ✓ FORMULA CHECK (CRITICAL):\n"
            "   - Count requirements: ___ (N)\n"
            "   - Count test cases: ___ (must equal N × 3)\n"
            "   - If test case count ≠ requirements × 3 → IMMEDIATE REJECT\n\n"
            
            "2. ✓ TRACEABILITY MATRIX (CRITICAL):\n"
            "   - Create a matrix showing each requirement and its 3 test cases\n"
            "   - REQ-001: [Positive test], [Negative test], [Edge Case test]\n"
            "   - REQ-002: [Positive test], [Negative test], [Edge Case test]\n"
            "   - If ANY requirement has ≠ 3 test cases → REJECT\n"
            "   - If ANY test case references non-existent requirement → REJECT\n\n"
            
            "3. ✓ STRUCTURAL COMPLETENESS (CRITICAL):\n"
            "   Every test case MUST have ALL of these fields:\n"
            "   - requirement_id: Single requirement ID\n"
            "   - requirement_desc: Description of what's tested\n"
            "   - title: Following 'REQ-XXX [Type]: Description' format\n"
            "   - priority: 1, 2, or 3\n"
            "   - test_type: Exactly 'Positive', 'Negative', or 'Edge Case'\n"
            "   - tags: Array of tags\n"
            "   - steps: Array with 1-6 steps (each with action and expected)\n"
            "   If ANY test case is missing ANY field → REJECT\n\n"
            
            "4. ✓ TEST TYPE DISTRIBUTION:\n"
            "   For N requirements:\n"
            "   - Positive tests: N (one per requirement)\n"
            "   - Negative tests: N (one per requirement)\n"
            "   - Edge Case tests: N (one per requirement)\n"
            "   - Total: N × 3\n"
            "   If distribution is wrong → REJECT\n\n"
            
            "5. ✓ STEP QUALITY:\n"
            "   - Each test case should have 1-6 steps (appropriate to complexity)\n"
            "   - Simple validation: 1-3 steps is acceptable\n"
            "   - Complex workflow: 4-6 steps is appropriate\n"
            "   - Each step must have clear 'action' and 'expected' fields\n"
            "   - Steps should be specific and actionable\n\n"
            
            "6. ✓ NAMING CONVENTION:\n"
            "   - Titles should follow: 'REQ-XXX [Type]: Description'\n"
            "   - Example: 'REQ-001 Positive: Valid email accepted'\n"
            "   - Example: 'REQ-001 Negative: Invalid email rejected'\n"
            "   - Example: 'REQ-001 Edge Case: Email with special characters'\n\n"
            
            "APPROVAL CRITERIA:\n"
            "APPROVE (approved=true) only if ALL of these are true:\n"
            "✓ Formula is correct: Test case count = Requirements × 3\n"
            "✓ Every requirement has exactly 3 test cases (1 Positive, 1 Negative, 1 Edge Case)\n"
            "✓ Every test case has ALL required fields\n"
            "✓ Test types are correctly distributed\n"
            "✓ Step counts are appropriate (1-6 per test case)\n"
            "✓ Naming convention is followed\n"
            "✓ No critical or major issues found\n\n"
            
            "REJECT (approved=false) if ANY of these are true:\n"
            "✗ Formula is wrong: Test case count ≠ Requirements × 3\n"
            "✗ Any requirement has ≠ 3 test cases\n"
            "✗ Any test case is missing required fields (check Pydantic schema)\n"
            "✗ Test type distribution is incorrect\n"
            "✗ Any test case has 0 steps or > 6 steps\n"
            "✗ Test case titles don't follow naming convention\n\n"
            
            "OUTPUT STRUCTURE:\n"
            "Provide detailed analysis including:\n"
            "- Traceability matrix (requirement → 3 test cases)\n"
            "- Formula verification (N requirements × 3 = actual count)\n"
            "- List of any issues found with specific test cases\n"
            "- Clear recommendation (approve or specific fixes needed)\n\n"
            
            "Be thorough and specific. Your review determines production readiness."
        )
        
        user_prompt = (
            f"ORIGINAL TICKET:\n"
            f"Summary: {summary}\n"
            f"Description: {body[:2000]}\n"
            f"Acceptance Criteria: {chr(10).join(ac_blocks) if ac_blocks else '(none)'}\n\n"
            
            f"GENERATED REQUIREMENTS ({len(requirements)}):\n"
        )
        
        for req in requirements:
            user_prompt += f"- {req.get('id')}: {req.get('description')}\n"
        
        user_prompt += f"\n\nGENERATED TEST CASES ({len(generated_test_cases)}):\n"
        
        # Group test cases by requirement for easier review
        by_req = {}
        for tc in generated_test_cases:
            req_id = tc.requirement_id or "UNMAPPED"
            if req_id not in by_req:
                by_req[req_id] = []
            by_req[req_id].append(tc)
        
        for req_id, tcs in by_req.items():
            user_prompt += f"\n{req_id} ({len(tcs)} test cases):\n"
            for tc in tcs:
                user_prompt += f"  • [{tc.test_type}] {tc.title}\n"
                user_prompt += f"    Priority: {tc.priority} | Steps: {len(tc.steps)}\n"
                if len(tc.steps) <= 3:
                    # Show steps if there are few (potential issue)
                    for i, step in enumerate(tc.steps, 1):
                        user_prompt += f"      {i}. {step.action[:80]}...\n"
        
        expected_count = len(requirements) * 3
        actual_count = len(generated_test_cases)
        formula_correct = (expected_count == actual_count)
        
        user_prompt += (
            "\n\n=== REVIEW CHECKLIST ===\n\n"
            f"1. FORMULA CHECK:\n"
            f"   Requirements: {len(requirements)}\n"
            f"   Expected test cases: {len(requirements)} × 3 = {expected_count}\n"
            f"   Actual test cases: {actual_count}\n"
            f"   Formula correct: {'✓ YES' if formula_correct else '✗ NO - REJECT'}\n\n"
            
            "2. TRACEABILITY MATRIX:\n"
            "   Create a matrix showing each requirement and its 3 test cases:\n"
            "   - Each requirement should have exactly 1 Positive, 1 Negative, 1 Edge Case\n"
            "   - Verify no requirements are missing test cases\n"
            "   - Verify no extra test cases exist\n\n"
            
            "3. STRUCTURAL COMPLETENESS:\n"
            "   Check that every test case has ALL required fields:\n"
            "   - requirement_id, requirement_desc, title, priority, test_type, tags\n"
            "   - steps (1-6 with action and expected)\n\n"
            
            "4. TEST TYPE DISTRIBUTION:\n"
            f"   - Positive tests: Should be {len(requirements)}\n"
            f"   - Negative tests: Should be {len(requirements)}\n"
            f"   - Edge Case tests: Should be {len(requirements)}\n"
            "   - Verify distribution is exact\n\n"
            
            "5. STEP QUALITY:\n"
            "   - Each test case should have 1-6 steps (appropriate to complexity)\n"
            "   - Steps should be clear, specific, and actionable\n\n"
            
            "6. NAMING CONVENTION:\n"
            "   - Verify titles follow: 'REQ-XXX [Type]: Description'\n\n"
            
            "Perform a thorough review. Build the traceability matrix.\n"
            "Be critical but fair. If there are issues, be specific about what needs to be fixed.\n"
            "Remember: You're enforcing the 1:3 formula - this is non-negotiable."
        )
        
        # Use structured outputs for critic review
        pydantic_model = CriticReviewResponse if PYDANTIC_AVAILABLE else None
        raw, ai_err = self.llm.complete_json(
            sys_prompt,
            user_prompt,
            max_tokens=3000,
            pydantic_model=pydantic_model
        )
        
        if ai_err:
            return (None, ai_err)
        
        # Parse the critic response
        critic_data = safe_json_extract(raw) if raw else None
        
        return (critic_data, None)
    
    def analyze_and_split_for_test_tickets(
        self,
        feature_context: Dict
    ) -> Tuple[Optional[Dict], Optional[str]]:
        """
        Analyze Epic + children and recommend how to split into test tickets.
        Returns (split_plan_dict, error_message)
        """
        sys_prompt = (
            "You are a Senior Business Analyst / Test Strategist.\n"
            "Your task is to analyze an Epic and its child tickets, then recommend how to structure test tickets.\n\n"
            
            "ANALYSIS APPROACH:\n"
            "1. Review all child tickets and group by functional area or user journey\n"
            "2. Each test ticket should cover 5-8 acceptance criteria (black-box)\n"
            "3. Group related functionality together\n"
            "4. Consider risk, complexity, and testability\n\n"
            
            "GROUPING STRATEGIES:\n"
            "- By functional area (e.g., 'Authentication', 'Data Migration', 'UI Components')\n"
            "- By user journey (e.g., 'New User Onboarding', 'Power User Workflow')\n"
            "- By feature module (e.g., 'Payment Processing', 'Report Generation')\n"
            "- By priority/risk (high-risk features get dedicated tickets)\n\n"
            
            "OUTPUT: Recommend 2-5 test tickets, each focused on a clear functional area."
        )
        
        # Build context
        if feature_context.get('type') == 'initiative':
            context_summary = (
                f"Initiative: {feature_context.get('initiative_summary', '')}\n"
                f"Epics: {', '.join(feature_context.get('epics', []))}\n"
                f"Total child tickets: {feature_context.get('total_children', 0)}\n\n"
                f"Child Tickets:\n"
            )
        else:
            context_summary = (
                f"Epic: {feature_context.get('epic_summary', '')}\n"
                f"Description: {feature_context.get('epic_desc', '')[:500]}\n\n"
                f"Child Tickets:\n"
            )
        
        # Add child ticket details
        children = feature_context.get('children', [])[:20]  # Limit to 20 for token management
        for child in children:
            context_summary += f"- {child.get('key', '')}: {child.get('summary', '')}\n"
            desc = child.get('desc', '')
            if desc:
                context_summary += f"  {desc[:200]}...\n"
        
        if feature_context.get('total_children', 0) > 20:
            context_summary += f"\n... and {feature_context.get('total_children') - 20} more tickets\n"
        
        user_prompt = (
            f"{context_summary}\n\n"
            "Based on the above, recommend how to structure test tickets.\n"
            "Each test ticket should have a clear functional focus and cover 5-8 black-box acceptance criteria.\n"
            "Provide your recommended structure."
        )
        
        pydantic_model = TestTicketSplitResponse if PYDANTIC_AVAILABLE else None
        raw, ai_err = self.llm.complete_json(
            sys_prompt,
            user_prompt,
            max_tokens=2000,
            pydantic_model=pydantic_model
        )
        
        if ai_err:
            return (None, ai_err)
        
        split_data = safe_json_extract(raw) if raw else None
        return (split_data, None)
    
    def generate_test_ticket(
        self,
        epic_name: str,
        functional_area: str,
        child_tickets: List[Dict],
        feature_context: Dict,
        previous_attempt: Optional[str] = None,
        reviewer_feedback: Optional[Dict] = None
    ) -> Tuple[Optional[Dict], Optional[str]]:
        """
        Generate a single test ticket using BA/PO persona.
        Returns (ticket_content_dict, error_message)
        """
        
        if previous_attempt and reviewer_feedback:
            # Refinement mode
            sys_prompt = (
                "You are a Senior Business Analyst / Product Owner.\n"
                "You previously created a test ticket that was REJECTED by the Product Manager.\n"
                "Study the feedback carefully and create an IMPROVED version.\n\n"
                
                "CRITICAL: IGNORE OUT-OF-SCOPE FUNCTIONALITY:\n"
                "- DO NOT include any functionality marked as 'removed from scope', 'out of scope', or 'not in scope'\n"
                "- If delete operations are marked as removed, DO NOT create test cases for delete\n"
                "- Only test IN-SCOPE functionality\n\n"
                
                "CRITICAL REQUIREMENTS:\n"
                "- Write in the Epic author's style (match tone, structure, terminology)\n"
                "- **MANDATORY**: Include 'Source Tickets' section at END of description (format: 'Source Tickets: KEY-123: Title, ...')\n"
                "- **DO NOT OMIT** the source tickets - this is required for traceability\n"
                "- MATCH the Epic's AC format (bullets, numbered, Given-When-Then, etc.)\n"
                "- DO NOT convert simple bullets to Given-When-Then unless Epic uses it\n"
                "- Focus on BLACK-BOX acceptance criteria (user-facing behavior)\n"
                "- NO technical implementation details\n"
                "- 5-8 testable acceptance criteria\n"
                "- Each AC must be verifiable by a manual tester\n\n"
                
                "Fix the issues mentioned in the review feedback."
            )
            
            user_prompt = (
                f"=== PREVIOUS ATTEMPT (REJECTED) ===\n{previous_attempt}\n\n"
                f"=== REVIEWER FEEDBACK ===\n"
                f"Quality Score: {reviewer_feedback.get('quality_score', 0)}/100\n"
                f"Issues:\n"
            )
            for issue in reviewer_feedback.get('issues', []):
                user_prompt += f"- {issue}\n"
            user_prompt += "\nRecommendations:\n"
            for rec in reviewer_feedback.get('recommendations', []):
                user_prompt += f"- {rec}\n"
            
            user_prompt += f"\n\n=== CREATE IMPROVED VERSION ===\n"
            user_prompt += f"Epic Name: {epic_name}\n"
            user_prompt += f"Functional Area: {functional_area}\n\n"
            
        else:
            # First attempt
            sys_prompt = (
                "You are a Senior Business Analyst / Product Owner tasked with creating test tickets.\n\n"
                
                "YOUR ROLE:\n"
                "- Analyze the Epic and child tickets to understand the feature\n"
                "- Create a comprehensive test ticket for this functional area\n"
                "- Match the Epic author's writing style (tone, terminology, structure)\n"
                "- Focus on BLACK-BOX acceptance criteria for manual testing\n\n"
                
                "CRITICAL: IGNORE OUT-OF-SCOPE FUNCTIONALITY:\n"
                "- If you see ANY mention of 'removed from scope', 'out of scope', or 'not in scope', DO NOT include that functionality in the test ticket\n"
                "- If delete operations are marked as removed, DO NOT create test cases for delete\n"
                "- Only test functionality that is IN SCOPE\n"
                "- When in doubt, exclude rather than include\n\n"
                
                "CRITICAL REQUIREMENTS:\n"
                "1. Summary: Follow format '[Epic Name] - Testing - [Functional Area]'\n"
                "2. Description:\n"
                "   - Clear context about what's being tested, matching Epic's style\n"
                "   - **MANDATORY**: At the END of the description, add a 'Source Tickets' section listing the child tickets used\n"
                "   - **REQUIRED FORMAT**: 'Source Tickets: KEY-123: Title, KEY-456: Title, ...'\n"
                "   - **DO NOT SKIP THIS** - every test ticket MUST include source tickets at the end\n"
                "3. Acceptance Criteria: 5-8 black-box criteria, each:\n"
                "   - Focuses on USER-FACING behavior\n"
                "   - Is VERIFIABLE by a manual tester\n"
                "   - Avoids technical implementation details\n"
                "   - MATCHES the Epic's AC format (bullet points, numbered list, Given-When-Then, etc.)\n"
                "   - DO NOT convert simple bullets to Given-When-Then unless the Epic uses that format\n\n"
                
                "WRITING STYLE DETECTION:\n"
                "- If Epic uses 'As a... I want... so that...' → Use user story format\n"
                "- If Epic ACs are simple bullets → Keep them as simple bullets\n"
                "- If Epic ACs use 'Given... When... Then...' → Match that format\n"
                "- If Epic ACs are numbered lists → Keep numbered format\n"
                "- NEVER impose a format the Epic doesn't use\n\n"
                
                "BLACK-BOX vs WHITE-BOX:\n"
                "✅ BLACK-BOX: 'User can select payment method from dropdown'\n"
                "❌ WHITE-BOX: 'API calls /v2/payments endpoint'\n"
                "✅ BLACK-BOX: 'System displays error message when invalid email entered'\n"
                "❌ WHITE-BOX: 'Database updates user_payments table'\n\n"
                
                "This ticket should score 80%+ on the Ticket Readiness assessment."
            )
            
            user_prompt = (
                f"Epic Name: {epic_name}\n"
                f"Functional Area: {functional_area}\n\n"
            )
        
        # Add feature context
        if feature_context.get('type') == 'initiative':
            user_prompt += f"Initiative: {feature_context.get('initiative_summary', '')}\n\n"
        else:
            user_prompt += f"Epic Description:\n{feature_context.get('epic_desc', '')[:800]}\n\n"
        
        # Add relevant child tickets
        user_prompt += "Relevant Child Tickets:\n"
        for child in child_tickets[:10]:  # Limit to 10 for token management
            user_prompt += f"- {child.get('key', '')}: {child.get('summary', '')}\n"
            desc = child.get('desc', '')
            if desc:
                user_prompt += f"  {desc[:300]}...\n\n"
        
        user_prompt += "\nCreate the test ticket now."
        
        pydantic_model = TestTicketContent if PYDANTIC_AVAILABLE else None
        raw, ai_err = self.llm.complete_json(
            sys_prompt,
            user_prompt,
            max_tokens=3000,
            pydantic_model=pydantic_model
        )
        
        if ai_err:
            return (None, ai_err)
        
        ticket_data = safe_json_extract(raw) if raw else None
        return (ticket_data, raw)  # Return both parsed and raw for storage
    
    def review_test_ticket(
        self,
        ticket_content: Dict,
        epic_name: str,
        feature_context: Dict
    ) -> Tuple[Optional[Dict], Optional[str]]:
        """
        Review a test ticket using BA/PM persona.
        Returns (review_dict, error_message)
        """
        sys_prompt = (
            "You are a Senior Business Analyst / Product Manager reviewing test tickets.\n\n"
            
            "YOUR ROLE:\n"
            "- Assess test ticket quality and readiness\n"
            "- Enforce black-box focus (no implementation details)\n"
            "- Verify acceptance criteria are testable by manual testers\n"
            "- Check writing style matches the Epic\n"
            "- Ensure ticket would score 80%+ on Readiness assessment\n\n"
            
            "APPROVAL CRITERIA:\n"
            "✓ 5-8 black-box acceptance criteria\n"
            "✓ All AC are manually testable (no technical jargon)\n"
            "✓ Clear description with sufficient context\n"
            "✓ Follows Epic's writing style\n"
            "✓ Would score 80%+ on Ticket Readiness\n"
            "✓ No implementation details (APIs, databases, code)\n\n"
            
            "REJECTION REASONS:\n"
            "✗ Too few or too many AC (not 5-8)\n"
            "✗ Contains technical implementation details\n"
            "✗ AC not testable by manual tester\n"
            "✗ Insufficient context or description\n"
            "✗ Style/format doesn't match Epic (e.g., Epic uses bullets but ticket uses Given-When-Then)\n"
            "✗ Would score <80% on Readiness\n\n"
            
            "If approved=True: quality_score 80-100\n"
            "If approved=False: provide specific issues and recommendations"
        )
        
        user_prompt = (
            f"Epic Name: {epic_name}\n\n"
            f"EPIC ACCEPTANCE CRITERIA (for style reference):\n"
        )
        
        # Add Epic AC samples for style comparison
        if feature_context.get('type') == 'epic':
            epic_ac = feature_context.get('epic_ac', [])
            if epic_ac:
                for i, ac in enumerate(epic_ac[:3], 1):  # Show first 3 ACs as examples
                    user_prompt += f"  {i}. {ac[:200]}...\n"
            else:
                user_prompt += "  (No explicit AC in Epic - match description style)\n"
        user_prompt += "\n"
        
        user_prompt += (
            f"TEST TICKET TO REVIEW:\n\n"
            f"Summary: {ticket_content.get('summary', '')}\n\n"
            f"Description:\n{ticket_content.get('description', '')}\n\n"
            f"Acceptance Criteria:\n"
        )
        
        for i, ac in enumerate(ticket_content.get('acceptance_criteria', []), 1):
            user_prompt += f"{i}. {ac}\n"
        
        user_prompt += "\n\nReview this test ticket and determine if it's ready or needs revision."
        
        pydantic_model = TestTicketReview if PYDANTIC_AVAILABLE else None
        raw, ai_err = self.llm.complete_json(
            sys_prompt,
            user_prompt,
            max_tokens=2000,
            pydantic_model=pydantic_model
        )
        
        if ai_err:
            return (None, ai_err)
        
        review_data = safe_json_extract(raw) if raw else None
        return (review_data, None)

# ---------- Exporter ----------
class Exporter:
    def __init__(self, out_dir: Path):
        self.out_dir=out_dir; self.out_dir.mkdir(parents=True, exist_ok=True)
    def write_catalog_csv(self, key: str, summary: str, cases: List[TestCase]) -> Path:
        p=self.out_dir/"testcases.csv"
        with p.open("w", newline="", encoding="utf-8") as f:
            w=csv.writer(f); w.writerow(["Jira Key","Title","Test Type","Priority","Tags","Objective","Num Steps"])
            for c in cases: w.writerow([key,c.title,c.test_type,c.priority,",".join(c.tags),c.objective,len(c.steps)])
        return p
    def write_case_tsvs(self, key: str, summary: str, cases: List[TestCase]) -> List[Path]:
        folder=self.out_dir/slugify(f"{key}-{summary}"); folder.mkdir(parents=True, exist_ok=True)
        paths=[]
        for i,c in enumerate(cases,1):
            fn=folder/f"TC-{i:02d}-{slugify(c.title)[:50]}.tsv"
            with fn.open("w", encoding="utf-8") as f:
                f.write("Action\tExpected Result\n")
                for s in c.steps: f.write(f"{s.action}\t{s.expected}\n")
            paths.append(fn)
        return paths
    def write_azure_devops_csv(self, cases: List[TestCase]) -> Path:
        p=self.out_dir/"azure-devops-testcases.csv"
        with p.open("w", newline="", encoding="utf-8") as f:
            w=csv.writer(f); w.writerow(["ID","Work Item Type","Title","Test Step","Step Action","Step Expected"])
            for c in cases:
                title=(c.title or "Untitled").strip()[:128]; steps=c.steps or [TestStep("Execute scenario","Expected behavior as defined.")]
                for idx,s in enumerate(steps,1): w.writerow(["","Test Case",title,idx,s.action or "",s.expected or ""])
        return p



# ========== ENHANCED STYLESHEET WITH TRANSITIONS ==========
NEBULA_QSS = """
/* ========== BASE STYLES ========== */
* { 
    color: #E8EBF0; 
    font-family: 'Inter', 'Plus Jakarta Sans', 'Segoe UI', system-ui, -apple-system, Arial; 
    font-weight: 600; 
}

QMainWindow, QWidget { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, 
        stop:0 #0A0E15, 
        stop:0.5 #0C1018, 
        stop:1 #0A0D14);
}

/* ========== TOP BAR - GLASSMORPHISM ========== */
#TopBar { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(18, 24, 38, 0.85), 
        stop:1 rgba(12, 18, 30, 0.75)); 
    border: 1px solid rgba(59, 130, 246, 0.2);
    border-radius: 20px; 
    padding: 12px;
}

#TopBar QLabel { 
    color: #F8FAFC; 
}

/* ========== TAB NAVIGATION - MODERN PILLS ========== */
#TopTabs QPushButton { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(30, 41, 59, 0.6), 
        stop:1 rgba(20, 30, 48, 0.5)); 
    border: 1px solid rgba(71, 85, 105, 0.4); 
    border-radius: 20px; 
    padding: 14px 26px; 
    margin: 0 8px; 
    color: #CBD5E1;
    font-weight: 700;
    font-size: 13px;
    letter-spacing: 0.5px;
    min-height: 42px;
}

#TopTabs QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(51, 65, 85, 0.7), 
        stop:1 rgba(30, 41, 59, 0.6));
    border: 1px solid rgba(96, 165, 250, 0.5);
    color: #F1F5F9;
}

#TopTabs QPushButton[active="true"] { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
        stop:0 #2563EB, 
        stop:0.5 #3B82F6,
        stop:1 #1D4ED8);
    border: 1px solid rgba(147, 197, 253, 0.8);
    color: #FFFFFF;
    font-weight: 800;
}

/* ========== INPUT FIELDS - ENHANCED DEPTH ========== */
QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(15, 23, 42, 0.9),
        stop:1 rgba(20, 30, 50, 0.85)); 
    border: 2px solid rgba(51, 65, 85, 0.5); 
    border-radius: 16px; 
    padding: 14px 18px;
    selection-background-color: rgba(59, 130, 246, 0.4);
    color: #E2E8F0;
    font-size: 14px;
    min-height: 40px;
}

QLineEdit:hover, QTextEdit:hover, QPlainTextEdit:hover, QSpinBox:hover {
    border: 2px solid rgba(96, 165, 250, 0.6);
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(20, 30, 50, 0.95),
        stop:1 rgba(25, 35, 60, 0.9));
}

QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus, QSpinBox:focus { 
    border: 2px solid #3B82F6;
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(25, 35, 60, 1),
        stop:1 rgba(30, 45, 75, 0.98));
    color: #F8FAFC;
}

/* ========== COMBO BOX - ENHANCED ========== */
QComboBox {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(15, 23, 42, 0.9),
        stop:1 rgba(20, 30, 50, 0.85)); 
    border: 2px solid rgba(51, 65, 85, 0.5); 
    border-radius: 16px; 
    padding: 10px 14px;
    color: #E2E8F0;
}

QComboBox:hover {
    border: 2px solid rgba(96, 165, 250, 0.6);
}

QComboBox:focus {
    border: 2px solid #3B82F6;
}

QComboBox::drop-down {
    border: none;
    width: 35px;
}

QComboBox::down-arrow {
    image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTYiIGhlaWdodD0iMTAiIHZpZXdCb3g9IjAgMCAxNiAxMCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj48cGF0aCBkPSJNMSAxTDggOEwxNSAxIiBzdHJva2U9IiM5NEE1QjgiIHN0cm9rZS13aWR0aD0iMiIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+PC9zdmc+);
}

QComboBox QAbstractItemView {
    background: rgba(20, 30, 50, 0.98);
    border: 1px solid rgba(59, 130, 246, 0.3);
    border-radius: 14px;
    selection-background-color: #2563EB;
    color: #E2E8F0;
    padding: 6px;
}

/* ========== BUTTONS - PREMIUM LOOK ========== */
QPushButton { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(30, 41, 59, 0.8), 
        stop:1 rgba(20, 30, 48, 0.7)); 
    border: 1px solid rgba(71, 85, 105, 0.6); 
    border-radius: 18px; 
    padding: 14px 28px; 
    color: #E2E8F0;
    font-weight: 700;
    font-size: 14px;
    letter-spacing: 0.3px;
    min-height: 40px;
}

QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(51, 65, 85, 0.9), 
        stop:1 rgba(30, 41, 59, 0.8));
    border: 1px solid rgba(96, 165, 250, 0.7);
    color: #FFFFFF;
}

QPushButton:pressed {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(15, 23, 42, 0.95), 
        stop:1 rgba(20, 30, 48, 0.9));
    border: 1px solid rgba(59, 130, 246, 0.8);
    padding-top: 13px;
}

QPushButton:disabled {
    background: rgba(15, 23, 42, 0.5);
    border: 1px solid rgba(51, 65, 85, 0.3);
    color: #64748B;
}

/* ========== DROPDOWN MENUS ========== */
QMenu {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(18, 24, 38, 0.95), 
        stop:1 rgba(12, 18, 30, 0.90));
    border: 1px solid rgba(59, 130, 246, 0.3);
    border-radius: 12px;
    padding: 6px;
    color: #E2E8F0;
}

QMenu::item {
    background: transparent;
    padding: 10px 20px;
    border-radius: 8px;
    margin: 2px 0px;
}

QMenu::item:selected {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(59, 130, 246, 0.3),
        stop:1 rgba(37, 99, 235, 0.2));
    color: #FFFFFF;
}

QMenu::separator {
    height: 1px;
    background: rgba(71, 85, 105, 0.3);
    margin: 4px 10px;
}

/* ========== ACCENT BUTTON - VIBRANT GLOW ========== */
QPushButton#accent { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
        stop:0 #2563EB, 
        stop:0.5 #3B82F6, 
        stop:1 #1D4ED8);
    border: 1px solid rgba(147, 197, 253, 0.6);
    color: #FFFFFF;
    font-weight: 800;
    padding: 15px 32px;
    font-size: 15px;
    min-height: 44px;
}

QPushButton#accent:hover {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
        stop:0 #1D4ED8, 
        stop:0.5 #2563EB, 
        stop:1 #1E40AF);
    border: 1px solid rgba(147, 197, 253, 0.9);
}

QPushButton#accent:pressed {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
        stop:0 #1E40AF, 
        stop:1 #1D4ED8);
}

/* ========== CARDS - GLASSMORPHISM DEPTH ========== */
QFrame[card="true"] { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(20, 30, 48, 0.75), 
        stop:1 rgba(15, 23, 42, 0.65));
    border: 1px solid rgba(71, 85, 105, 0.4);
    border-radius: 24px;
    padding: 2px;
}

QFrame[card="true"]:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(30, 41, 59, 0.8), 
        stop:1 rgba(20, 30, 48, 0.7));
    border: 1px solid rgba(96, 165, 250, 0.5);
}

/* ========== TILE CARDS - MODERN ELEVATION ========== */
QFrame[card="true"][tile="true"] {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(25, 35, 58, 0.9), 
        stop:1 rgba(18, 26, 45, 0.85));
    border: 1px solid rgba(71, 85, 105, 0.5);
    border-radius: 22px;
}

QFrame[card="true"][tile="true"]:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(30, 45, 70, 0.95), 
        stop:1 rgba(22, 32, 55, 0.9));
    border: 1px solid rgba(96, 165, 250, 0.6);
}

QFrame[card="true"][tile="true"][selected="true"] {
    border: 2px solid #3B82F6;
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
        stop:0 rgba(37, 99, 235, 0.25), 
        stop:1 rgba(29, 78, 216, 0.2));
}

/* ========== STATUS BADGES - VIBRANT NEON ========== */
.pill { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 rgba(30, 58, 138, 0.6),
        stop:1 rgba(37, 99, 235, 0.5)); 
    border: 1px solid rgba(96, 165, 250, 0.6); 
    color: #93C5FD; 
    padding: 6px 14px; 
    border-radius: 999px; 
    font-size: 11px;
    font-weight: 800;
    letter-spacing: 0.5px;
}

QLabel[priority="High"] { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 rgba(220, 38, 38, 0.35), 
        stop:1 rgba(239, 68, 68, 0.3));
    border: 1px solid rgba(248, 113, 113, 0.7); 
    color: #FCA5A5; 
    padding: 6px 14px; 
    border-radius: 999px; 
    font-size: 11px;
    font-weight: 800;
    letter-spacing: 0.5px;
}

QLabel[priority="Medium"] { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 rgba(245, 158, 11, 0.35), 
        stop:1 rgba(251, 191, 36, 0.3));
    border: 1px solid rgba(252, 211, 77, 0.7); 
    color: #FDE047; 
    padding: 6px 14px; 
    border-radius: 999px; 
    font-size: 11px;
    font-weight: 800;
    letter-spacing: 0.5px;
}

QLabel[priority="Low"] { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 rgba(22, 163, 74, 0.35), 
        stop:1 rgba(34, 197, 94, 0.3));
    border: 1px solid rgba(74, 222, 128, 0.7); 
    color: #86EFAC; 
    padding: 6px 14px; 
    border-radius: 999px; 
    font-size: 11px;
    font-weight: 800;
    letter-spacing: 0.5px;
}

/* ========== PROGRESS BARS - ANIMATED GRADIENT ========== */
QProgressBar {
    border: 2px solid rgba(59, 130, 246, 0.4);
    border-radius: 16px;
    text-align: center;
    background: rgba(15, 23, 42, 0.95);
    height: 36px;
    color: #FFFFFF;
    font-size: 14px;
    font-weight: 700;
}

QProgressBar::chunk {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 #3B82F6, 
        stop:0.3 #60A5FA, 
        stop:0.6 #3B82F6,
        stop:1 #2563EB);
    border-radius: 14px;
    margin: 2px;
}

/* ========== TABLES - MODERN GRID ========== */
QTableWidget { 
    background: rgba(15, 23, 42, 0.9); 
    border: 1px solid rgba(71, 85, 105, 0.4); 
    border-radius: 16px; 
    gridline-color: rgba(51, 65, 85, 0.3); 
    selection-background-color: rgba(37, 99, 235, 0.5); 
    selection-color: #FFFFFF; 
}

QHeaderView::section { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(30, 41, 59, 0.9),
        stop:1 rgba(20, 30, 48, 0.85)); 
    color: #CBD5E1; 
    border: none; 
    border-bottom: 2px solid rgba(59, 130, 246, 0.3); 
    padding: 10px 12px;
    font-weight: 800;
    font-size: 11px;
    letter-spacing: 0.8px;
}

/* ========== CHECKBOXES - MODERN STYLE ========== */
QCheckBox { 
    spacing: 10px; 
}

QCheckBox::indicator { 
    width: 22px; 
    height: 22px; 
    border-radius: 8px; 
    border: 2px solid rgba(71, 85, 105, 0.6); 
    background: rgba(15, 23, 42, 0.8); 
}

QCheckBox::indicator:hover {
    border: 2px solid rgba(59, 130, 246, 0.8);
    background: rgba(20, 30, 50, 0.9);
}

QCheckBox::indicator:checked { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, 
        stop:0 #2563EB, 
        stop:1 #3B82F6);
    image: url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' width='22' height='22' viewBox='0 0 22 22'><path fill='none' stroke='%23FFFFFF' stroke-width='2.5' stroke-linecap='round' d='M6 11 L10 15 L16 7'/></svg>"); 
    border: 2px solid rgba(147, 197, 253, 0.9); 
}

/* ========== TYPOGRAPHY - REFINED ========== */
QLabel[h1="true"] { 
    font-family: 'Plus Jakarta Sans', 'Inter'; 
    font-size: 34px; 
    font-weight: 900; 
    letter-spacing: -0.5px;
    color: #F8FAFC;
}

QLabel[h2="true"] { 
    font-family: 'Plus Jakarta Sans', 'Inter'; 
    font-size: 26px; 
    font-weight: 800; 
    letter-spacing: -0.3px;
    color: #F1F5F9;
    margin: 6px 0;
}

QLabel[muted="true"] { 
    color: #94A3B8; 
    font-weight: 500;
}

/* ========== SCROLL BARS - SLEEK MINIMAL ========== */
QScrollBar:vertical {
    background: rgba(15, 23, 42, 0.5);
    width: 14px;
    border-radius: 7px;
    margin: 3px;
}

QScrollBar::handle:vertical {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 rgba(59, 130, 246, 0.4),
        stop:1 rgba(96, 165, 250, 0.3));
    border-radius: 6px;
    min-height: 40px;
}

QScrollBar::handle:vertical:hover {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 rgba(59, 130, 246, 0.6),
        stop:1 rgba(96, 165, 250, 0.5));
}

QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
    height: 0px;
}

QScrollBar:horizontal {
    background: rgba(15, 23, 42, 0.5);
    height: 14px;
    border-radius: 7px;
    margin: 3px;
}

QScrollBar::handle:horizontal {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(59, 130, 246, 0.4),
        stop:1 rgba(96, 165, 250, 0.3));
    border-radius: 6px;
    min-width: 40px;
}

QScrollBar::handle:horizontal:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 rgba(59, 130, 246, 0.6),
        stop:1 rgba(96, 165, 250, 0.5));
}

QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
    width: 0px;
}

/* ========== TEXT BROWSERS - ENHANCED ========== */
QTextBrowser {
    background: rgba(15, 23, 42, 0.85);
    border: 1px solid rgba(71, 85, 105, 0.4);
    border-radius: 18px;
    padding: 18px;
    color: #E2E8F0;
}

QTextBrowser:focus {
    border: 1px solid rgba(96, 165, 250, 0.6);
}

/* ========== SCROLL AREA - TRANSPARENT ========== */
QScrollArea {
    border: none;
    background: transparent;
}

QScrollArea > QWidget > QWidget {
    background: transparent;
}
"""

# ========== ANIMATION UTILITIES ==========

class AnimatedButton(QPushButton):
    """Enhanced button with hover animations and glow effects."""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setMouseTracking(True)
        self._setup_glow()
        self._setup_animations()
    
    def _setup_glow(self):
        """Add glow effect to button."""
        self.glow = QGraphicsDropShadowEffect(self)
        self.glow.setBlurRadius(0)
        self.glow.setOffset(0, 0)
        self.glow.setColor(QColor(59, 130, 246, 0))
        self.setGraphicsEffect(self.glow)
    
    def _setup_animations(self):
        """Setup property animations for smooth transitions."""
        self.default_height = 40
        self.hover_offset = -3
        
    def enterEvent(self, event):
        """Animate on hover enter."""
        self.glow_anim = QPropertyAnimation(self.glow, b"blurRadius")
        self.glow_anim.setDuration(200)
        self.glow_anim.setStartValue(self.glow.blurRadius())
        self.glow_anim.setEndValue(20)
        self.glow_anim.setEasingCurve(QEasingCurve.OutCubic)
        
        self.color_anim = QPropertyAnimation(self.glow, b"color")
        self.color_anim.setDuration(200)
        self.color_anim.setStartValue(self.glow.color())
        self.color_anim.setEndValue(QColor(59, 130, 246, 150))
        self.color_anim.setEasingCurve(QEasingCurve.OutCubic)
        
        self.glow_anim.start()
        self.color_anim.start()
        super().enterEvent(event)
    
    def leaveEvent(self, event):
        """Animate on hover leave."""
        self.glow_anim = QPropertyAnimation(self.glow, b"blurRadius")
        self.glow_anim.setDuration(200)
        self.glow_anim.setStartValue(self.glow.blurRadius())
        self.glow_anim.setEndValue(0)
        self.glow_anim.setEasingCurve(QEasingCurve.InCubic)
        
        self.color_anim = QPropertyAnimation(self.glow, b"color")
        self.color_anim.setDuration(200)
        self.color_anim.setStartValue(self.glow.color())
        self.color_anim.setEndValue(QColor(59, 130, 246, 0))
        self.color_anim.setEasingCurve(QEasingCurve.InCubic)
        
        self.glow_anim.start()
        self.color_anim.start()
        super().leaveEvent(event)


class AnimatedTile(QFrame):
    """Enhanced tile with smooth animations and micro-interactions."""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setProperty("card", True)
        self.setProperty("tile", True)
        self.setMouseTracking(True)
        
        # Store original geometry
        self._original_y = 0
        self._is_animating = False
        
        # Setup effects
        self._setup_shadow()
        self._setup_animations()
    
    def _setup_shadow(self):
        """Enhanced shadow effect."""
        try:
            self.shadow = QGraphicsDropShadowEffect(self)
            self.shadow.setBlurRadius(28)
            self.shadow.setOffset(0, 10)
            self.shadow.setColor(QColor(0, 0, 0, 60))
            self.setGraphicsEffect(self.shadow)
        except Exception:
            self.shadow = None
    
    def __del__(self):
        """Cleanup animations on deletion."""
        try:
            if hasattr(self, 'hover_group') and self.hover_group is not None:
                self.hover_group.stop()
        except (RuntimeError, AttributeError):
            pass
    
    def _setup_animations(self):
        """Setup smooth animations."""
        # Shadow animation
        self.shadow_anim = QPropertyAnimation(self.shadow, b"blurRadius")
        self.shadow_anim.setDuration(300)
        self.shadow_anim.setEasingCurve(QEasingCurve.OutCubic)
        
        self.shadow_offset_anim = QPropertyAnimation(self.shadow, b"offset")
        self.shadow_offset_anim.setDuration(300)
        self.shadow_offset_anim.setEasingCurve(QEasingCurve.OutCubic)
        
        self.shadow_color_anim = QPropertyAnimation(self.shadow, b"color")
        self.shadow_color_anim.setDuration(300)
        self.shadow_color_anim.setEasingCurve(QEasingCurve.OutCubic)
        
        # Group animations
        self.hover_group = QParallelAnimationGroup(self)
        self.hover_group.addAnimation(self.shadow_anim)
        self.hover_group.addAnimation(self.shadow_offset_anim)
        self.hover_group.addAnimation(self.shadow_color_anim)
        
        # Reset animation flag when complete
        self.hover_group.finished.connect(lambda: setattr(self, '_is_animating', False))
    
    def enterEvent(self, event):
        """Animate tile lift on hover."""
        try:
            if (not self._is_animating and 
                self.shadow is not None and 
                hasattr(self, 'hover_group') and
                self.graphicsEffect() is not None):
                
                self._is_animating = True
                
                # Stop any running animations
                if self.hover_group.state() == QParallelAnimationGroup.Running:
                    self.hover_group.stop()
                
                # Enhance shadow
                self.shadow_anim.setStartValue(self.shadow.blurRadius())
                self.shadow_anim.setEndValue(40)
                
                self.shadow_offset_anim.setStartValue(self.shadow.offset())
                self.shadow_offset_anim.setEndValue(QPoint(0, 16))
                
                self.shadow_color_anim.setStartValue(self.shadow.color())
                self.shadow_color_anim.setEndValue(QColor(59, 130, 246, 100))
                
                self.hover_group.start()
        except (RuntimeError, AttributeError):
            # Shadow effect was deleted, ignore
            pass
        
        super().enterEvent(event)

    def leaveEvent(self, event):
        """Animate tile drop on hover leave."""
        try:
            if (not self._is_animating and 
                self.shadow is not None and 
                hasattr(self, 'hover_group') and
                self.graphicsEffect() is not None):
                
                self._is_animating = True
                
                # Stop any running animations
                if self.hover_group.state() == QParallelAnimationGroup.Running:
                    self.hover_group.stop()
                
                # Reset shadow
                self.shadow_anim.setStartValue(self.shadow.blurRadius())
                self.shadow_anim.setEndValue(28)
                
                self.shadow_offset_anim.setStartValue(self.shadow.offset())
                self.shadow_offset_anim.setEndValue(QPoint(0, 10))
                
                self.shadow_color_anim.setStartValue(self.shadow.color())
                self.shadow_color_anim.setEndValue(QColor(0, 0, 0, 60))
                
                self.hover_group.start()
        except (RuntimeError, AttributeError):
            # Shadow effect was deleted, ignore
            pass
        
        super().leaveEvent(event)


class PulsingBadge(QLabel):
    """Animated badge with subtle pulsing glow."""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self._setup_pulse_animation()
    
    def _setup_pulse_animation(self):
        """Create subtle pulsing effect."""
        self.glow = QGraphicsDropShadowEffect(self)
        self.glow.setBlurRadius(12)
        self.glow.setOffset(0, 0)
        self.glow.setColor(QColor(59, 130, 246, 100))
        self.setGraphicsEffect(self.glow)
        
        # Pulse animation
        self.pulse_anim = QPropertyAnimation(self.glow, b"blurRadius")
        self.pulse_anim.setDuration(1500)
        self.pulse_anim.setStartValue(8)
        self.pulse_anim.setEndValue(16)
        self.pulse_anim.setEasingCurve(QEasingCurve.InOutSine)
        self.pulse_anim.setLoopCount(-1)  # Infinite loop
        self.pulse_anim.start()


class FadeInWidget(QWidget):
    """Widget that fades in when shown."""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.opacity_effect = QGraphicsOpacityEffect(self)
        self.setGraphicsEffect(self.opacity_effect)
        self.opacity_effect.setOpacity(0)
    
    def fade_in(self, duration=400):
        """Fade in animation."""
        self.anim = QPropertyAnimation(self.opacity_effect, b"opacity")
        self.anim.setDuration(duration)
        self.anim.setStartValue(0)
        self.anim.setEndValue(1)
        self.anim.setEasingCurve(QEasingCurve.OutCubic)
        self.anim.start()
    
    def showEvent(self, event):
        """Auto fade-in on show."""
        super().showEvent(event)
        QTimer.singleShot(50, lambda: self.fade_in())


class AnimatedProgressBar(QProgressBar):
    """Progress bar with smooth animations and glow."""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._setup_glow()
    
    def _setup_glow(self):
        """Add glowing effect to progress bar."""
        self.glow = QGraphicsDropShadowEffect(self)
        self.glow.setBlurRadius(20)
        self.glow.setOffset(0, 0)
        self.glow.setColor(QColor(59, 130, 246, 180))
        self.setGraphicsEffect(self.glow)
        
        # Pulse animation
        self.pulse = QPropertyAnimation(self.glow, b"blurRadius")
        self.pulse.setDuration(1000)
        self.pulse.setStartValue(15)
        self.pulse.setEndValue(25)
        self.pulse.setEasingCurve(QEasingCurve.InOutSine)
        self.pulse.setLoopCount(-1)
        self.pulse.start()
    
    def setValue(self, value):
        """Smooth value transition."""
        if hasattr(self, '_current_value') and self.value() != value:
            self._value_anim = QPropertyAnimation(self, b"value")
            self._value_anim.setDuration(300)
            self._value_anim.setStartValue(self.value())
            self._value_anim.setEndValue(value)
            self._value_anim.setEasingCurve(QEasingCurve.OutCubic)
            self._value_anim.start()
        else:
            super().setValue(value)
        self._current_value = value


# ========== ENHANCED VISUAL HELPERS ==========

def apply_enhanced_shadow(widget: QWidget, blur=32, offset_y=12, color="#000000", opacity=0.3):
    """Apply enhanced shadow with better depth."""
    effect = QGraphicsDropShadowEffect(widget)
    effect.setBlurRadius(blur)
    effect.setOffset(0, offset_y)
    c = QColor(color)
    c.setAlphaF(opacity)
    effect.setColor(c)
    widget.setGraphicsEffect(effect)


def apply_animated_glow(widget: QWidget, color="#3B82F6", min_blur=15, max_blur=30):
    """Apply animated glow effect that pulses."""
    glow = QGraphicsDropShadowEffect(widget)
    glow.setBlurRadius(min_blur)
    glow.setOffset(0, 0)
    c = QColor(color)
    c.setAlphaF(0.6)
    glow.setColor(c)
    widget.setGraphicsEffect(glow)
    
    # Animate the glow
    anim = QPropertyAnimation(glow, b"blurRadius")
    anim.setDuration(1500)
    anim.setStartValue(min_blur)
    anim.setEndValue(max_blur)
    anim.setEasingCurve(QEasingCurve.InOutSine)
    anim.setLoopCount(-1)
    anim.start()
    
    # Store animation to prevent garbage collection
    widget._glow_animation = anim
    return anim


def apply_selection_glow(widget: QWidget, selected: bool):
    """Apply or remove selection glow effect."""
    if selected:
        glow = QGraphicsDropShadowEffect(widget)
        glow.setBlurRadius(25)
        glow.setOffset(0, 0)
        glow.setColor(QColor(59, 130, 246, 150))
        widget.setGraphicsEffect(glow)
        
        # Animate glow intensity
        anim = QPropertyAnimation(glow, b"color")
        anim.setDuration(300)
        anim.setStartValue(QColor(59, 130, 246, 0))
        anim.setEndValue(QColor(59, 130, 246, 150))
        anim.setEasingCurve(QEasingCurve.OutCubic)
        anim.start()
        widget._selection_anim = anim
    else:
        # Fade out glow
        if widget.graphicsEffect():
            effect = widget.graphicsEffect()
            anim = QPropertyAnimation(effect, b"color")
            anim.setDuration(300)
            anim.setStartValue(QColor(59, 130, 246, 150))
            anim.setEndValue(QColor(59, 130, 246, 0))
            anim.setEasingCurve(QEasingCurve.InCubic)
            anim.finished.connect(lambda: widget.setGraphicsEffect(None))
            anim.start()
            widget._selection_anim = anim


def create_staggered_fade_in(widgets: list, delay_ms=80):
    """Create staggered fade-in animation for multiple widgets."""
    animations = []
    
    for i, widget in enumerate(widgets):
        opacity = QGraphicsOpacityEffect(widget)
        widget.setGraphicsEffect(opacity)
        opacity.setOpacity(0)
        
        anim = QPropertyAnimation(opacity, b"opacity")
        anim.setDuration(400)
        anim.setStartValue(0)
        anim.setEndValue(1)
        anim.setEasingCurve(QEasingCurve.OutCubic)
        
        # Delay each animation
        QTimer.singleShot(i * delay_ms, anim.start)
        animations.append(anim)
    
    return animations


def animate_widget_property(widget: QWidget, property_name: bytes, start_val, end_val, 
                            duration=300, easing=QEasingCurve.OutCubic):
    """Generic property animation helper."""
    anim = QPropertyAnimation(widget, property_name)
    anim.setDuration(duration)
    anim.setStartValue(start_val)
    anim.setEndValue(end_val)
    anim.setEasingCurve(easing)
    anim.start()
    return anim


# ========== ENHANCED TILE WITH ALL EFFECTS ==========

COMPACT_WIDTH = 310

class EnhancedTestCaseTile(QFrame):
    def __init__(self, case, on_edit, on_remove, on_toggle_expand, parent=None):
        super().__init__(parent)
        self.setProperty("card", True)
        self.setProperty("tile", True)
        apply_soft_shadow(self, radius=28, dy=10, opacity=0.18)
        
        self.case=case; self.on_edit=on_edit; self.on_remove=on_remove; self.on_toggle_expand=on_toggle_expand
        self.expanded=False
        self.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Minimum)
        self.setMinimumWidth(COMPACT_WIDTH); self.setMaximumWidth(COMPACT_WIDTH)
        
        self._build_ui()
        
    def _build_ui(self):
        """Build tile UI."""
        root = QVBoxLayout(self)
        root.setContentsMargins(14, 14, 14, 14)
        root.setSpacing(8)
        
        # Header with Test Case ID, Priority, and Checkbox
        hdr = QHBoxLayout()
        hdr.setSpacing(8)
        
        self.id_lbl = QLabel("TC-001")
        self.id_lbl.setStyleSheet("font-size:16px; font-weight:800; background:transparent;")
        
        # Map priority numbers to text
        priority_map = {1: "High", 2: "Medium", 3: "Low", 4: "Low"}
        priority_text = priority_map.get(self.case.priority, "Medium")
        self.priority_pill = QLabel(priority_text)
        self.priority_pill.setProperty("priority", priority_text)
        self.priority_pill.setObjectName("pill")
        
        self.checkbox = QCheckBox()
        self.checkbox.setStyleSheet("background:transparent;")
        self.checkbox.stateChanged.connect(self._on_selection_changed)
        
        hdr.addWidget(self.id_lbl, 1)
        hdr.addWidget(self.priority_pill, 0, Qt.AlignRight)
        hdr.addWidget(self.checkbox, 0, Qt.AlignRight)
        root.addLayout(hdr)
        
        # Title
        self.title_lbl = QLabel(self.case.title)
        self.title_lbl.setWordWrap(True)
        self.title_lbl.setStyleSheet("font-size:14px; font-weight:600; background:transparent;")
        root.addWidget(self.title_lbl)
        
        # Test type badge
        test_type_badge = QLabel(self.case.test_type)
        test_type_badge.setStyleSheet("font-size:10px; color:#A7B3C6; font-weight:600; background:transparent;")
        root.addWidget(test_type_badge)
        
        # Expand button
        self.expand_btn = QPushButton("Show details ▾")
        self.expand_btn.setStyleSheet("padding: 8px 16px; min-height: 24px;")
        self.expand_btn.clicked.connect(self.toggle_expand)
        root.addWidget(self.expand_btn, 0, Qt.AlignRight)
        
        # Details section (collapsible)
        self.details = QFrame()
        d = QVBoxLayout(self.details)
        d.setContentsMargins(10, 10, 10, 10)
        d.setSpacing(8)
        
        self.steps_tbl = QTableWidget(0, 3)
        self.steps_tbl.setHorizontalHeaderLabels(["#", "Action", "Expected"])
        self.steps_tbl.verticalHeader().setVisible(False)
        self.steps_tbl.horizontalHeader().setStretchLastSection(True)
        self.steps_tbl.setEditTriggers(QTableWidget.NoEditTriggers)
        self.refresh_steps_table()
        d.addWidget(self.steps_tbl)
        
        # Action buttons
        foot = QHBoxLayout()
        self.edit_btn = QPushButton("Edit")
        self.remove_btn = QPushButton("Remove")
        self.edit_btn.clicked.connect(lambda: self.on_edit(self))
        self.remove_btn.clicked.connect(lambda: self.on_remove(self))
        foot.addWidget(self.edit_btn)
        foot.addWidget(self.remove_btn)
        foot.addStretch(1)
        d.addLayout(foot)
        
        self.details.setVisible(False)
        root.addWidget(self.details)
    
    def _on_selection_changed(self, state):
        """Handle checkbox selection."""
        selected = state == Qt.Checked
        self.setProperty("selected", selected)
        self.style().unpolish(self)
        self.style().polish(self)
    
    def refresh_steps_table(self):
        """Refresh steps table content."""
        steps = self.case.steps or []
        self.steps_tbl.setRowCount(len(steps))
        for i, s in enumerate(steps, 1):
            num = QTableWidgetItem(str(i))
            num.setTextAlignment(Qt.AlignCenter)
            self.steps_tbl.setItem(i-1, 0, num)
            self.steps_tbl.setItem(i-1, 1, QTableWidgetItem(s.action))
            self.steps_tbl.setItem(i-1, 2, QTableWidgetItem(s.expected))
        self.steps_tbl.resizeColumnsToContents()
    
    def set_test_case_id(self, tc_id: int):
        """Set test case ID with requirement mapping."""
        req_id = self.case.requirement_id or "UNMAPPED"
        self.id_lbl.setText(f"{req_id} : TC-{tc_id:03d}")
    
    def toggle_expand(self):
        """Toggle expansion - EXACTLY like old version."""
        self.expanded = not self.expanded
        self.details.setVisible(self.expanded)
        self.expand_btn.setText("Hide details ▲" if self.expanded else "Show details ▾")
        self.on_toggle_expand(self, self.expanded)
    
    def sizeHint(self):
        return QSize(self.maximumWidth(), 160 if not self.expanded else 320)
    
    def enterEvent(self, event):
        """Add subtle hover effect."""
        self.setStyleSheet("""
            QFrame[tile="true"] {
                background:qlineargradient(x1:0,y1:0,x2:0,y2:1, 
                    stop:0 rgba(22,28,42,1), stop:1 rgba(16,22,35,.98));
                border:1px solid rgba(59,130,246,.4);
            }
        """)
        super().enterEvent(event)

    def leaveEvent(self, event):
        """Remove hover effect."""
        self.setStyleSheet("")  # Reset to stylesheet defaults
        super().leaveEvent(event)

class ExportTestCaseTile(QWidget):
    """Compact tile for export page - shows test case summary with checkbox."""
    def __init__(self, case: TestCase, test_id: int):
        super().__init__()
        self.case = case
        self.test_id = test_id
        
        # Main layout
        main = QVBoxLayout(self)
        main.setContentsMargins(0, 0, 0, 0)
        main.setSpacing(0)
        
        # Container with styling
        container = QFrame()
        container.setObjectName("exportTileContainer")
        container.setStyleSheet("""
            QFrame#exportTileContainer {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 rgba(30, 41, 59, 0.8),
                    stop:1 rgba(15, 23, 42, 0.8));
                border: 1px solid rgba(71, 85, 105, 0.3);
                border-radius: 16px;
            }
            QFrame#exportTileContainer:hover {
                border: 1px solid rgba(96, 165, 250, 0.5);
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 rgba(30, 41, 59, 0.95),
                    stop:1 rgba(15, 23, 42, 0.95));
            }
            QFrame#exportTileContainer QLabel {
                background: transparent;
            }
            QFrame#exportTileContainer QCheckBox {
                background: transparent;
            }
        """)
        
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)
        
        # Header row: checkbox + ID + priority
        header = QHBoxLayout()
        
        self.checkbox = QCheckBox()
        self.checkbox.setChecked(True)
        header.addWidget(self.checkbox)
        
        # Test case ID
        self.id_label = QLabel()
        self.id_label.setProperty("muted", True)
        self.id_label.setStyleSheet("font-weight: 600; font-size: 13px;")
        header.addWidget(self.id_label)
        
        header.addStretch(1)
        
        # Priority badge
        self.priority_badge = QLabel()
        self.priority_badge.setProperty("priority", 
            "High" if case.priority == 1 else "Medium" if case.priority == 2 else "Low")
        priority_text = "High" if case.priority == 1 else "Medium" if case.priority == 2 else "Low"
        self.priority_badge.setText(priority_text)
        self.priority_badge.setStyleSheet("""
            background: transparent;
            padding: 4px 12px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 700;
        """)
        header.addWidget(self.priority_badge)
        
        layout.addLayout(header)
        
        # Title
        title_label = QLabel(case.title)
        title_label.setWordWrap(True)
        title_label.setStyleSheet("font-weight: 600; font-size: 14px; color: #E2E8F0;")
        layout.addWidget(title_label)
        
        # Test type + tags
        meta_row = QHBoxLayout()
        
        type_badge = QLabel(case.test_type)
        type_badge.setStyleSheet("""
            background: transparent;
            color: #93C5FD;
            padding: 3px 10px;
            border-radius: 8px;
            font-size: 11px;
            font-weight: 600;
        """)
        meta_row.addWidget(type_badge)
        
        # Show first 2 tags
        for tag in case.tags[:2]:
            tag_badge = QLabel(tag)
            tag_badge.setStyleSheet("""
                background: transparent;
                color: #C4B5FD;
                padding: 3px 10px;
                border-radius: 8px;
                font-size: 10px;
            """)
            meta_row.addWidget(tag_badge)
        
        if len(case.tags) > 2:
            more_label = QLabel(f"+{len(case.tags) - 2}")
            more_label.setProperty("muted", True)
            more_label.setStyleSheet("font-size: 10px;")
            meta_row.addWidget(more_label)
        
        meta_row.addStretch(1)
        layout.addLayout(meta_row)
        
        # Step count
        steps_label = QLabel(f"{len(case.steps)} step(s)")
        steps_label.setProperty("muted", True)
        steps_label.setStyleSheet("font-size: 12px;")
        layout.addWidget(steps_label)
        
        main.addWidget(container)
        
        # Set requirement ID
        self.set_test_case_id()
    
    def set_test_case_id(self):
        """Set test case ID with requirement mapping."""
        req_id = self.case.requirement_id or "UNMAPPED"
        self.id_label.setText(f"{req_id} : TC-{self.test_id:03d}")

# ========== USAGE INSTRUCTIONS ==========
"""
To integrate these enhancements into your existing code:

1. Replace NEBULA_QSS variable with the one above
2. Add the new imports at the top of your file
3. Replace TestCaseTile class with EnhancedTestCaseTile
4. Update any QPushButton instances you want animated to use AnimatedButton
5. Replace your progress bars with AnimatedProgressBar
6. Use apply_enhanced_shadow() instead of apply_soft_shadow()

Example replacements in your MainWindow.__init__:

# Old:
self.cases.progress_bar = QProgressBar()

# New:
self.cases.progress_bar = AnimatedProgressBar()

# Old:
apply_soft_shadow(self.topbar)

# New:
apply_enhanced_shadow(self.topbar, blur=40, offset_y=16, opacity=0.25)

# For tile creation in populate_case_tiles():
# Old:
tile = TestCaseTile(case, on_edit=..., on_remove=..., on_toggle_expand=...)

# New:
tile = EnhancedTestCaseTile(case, on_edit=..., on_remove=..., on_toggle_expand=...)

# Add staggered fade-in for tiles after creation:
tiles = [...]  # your list of tiles
create_staggered_fade_in(tiles, delay_ms=60)
"""

# ---------- FlowLayout for test case tiles ----------
# FlowLayout - EXACTLY as the old working version
class FlowLayout(QLayout):
    def __init__(self, parent=None, margin=0, hspacing=12, vspacing=12):
        super().__init__(parent); self.itemList=[]; self.setContentsMargins(margin, margin, margin, margin)
        self._hspace=hspacing; self._vspace=vspacing
    def addItem(self, item): self.itemList.append(item)
    def addWidget(self, w): super().addWidget(w)
    def count(self): return len(self.itemList)
    def itemAt(self, idx): return self.itemList[idx] if 0<=idx<len(self.itemList) else None
    def takeAt(self, idx): return self.itemList.pop(idx) if 0<=idx<len(self.itemList) else None
    def expandingDirections(self): return Qt.Orientations(Qt.Orientation(0))
    def hasHeightForWidth(self): return True
    def heightForWidth(self, w): return self._doLayout(QRect(0,0,w,0), True)
    def setGeometry(self, rect): super().setGeometry(rect); self._doLayout(rect, False)
    def sizeHint(self): return self.minimumSize()
    def minimumSize(self):
        s=QSize()
        for i in self.itemList: s=s.expandedTo(i.minimumSize())
        m=self.contentsMargins(); s+=QSize(m.left()+m.right(), m.top()+m.bottom()); return s
    def _doLayout(self, rect, testOnly):
        x=rect.x(); y=rect.y(); lineH=0
        l,t,r,b=self.getContentsMargins(); x+=l; y+=t; maxw=rect.width()-l-r
        for item in self.itemList:
            spaceX=self._hspace; spaceY=self._vspace
            nextX=x+item.sizeHint().width()+spaceX
            if nextX - spaceX > rect.x()+maxw and lineH>0:
                x=rect.x()+l; y=y+lineH+spaceY; nextX=x+item.sizeHint().width()+spaceX; lineH=0
            if not testOnly: item.setGeometry(QRect(QPoint(x,y), item.sizeHint()))
            x=nextX; lineH=max(lineH, item.sizeHint().height())
        return y + lineH + b

# ---------- UI pieces ----------
class TopBar(QWidget):
    def __init__(self, title="Sign In", parent=None, on_back=None):
        super().__init__(parent); self.setObjectName("TopBar")
        h=QHBoxLayout(self); h.setContentsMargins(12, 8, 12, 8) # Symmetrical margins

        self.back=QPushButton("← Back")
        self.back.setMinimumWidth(100)
        self.back.setMinimumHeight(40)
        self.back.clicked.connect(on_back or (lambda: None))

        self.title=QLabel(title); self.title.setProperty("h1", True)
        self.ai_status=QLabel("AI: OFF"); self.ai_status.setStyleSheet("padding:6px 10px; background:#1a2740; border-radius:12px;")
        self.avatar=QLabel("CB"); self.avatar.setFixedSize(34,34); self.avatar.setStyleSheet("background:#1a2740; border-radius:17px; qproperty-alignment: 'AlignCenter';")

        # Add widgets with vertical center alignment
        h.addWidget(self.back, 0, Qt.AlignVCenter)
        h.addSpacing(8)
        h.addWidget(self.title, 0, Qt.AlignVCenter)
        h.addStretch(1)
        h.addWidget(self.ai_status, 0, Qt.AlignVCenter)
        h.addSpacing(12)
        h.addWidget(self.avatar, 0, Qt.AlignVCenter)

class TopTabs(QWidget):
    """Now includes 'Feature to analyze' and 'AI Feature Overview' tabs."""
    def __init__(self, parent=None):
        super().__init__(parent); self.setObjectName("TopTabs")
        h=QHBoxLayout(self); h.setContentsMargins(0,6,0,6); h.addStretch(1)
        self.btn_feature=QPushButton("Feature to analyze")
        self.btn_overview=QPushButton("AI Feature Overview")
        self.btn_ticket =QPushButton("Ticket to analyze")
        self.btn_readiness=QPushButton("Ticket Readiness")
        self.btn_cases  =QPushButton("Test Cases")
        self.btn_export =QPushButton("Export")
        for b in (self.btn_feature,self.btn_overview,self.btn_ticket,self.btn_readiness,self.btn_cases,self.btn_export):
            b.setCheckable(True); b.setProperty("active", False); h.addWidget(b)
        h.addStretch(1)
    def set_active(self, btn: QPushButton):
        for b in self.findChildren(QPushButton):
            b.setProperty("active", b is btn); b.style().unpolish(b); b.style().polish(b)
    def set_enabled_after_login(self, enabled: bool):
        for b in (self.btn_feature,self.btn_overview,self.btn_ticket,self.btn_readiness,self.btn_cases,self.btn_export): b.setEnabled(enabled)

# ---------- Dialogs ----------
class StepEditor(QDialog):
    def __init__(self, case: TestCase, parent=None):
        super().__init__(parent); self.setWindowTitle("Edit Test Case"); self.case=case
        v=QVBoxLayout(self); form=QFormLayout()
        self.title_edit=QLineEdit(case.title); self.priority_spin=QSpinBox(); self.priority_spin.setRange(1,4); self.priority_spin.setValue(case.priority)
        self.tags_edit=QLineEdit(",".join(case.tags))
        form.addRow("Title", self.title_edit); form.addRow("Priority (1-4)", self.priority_spin); form.addRow("Tags (comma)", self.tags_edit); v.addLayout(form)
        self.table=QTableWidget(len(case.steps),2); self.table.setHorizontalHeaderLabels(["Action","Expected Result"]); self.table.horizontalHeader().setStretchLastSection(True)
        for r,s in enumerate(case.steps): self.table.setItem(r,0,QTableWidgetItem(s.action)); self.table.setItem(r,1,QTableWidgetItem(s.expected))
        v.addWidget(self.table)
        btns=QHBoxLayout(); add_btn=QPushButton("Add Step"); del_btn=QPushButton("Delete Step"); btns.addWidget(add_btn); btns.addWidget(del_btn); v.addLayout(btns)
        add_btn.clicked.connect(lambda: (self.table.insertRow(self.table.rowCount()), None)); del_btn.clicked.connect(lambda: self.table.removeRow(self.table.currentRow()))
        bb=QDialogButtonBox(QDialogButtonBox.Ok|QDialogButtonBox.Cancel); bb.accepted.connect(self.accept); bb.rejected.connect(self.reject); v.addWidget(bb)
    def accept(self):
        self.case.title=self.title_edit.text().strip() or self.case.title
        self.case.priority=int(self.priority_spin.value())
        self.case.tags=[t.strip() for t in self.tags_edit.text().split(',') if t.strip()]
        steps=[]
        for r in range(self.table.rowCount()):
            a=(self.table.item(r,0).text() if self.table.item(r,0) else "").strip()
            e=(self.table.item(r,1).text() if self.table.item(r,1) else "").strip()
            if a or e: steps.append(TestStep(a or "(action)", e or "(expected)"))
        if steps: self.case.steps=steps
        super().accept()

class RawAIDialog(QDialog):
    def __init__(self, raw_text: str, ai_error: Optional[str], used_fallback: bool, parent=None):
        super().__init__(parent); self.setWindowTitle("Raw AI Output")
        v=QVBoxLayout(self); info=QLabel(f"Source: {'Fallback' if used_fallback else 'OpenAI JSON'}" + (f" | Error: {ai_error}" if ai_error else "")); v.addWidget(info)
        txt=QPlainTextEdit(); txt.setPlainText(raw_text or "(empty)"); txt.setReadOnly(True); v.addWidget(txt)
        bb=QDialogButtonBox(QDialogButtonBox.Close); bb.rejected.connect(self.reject); bb.accepted.connect(self.accept); v.addWidget(bb)

class TraceabilityDialog(QDialog):
    def __init__(self, test_cases: List[TestCase], parent=None):
        super().__init__(parent)
        self.setWindowTitle("Requirement Traceability Matrix")
        self.setMinimumSize(1000, 600)
        
        v = QVBoxLayout(self)
        
        # Header
        hdr = QLabel("📊 Requirement → Test Case Traceability")
        hdr.setProperty("h2", True)
        hdr.setStyleSheet("background:transparent;")
        v.addWidget(hdr)
        
        # Build traceability map
        req_map = {}
        for tc in test_cases:
            req_id = tc.requirement_id or "UNMAPPED"
            req_desc = tc.requirement_desc or "No requirement specified"
            if req_id not in req_map:
                req_map[req_id] = {"desc": req_desc, "cases": []}
            req_map[req_id]["cases"].append(tc)
        
        # Scroll area for requirements
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        container = QWidget()
        container_layout = QVBoxLayout(container)
        
        # Display each requirement with its test cases
        for req_id in sorted(req_map.keys()):
            req_data = req_map[req_id]
            
            # Requirement header
            req_frame = QFrame()
            req_frame.setProperty("card", True)
            req_frame.setStyleSheet("background: rgba(30, 41, 59, 0.8); border: 1px solid rgba(59, 130, 246, 0.4); border-radius: 12px; padding: 12px; margin: 4px;")
            req_layout = QVBoxLayout(req_frame)
            
            req_title = QLabel(f"<b>{req_id}</b>: {req_data['desc'][:120]}")
            req_title.setWordWrap(True)
            req_title.setStyleSheet("background:transparent; color: #93C5FD; font-size: 14px;")
            req_layout.addWidget(req_title)
            
            # Test case count by type
            pos_count = sum(1 for tc in req_data["cases"] if tc.test_type == "Positive")
            neg_count = sum(1 for tc in req_data["cases"] if tc.test_type == "Negative")
            edge_count = sum(1 for tc in req_data["cases"] if tc.test_type == "Edge Case")
            
            stats = QLabel(f"Coverage: {len(req_data['cases'])} test cases | ✓ {pos_count} Positive | ✗ {neg_count} Negative | ⚠ {edge_count} Edge")
            stats.setStyleSheet("background:transparent; color: #A7B3C6; font-size: 12px; padding: 4px 0;")
            req_layout.addWidget(stats)
            
            # List test cases
            for tc in req_data["cases"]:
                tc_label = QLabel(f"   • [{tc.test_type}] {tc.title}")
                tc_label.setWordWrap(True)
                tc_label.setStyleSheet("background:transparent; color: #E2E8F0; font-size: 11px; padding-left: 20px;")
                req_layout.addWidget(tc_label)
            
            container_layout.addWidget(req_frame)
        
        container_layout.addStretch(1)
        scroll.setWidget(container)
        v.addWidget(scroll)
        
        # Close button
        bb = QDialogButtonBox(QDialogButtonBox.Close)
        bb.rejected.connect(self.reject)
        v.addWidget(bb)

# ---------- Pages ----------
class LoginPage(QWidget):
    def __init__(self):
        super().__init__(); root=QVBoxLayout(self); root.setContentsMargins(0,0,0,0); root.addStretch(1)
        container=QFrame(); container.setProperty("card", True); container.setMinimumWidth(1080); container.setMaximumWidth(1080); apply_soft_shadow(container, radius=40, dy=18, opacity=0.22)
        c=QVBoxLayout(container); c.setContentsMargins(28,28,28,28); c.setSpacing(14)
        title=QLabel("Log in details"); title.setProperty("h1", True); title.setStyleSheet("background:transparent;"); c.addWidget(title, 0, Qt.AlignHCenter)
        subtitle=QLabel("Sign into Jira"); subtitle.setProperty("h2", True); subtitle.setStyleSheet("background:transparent;"); c.addWidget(subtitle, 0, Qt.AlignHCenter)
        form=QFormLayout(); form.setLabelAlignment(Qt.AlignRight)
        
        email_label = QLabel("Email")
        email_label.setStyleSheet("background:transparent;")
        self.email=QLineEdit(); self.email.setPlaceholderText("you@example.com")
        self.email.setText(os.getenv("JIRA_EMAIL", ""))
        
        token_label = QLabel("API Token")
        token_label.setStyleSheet("background:transparent;")
        self.token=QLineEdit(); self.token.setPlaceholderText("API Token"); self.token.setEchoMode(QLineEdit.Password)
        self.token.setText(os.getenv("JIRA_API_TOKEN", ""))
        
        openai_label = QLabel("OpenAI Key")
        openai_label.setStyleSheet("background:transparent;")
        self.openai=QLineEdit(); self.openai.setPlaceholderText("OpenAI API Key (required for AI)"); self.openai.setEchoMode(QLineEdit.Password)
        self.openai.setText(os.getenv("OPENAI_API_KEY", ""))
        
        from PySide6.QtWidgets import QComboBox
        model_label = QLabel("OpenAI Model")
        model_label.setStyleSheet("background:transparent;")
        self.model=QComboBox()
        # Add GPT-5 at the top
        self.model.addItems([
            "gpt-5",           # ← New GPT-5 model
            "gpt-4o", 
            "gpt-4o-mini", 
            "gpt-4-turbo", 
            "gpt-3.5-turbo", 
            "o1-preview", 
            "o1-mini"
        ])
        default_model = os.getenv("OPENAI_MODEL", "gpt-4o")  # ← Default to GPT-5
        self.model.setCurrentText(default_model)
        
        form.addRow(email_label, self.email); form.addRow(token_label, self.token); form.addRow(openai_label, self.openai); form.addRow(model_label, self.model)
        c.addLayout(form)
        c.addSpacing(20)
        
        self.next_btn = AnimatedButton("Continue"); self.next_btn.setObjectName("accent"); c.addWidget(self.next_btn)
        root.addWidget(container, 0, Qt.AlignHCenter); root.addStretch(2)
    
    def has_all_credentials(self) -> bool:
        """Check if all required credentials are filled."""
        return bool(self.email.text().strip() and self.token.text().strip())

class FeaturePage(QWidget):
    """Paste an Epic or Initiative key, load full hierarchy, keep Feature Context."""
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        hdr=QLabel("Feature to analyze"); hdr.setProperty("h2", True); v.addWidget(hdr)

        row=QHBoxLayout()
        self.key_edit=QLineEdit(); self.key_edit.setPlaceholderText("EPIC-123 or INIT-456")
        self.fetch_btn=QPushButton("Load Feature")
        row.addWidget(QLabel("Epic or Initiative Key")); row.addWidget(self.key_edit); row.addWidget(self.fetch_btn)
        v.addLayout(row)
        
        # Manual Epic entry for Initiatives
        self.manual_epic_row=QHBoxLayout()
        self.manual_epic_row.addWidget(QLabel("Manual Epic Keys (comma-separated):"))
        self.manual_epic_keys=QLineEdit(); self.manual_epic_keys.setPlaceholderText("UEX-31, UEX-38")
        self.load_epics_btn=QPushButton("Load Epics")
        self.load_epics_btn.setObjectName("accent")
        self.manual_epic_row.addWidget(self.manual_epic_keys)
        self.manual_epic_row.addWidget(self.load_epics_btn)
        v.addLayout(self.manual_epic_row)
        
        # Initially hide manual epic entry
        self.manual_epic_keys.setVisible(False)
        self.load_epics_btn.setVisible(False)
        for i in range(self.manual_epic_row.count()):
            widget = self.manual_epic_row.itemAt(i).widget()
            if widget and isinstance(widget, QLabel):
                widget.setVisible(False)
        
        # Epic loading status
        self.epic_status=QLabel("")
        self.epic_status.setProperty("muted", True)
        self.epic_status.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        self.epic_status.setVisible(False)
        v.addWidget(self.epic_status)

        # Feature details section
        self.viewer=QTextBrowser(); self.viewer.setOpenExternalLinks(True)
        self.viewer.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:16px; padding:12px;")
        v.addWidget(self.viewer)
        
        row3=QHBoxLayout()
        self.to_overview_btn=QPushButton("Next → AI Feature Overview"); self.to_overview_btn.setObjectName("accent"); self.to_overview_btn.setEnabled(False)
        row3.addStretch(1); row3.addWidget(self.to_overview_btn)
        v.addLayout(row3)
    
    def show_manual_epic_entry(self):
        """Show the manual epic entry fields."""
        self.manual_epic_keys.setVisible(True)
        self.load_epics_btn.setVisible(True)
        for i in range(self.manual_epic_row.count()):
            widget = self.manual_epic_row.itemAt(i).widget()
            if widget and isinstance(widget, QLabel):
                widget.setVisible(True)
    
    def hide_manual_epic_entry(self):
        """Hide the manual epic entry fields."""
        self.manual_epic_keys.setVisible(False)
        self.load_epics_btn.setVisible(False)
        for i in range(self.manual_epic_row.count()):
            widget = self.manual_epic_row.itemAt(i).widget()
            if widget and isinstance(widget, QLabel):
                widget.setVisible(False)

class TicketPage(QWidget):
    def __init__(self):
        super().__init__(); v=QVBoxLayout(self); v.setSpacing(10)
        hdr=QLabel("Ticket to analyze"); hdr.setProperty("h2", True); v.addWidget(hdr)
        
        # Jira key input and fetch button with dropdown
        row=QHBoxLayout()
        row.addWidget(QLabel("Jira Key"))
        self.key_edit=QLineEdit()
        self.key_edit.setPlaceholderText("ABC-123 (If no ticket exists, select the dropdown button and select Create Test Ticket(s) to create some!) ")
        row.addWidget(self.key_edit)
        
        # Create split button with dropdown menu
        self.fetch_btn=QPushButton("Analyze Existing Ticket")
        self.fetch_btn.setObjectName("primary")
        
        # Create dropdown menu
        from PySide6.QtWidgets import QMenu
        from PySide6.QtGui import QAction
        self.ticket_menu = QMenu(self)
        
        self.action_analyze = QAction("Analyze Existing Ticket", self)
        self.action_create = QAction("Create Test Ticket(s)", self)
        self.action_view_created = QAction("View Created Tickets", self)
        
        self.ticket_menu.addAction(self.action_analyze)
        self.ticket_menu.addAction(self.action_create)
        self.ticket_menu.addAction(self.action_view_created)
        
        # Initially hide "View Created Tickets" until tickets are generated
        self.action_view_created.setVisible(False)
        
        # Dropdown button
        self.dropdown_btn = QPushButton("▼")
        self.dropdown_btn.setObjectName("primary")  # Match main button style
        self.dropdown_btn.setFixedWidth(50)  # Wider to be clearly rectangular
        self.dropdown_btn.setFixedHeight(42)  # Match standard button height
        self.dropdown_btn.setStyleSheet("""
            QPushButton {
                border-radius: 12px;
                padding: 0px;
                font-size: 12px;
                text-align: center;
            }
            QPushButton::menu-indicator {
                image: none;
                width: 0px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 rgba(96, 165, 250, 0.5),
                    stop:1 rgba(59, 130, 246, 0.4));
            }
        """)
        self.dropdown_btn.setMenu(self.ticket_menu)
        
        row.addWidget(self.fetch_btn)
        row.addWidget(self.dropdown_btn)
        v.addLayout(row)
        
        self.viewer=QTextBrowser(); self.viewer.setOpenExternalLinks(True)
        self.viewer.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:16px; padding:12px;")
        v.addWidget(self.viewer)
        actions=QHBoxLayout()
        self.clear_btn=QPushButton("Clear / New Ticket")
        self.assess_btn=QPushButton("Next → Assess Ticket Readiness"); self.assess_btn.setObjectName("accent"); self.assess_btn.setEnabled(False)
        actions.addWidget(self.clear_btn); actions.addStretch(1); actions.addWidget(self.assess_btn)
        v.addLayout(actions)

class CreateTestTicketsPage(QWidget):
    """New page for creating test tickets from Epic context."""
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        
        # Header
        hdr=QLabel("Create Test Ticket(s)"); hdr.setProperty("h2", True)
        v.addWidget(hdr)
        
        # Epic info banner
        self.epic_info=QLabel("No feature loaded"); self.epic_info.setProperty("muted", True)
        self.epic_info.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        v.addWidget(self.epic_info)
        
        # Progress bar
        self.progress_bar = AnimatedProgressBar()
        self.progress_bar.setVisible(False)
        v.addWidget(self.progress_bar)
        
        # Status/Results area
        self.status_label = QLabel("Click 'Analyze and Recommend' to begin...")
        self.status_label.setProperty("muted", True)
        self.status_label.setWordWrap(True)  # Enable word wrapping for long strategy text
        self.status_label.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        v.addWidget(self.status_label)
        
        # Scroll area for splits (will be populated dynamically)
        self.splits_scroll = QScrollArea()
        self.splits_scroll.setWidgetResizable(True)
        self.splits_scroll.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:16px;")
        self.splits_widget = QWidget()
        self.splits_layout = QVBoxLayout(self.splits_widget)
        self.splits_scroll.setWidget(self.splits_widget)
        self.splits_scroll.setVisible(False)
        v.addWidget(self.splits_scroll)
        
        # Action buttons
        actions=QHBoxLayout()
        self.back_btn=QPushButton("← Back to Feature")
        self.analyze_btn=QPushButton("Analyze and Recommend"); self.analyze_btn.setObjectName("accent")
        self.generate_btn=QPushButton("Generate Test Ticket(s)"); self.generate_btn.setObjectName("accent")
        self.generate_btn.setVisible(False)
        actions.addWidget(self.back_btn)
        actions.addStretch(1)
        actions.addWidget(self.analyze_btn)
        actions.addWidget(self.generate_btn)
        v.addLayout(actions)
    
    def clear_splits(self):
        """Clear all split widgets."""
        while self.splits_layout.count():
            child = self.splits_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
    
    def add_split_widget(self, split_data: Dict, index: int):
        """Add a widget for a recommended split."""
        frame = QFrame()
        frame.setProperty("card", True)
        frame.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:12px; padding:12px;")
        
        layout = QVBoxLayout(frame)
        
        # Title
        title = QLabel(f"{index}. Testing - {split_data.get('functional_area', 'Unknown')}")
        title.setStyleSheet("font-size:16px; font-weight:600; color:#EAECEF;")
        layout.addWidget(title)
        
        # Details
        child_tickets = split_data.get('child_tickets', [])
        ac_count = split_data.get('estimated_ac_count', 0)
        rationale = split_data.get('rationale', '')
        
        details = QLabel(
            f"Covers: {', '.join(child_tickets[:5])}" + 
            (f" (+{len(child_tickets)-5} more)" if len(child_tickets) > 5 else "") +
            f"\nEstimated AC: {ac_count}\nRationale: {rationale}"
        )
        details.setProperty("muted", True)
        details.setWordWrap(True)
        layout.addWidget(details)
        
        self.splits_layout.addWidget(frame)

class TestTicketsResultsPage(QWidget):
    """Page showing generated test tickets with analysis options."""
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        
        # Header
        hdr=QLabel("Test Tickets Created"); hdr.setProperty("h2", True)
        v.addWidget(hdr)
        
        # Summary banner
        self.summary_label=QLabel("No tickets created yet"); self.summary_label.setProperty("muted", True)
        self.summary_label.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        v.addWidget(self.summary_label)
        
        # Progress bar (for sequential generation)
        self.progress_bar = AnimatedProgressBar()
        self.progress_bar.setVisible(False)
        v.addWidget(self.progress_bar)
        
        # Scroll area for ticket cards
        self.tickets_scroll = QScrollArea()
        self.tickets_scroll.setWidgetResizable(True)
        self.tickets_scroll.setStyleSheet("background:transparent; border:none;")
        self.tickets_widget = QWidget()
        self.tickets_layout = QVBoxLayout(self.tickets_widget)
        self.tickets_scroll.setWidget(self.tickets_widget)
        v.addWidget(self.tickets_scroll)
        
        # Action buttons
        actions=QHBoxLayout()
        self.back_btn=QPushButton("← Back to Feature")
        self.export_all_btn=QPushButton("Export All to Jira Format")
        actions.addWidget(self.back_btn)
        actions.addStretch(1)
        actions.addWidget(self.export_all_btn)
        v.addLayout(actions)
    
    def clear_tickets(self):
        """Clear all ticket widgets."""
        while self.tickets_layout.count():
            child = self.tickets_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
    
    def add_ticket_widget(self, ticket: GeneratedTestTicket, analyze_callback, copy_callback, regenerate_callback):
        """Add a widget for a generated test ticket."""
        frame = QFrame()
        frame.setProperty("card", True)
        frame.setObjectName(f"ticket_card_{ticket.id}")  # Add object name for easy lookup
        
        # Different styling based on status - frame only, buttons inherit from global
        if ticket.analyzed:
            frame.setStyleSheet("""
                QFrame {
                    background:#0E1E0E; 
                    border:2px solid #22C55E; 
                    border-radius:12px; 
                    padding:16px;
                }
                QPushButton { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(30, 41, 59, 0.8), 
                        stop:1 rgba(20, 30, 48, 0.7)); 
                    border: 1px solid rgba(71, 85, 105, 0.6); 
                    border-radius: 18px; 
                    padding: 14px 28px; 
                    color: #E2E8F0;
                    font-weight: 700;
                    font-size: 14px;
                    min-height: 40px;
                }
                QPushButton:hover { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(51, 65, 85, 0.9), 
                        stop:1 rgba(30, 41, 59, 0.8));
                    border: 1px solid rgba(96, 165, 250, 0.7);
                    color: #FFFFFF;
                }
                QPushButton:pressed {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(15, 23, 42, 0.95), 
                        stop:1 rgba(20, 30, 48, 0.9));
                    border: 1px solid rgba(59, 130, 246, 0.8);
                }
                QPushButton#accent { 
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #2563EB, 
                        stop:0.5 #3B82F6, 
                        stop:1 #1D4ED8);
                    border: 1px solid rgba(147, 197, 253, 0.6);
                    color: #FFFFFF;
                    font-weight: 800;
                    padding: 15px 32px;
                    font-size: 15px;
                    min-height: 44px;
                }
                QPushButton#accent:hover {
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #1D4ED8, 
                        stop:0.5 #2563EB, 
                        stop:1 #1E40AF);
                    border: 1px solid rgba(147, 197, 253, 0.9);
                }
            """)
        elif ticket.quality_score >= 80:
            frame.setStyleSheet("""
                QFrame {
                    background:#0E141E; 
                    border:2px solid #3B82F6; 
                    border-radius:12px; 
                    padding:16px;
                }
                QPushButton { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(30, 41, 59, 0.8), 
                        stop:1 rgba(20, 30, 48, 0.7)); 
                    border: 1px solid rgba(71, 85, 105, 0.6); 
                    border-radius: 18px; 
                    padding: 14px 28px; 
                    color: #E2E8F0;
                    font-weight: 700;
                    font-size: 14px;
                    min-height: 40px;
                }
                QPushButton:hover { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(51, 65, 85, 0.9), 
                        stop:1 rgba(30, 41, 59, 0.8));
                    border: 1px solid rgba(96, 165, 250, 0.7);
                    color: #FFFFFF;
                }
                QPushButton:pressed {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(15, 23, 42, 0.95), 
                        stop:1 rgba(20, 30, 48, 0.9));
                    border: 1px solid rgba(59, 130, 246, 0.8);
                }
                QPushButton#accent { 
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #2563EB, 
                        stop:0.5 #3B82F6, 
                        stop:1 #1D4ED8);
                    border: 1px solid rgba(147, 197, 253, 0.6);
                    color: #FFFFFF;
                    font-weight: 800;
                    padding: 15px 32px;
                    font-size: 15px;
                    min-height: 44px;
                }
                QPushButton#accent:hover {
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #1D4ED8, 
                        stop:0.5 #2563EB, 
                        stop:1 #1E40AF);
                    border: 1px solid rgba(147, 197, 253, 0.9);
                }
            """)
        else:
            frame.setStyleSheet("""
                QFrame {
                    background:#1E0E0E; 
                    border:2px solid #F59E0B; 
                    border-radius:12px; 
                    padding:16px;
                }
                QPushButton { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(30, 41, 59, 0.8), 
                        stop:1 rgba(20, 30, 48, 0.7)); 
                    border: 1px solid rgba(71, 85, 105, 0.6); 
                    border-radius: 18px; 
                    padding: 14px 28px; 
                    color: #E2E8F0;
                    font-weight: 700;
                    font-size: 14px;
                    min-height: 40px;
                }
                QPushButton:hover { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(51, 65, 85, 0.9), 
                        stop:1 rgba(30, 41, 59, 0.8));
                    border: 1px solid rgba(96, 165, 250, 0.7);
                    color: #FFFFFF;
                }
                QPushButton:pressed {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 rgba(15, 23, 42, 0.95), 
                        stop:1 rgba(20, 30, 48, 0.9));
                    border: 1px solid rgba(59, 130, 246, 0.8);
                }
                QPushButton#accent { 
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #2563EB, 
                        stop:0.5 #3B82F6, 
                        stop:1 #1D4ED8);
                    border: 1px solid rgba(147, 197, 253, 0.6);
                    color: #FFFFFF;
                    font-weight: 800;
                    padding: 15px 32px;
                    font-size: 15px;
                    min-height: 44px;
                }
                QPushButton#accent:hover {
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:0.5, 
                        stop:0 #1D4ED8, 
                        stop:0.5 #2563EB, 
                        stop:1 #1E40AF);
                    border: 1px solid rgba(147, 197, 253, 0.9);
                }
            """)
        
        layout = QVBoxLayout(frame)
        layout.setSpacing(8)
        
        # Status icon and title
        header_layout = QHBoxLayout()
        status_icon = "✅" if ticket.analyzed else "⏳" if ticket.quality_score >= 80 else "⚠️"
        title_label = QLabel(f"{status_icon} {ticket.title}")
        title_label.setStyleSheet("font-size:16px; font-weight:600; color:#EAECEF; border:none; background:transparent;")
        header_layout.addWidget(title_label)
        header_layout.addStretch(1)
        layout.addLayout(header_layout)
        
        # Quality score
        quality_text = f"Quality: {ticket.quality_score}% "
        if ticket.quality_score >= 90:
            quality_text += "(Excellent)"
            quality_color = "#22C55E"
        elif ticket.quality_score >= 80:
            quality_text += "(Good)"
            quality_color = "#3B82F6"
        elif ticket.quality_score >= 70:
            quality_text += "(Fair)"
            quality_color = "#F59E0B"
        else:
            quality_text += "(Needs Work)"
            quality_color = "#EF4444"
        
        quality_label = QLabel(quality_text)
        quality_label.setStyleSheet(f"color:{quality_color}; font-weight:600; border:none; background:transparent;")
        layout.addWidget(quality_label)
        
        # Summary preview
        summary_preview = ticket.summary[:100] + ("..." if len(ticket.summary) > 100 else "")
        summary_label = QLabel(f"Summary: {summary_preview}")
        summary_label.setProperty("muted", True)
        summary_label.setStyleSheet("border:none; background:transparent;")
        summary_label.setWordWrap(True)
        layout.addWidget(summary_label)
        
        # AC count
        ac_label = QLabel(f"AC Count: {ticket.ac_count}")
        ac_label.setProperty("muted", True)
        ac_label.setStyleSheet("border:none; background:transparent;")
        layout.addWidget(ac_label)
        
        # Source tickets
        if ticket.child_tickets:
            source_parts = []
            for child in ticket.child_tickets[:5]:  # Show first 5
                key = child.get('key', '')
                summary = child.get('summary', '')
                if key:
                    source_parts.append(f"{key}: {summary[:40]}..." if len(summary) > 40 else f"{key}: {summary}")
            
            remaining = len(ticket.child_tickets) - 5
            if remaining > 0:
                source_parts.append(f"+{remaining} more")
            
            source_text = "; ".join(source_parts)
            source_label = QLabel(f"📋 Source: {source_text}")
            source_label.setProperty("muted", True)
            source_label.setStyleSheet("border:none; background:transparent; font-size:11px; color:#94A3B8;")
            source_label.setWordWrap(True)
            layout.addWidget(source_label)
        
        # Analysis status
        if ticket.analyzed:
            status_label = QLabel(f"✅ ANALYZED - {len(ticket.test_cases or [])} test cases generated")
            status_label.setStyleSheet("color:#22C55E; font-weight:600; border:none; background:transparent;")
            layout.addWidget(status_label)
        else:
            status_label = QLabel("NOT YET ANALYZED")
            status_label.setStyleSheet("color:#F59E0B; border:none; background:transparent;")
            layout.addWidget(status_label)
        
        # Progress bar for regeneration (hidden by default)
        progress_bar = AnimatedProgressBar()
        progress_bar.setVisible(False)
        progress_bar.setObjectName(f"progress_bar_{ticket.id}")
        layout.addWidget(progress_bar)
        
        # Action buttons
        button_layout = QHBoxLayout()
        button_layout.setSpacing(8)
        button_container = QWidget()
        button_container.setObjectName(f"button_container_{ticket.id}")
        button_container.setLayout(button_layout)
        
        view_btn = QPushButton("View Full Ticket")
        copy_btn = QPushButton("Copy to Clipboard")
        
        if ticket.analyzed:
            analyze_btn = QPushButton("Re-analyze")
            view_cases_btn = QPushButton("View Test Cases")
            button_layout.addWidget(view_btn)
            button_layout.addWidget(copy_btn)
            button_layout.addWidget(analyze_btn)
            button_layout.addWidget(view_cases_btn)
            analyze_btn.clicked.connect(lambda checked=False, t=ticket: analyze_callback(t, reanalyze=True))
        else:
            analyze_btn = QPushButton("Analyze This Ticket →")
            analyze_btn.setObjectName("accent")
            regen_btn = QPushButton("Regenerate") if ticket.quality_score < 80 else None
            
            button_layout.addWidget(view_btn)
            button_layout.addWidget(copy_btn)
            if regen_btn:
                button_layout.addWidget(regen_btn)
                regen_btn.clicked.connect(lambda checked=False, t=ticket: regenerate_callback(t))
            button_layout.addWidget(analyze_btn)
            analyze_btn.clicked.connect(lambda checked=False, t=ticket: analyze_callback(t, reanalyze=False))
        
        view_btn.clicked.connect(lambda checked=False, t=ticket: self.show_ticket_details(t))
        copy_btn.clicked.connect(lambda checked=False, t=ticket: copy_callback(t))
        
        layout.addWidget(button_container)
        
        self.tickets_layout.addWidget(frame)
    
    def show_ticket_details(self, ticket: GeneratedTestTicket):
        """Show full ticket details in a dialog."""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Ticket Details - {ticket.title}")
        dialog.resize(800, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Scrollable text browser
        viewer = QTextBrowser()
        viewer.setOpenExternalLinks(False)
        viewer.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:12px; padding:12px;")
        
        # Build HTML
        html_parts = [JIRA_HTML_CSS, "<div class='jira'>"]
        html_parts.append(f"<h2>{_html_escape(ticket.summary)}</h2>")
        html_parts.append(f"<p>{_html_escape(ticket.description).replace(chr(10), '<br/>')}</p>")
        
        if ticket.acceptance_criteria:
            html_parts.append("<h3>Acceptance Criteria</h3><ul>")
            for ac in ticket.acceptance_criteria:
                html_parts.append(f"<li>{_html_escape(ac)}</li>")
            html_parts.append("</ul>")
        
        html_parts.append("</div>")
        viewer.setHtml("".join(html_parts))
        
        layout.addWidget(viewer)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()

class ReadinessPage(QWidget):
    """New: Ticket Readiness Assessment page."""
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        hdr=QLabel("Ticket Readiness Assessment"); hdr.setProperty("h2", True); v.addWidget(hdr)
        
        # Ticket info banner
        self.ticket_info=QLabel("No ticket loaded"); self.ticket_info.setProperty("muted", True)
        self.ticket_info.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        v.addWidget(self.ticket_info)
        
        # Progress bar
        self.progress_bar = AnimatedProgressBar()

        self.progress_bar.setVisible(False)
        v.addWidget(self.progress_bar)
        
        # Readiness score card
        score_frame = QFrame(); score_frame.setProperty("card", True)
        score_layout = QHBoxLayout(score_frame); score_layout.setContentsMargins(20, 20, 20, 20)
        
        score_left = QVBoxLayout()
        self.score_label = QLabel("Readiness Score")
        self.score_label.setStyleSheet("font-size:16px; font-weight:600; color:#A7B3C6; background:transparent;")
        self.score_value = QLabel("---")
        self.score_value.setStyleSheet("font-size:48px; font-weight:800; color:#EAECEF; background:transparent;")
        score_left.addWidget(self.score_label)
        score_left.addWidget(self.score_value)
        
        self.score_badge = QLabel("Not Assessed")
        self.score_badge.setProperty("class", "pill")
        self.score_badge.setStyleSheet("""
            font-size:20px; 
            padding:10px 20px; 
            font-weight:800;
            letter-spacing:0.5px;
        """)
        
        score_layout.addLayout(score_left)
        score_layout.addStretch(1)
        score_layout.addWidget(self.score_badge, 0, Qt.AlignCenter)
        v.addWidget(score_frame)
        
        # Assessment viewer
        self.assessment_viewer=QTextBrowser(); self.assessment_viewer.setOpenExternalLinks(False)
        self.assessment_viewer.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:16px; padding:12px;")
        self.assessment_viewer.setHtml(plain_to_html("Load and assess a ticket to see readiness analysis..."))
        # Ensure it can handle long content
        self.assessment_viewer.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.assessment_viewer.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        v.addWidget(self.assessment_viewer)
        
        # Navigation buttons
        row=QHBoxLayout()
        self.back_btn=QPushButton("← Back to Ticket")
        self.refresh_btn=QPushButton("Re-assess Ticket")
        self.to_cases_btn=QPushButton("Generate Test Cases →"); self.to_cases_btn.setObjectName("accent")
        self.to_cases_btn.setEnabled(False)
        row.addWidget(self.back_btn); row.addWidget(self.refresh_btn); row.addStretch(1); row.addWidget(self.to_cases_btn)
        v.addLayout(row)

class OverviewPage(QWidget):
    """New: Dedicated page for AI Feature Overview."""
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        hdr=QLabel("AI Feature Overview"); hdr.setProperty("h2", True); v.addWidget(hdr)
        
        # Feature info banner
        self.feature_info=QLabel("No feature loaded"); self.feature_info.setProperty("muted", True)
        self.feature_info.setStyleSheet("padding:10px; background:#0E141E; border:1px solid #192334; border-radius:12px;")
        v.addWidget(self.feature_info)
        
        # Progress bar
        self.progress_bar = AnimatedProgressBar()

        self.progress_bar.setVisible(False)
        v.addWidget(self.progress_bar)
        
        # Overview viewer
        self.overview_viewer=QTextBrowser(); self.overview_viewer.setOpenExternalLinks(False)
        self.overview_viewer.setStyleSheet("background:#0E141E; border:1px solid #192334; border-radius:16px; padding:12px;")
        self.overview_viewer.setHtml(plain_to_html("Load a feature to see AI-generated overview..."))
        v.addWidget(self.overview_viewer)
        
        # Navigation buttons
        row=QHBoxLayout()
        self.back_btn=QPushButton("← Back to Feature")
        self.refresh_btn=QPushButton("Regenerate Overview")
        self.to_ticket_btn=QPushButton("Next → Ticket to analyze"); self.to_ticket_btn.setObjectName("accent")
        row.addWidget(self.back_btn); row.addWidget(self.refresh_btn); row.addStretch(1); row.addWidget(self.to_ticket_btn)
        v.addLayout(row)



class CasesPage(QWidget):
    def __init__(self):
        super().__init__()
        v=QVBoxLayout(self); v.setSpacing(10)
        hdr=QLabel("Test Cases"); hdr.setProperty("h2", True); v.addWidget(hdr)
        
        # Progress bar
        self.progress_bar = AnimatedProgressBar()
        self.progress_bar.setVisible(False)
        v.addWidget(self.progress_bar)
        
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollArea > QWidget > QWidget {
                background: #0E141E;
                border: 1px solid #192334;
                border-radius: 16px;
            }
        """)

        self.container = QWidget()
        self.flow = FlowLayout(self.container, hspacing=12, vspacing=12)
        self.container.setLayout(self.flow)
        self.scroll.setWidget(self.container)
        v.addWidget(self.scroll, 1)

        row=QHBoxLayout()
        self.back_btn=QPushButton("← Back to Ticket")
        self.traceability_btn=QPushButton("📊 View Traceability Matrix")
        self.select_all_btn=QPushButton("Select All")
        self.select_none_btn=QPushButton("Deselect All")
        self.view_ai_btn=QPushButton("View Raw AI Output…")
        self.edit_btn=QPushButton("Edit Selected")
        self.remove_btn=QPushButton("Remove Selected")
        self.to_export_btn=QPushButton("Go to Export →"); self.to_export_btn.setObjectName("accent")
        row.addWidget(self.back_btn); row.addWidget(self.traceability_btn); row.addWidget(self.select_all_btn); row.addWidget(self.select_none_btn); row.addStretch(1); row.addWidget(self.view_ai_btn); row.addWidget(self.edit_btn); row.addWidget(self.remove_btn); row.addWidget(self.to_export_btn)
        v.addLayout(row)
        self.tiles: List[EnhancedTestCaseTile]=[]
    
    def viewport_width(self)->int:
        m=self.flow.contentsMargins(); return max(600, self.scroll.viewport().width()-m.left()-m.right()-12)
    
    def on_tile_toggle(self, tile: EnhancedTestCaseTile, expanded: bool):
        if expanded:
            w=self.viewport_width(); tile.setMaximumWidth(w); tile.setMinimumWidth(w)
        else:
            tile.setMaximumWidth(COMPACT_WIDTH); tile.setMinimumWidth(COMPACT_WIDTH)
        self.container.updateGeometry(); self.flow.invalidate(); self.container.adjustSize()

    def set_cases(self, cases: List[TestCase], on_edit, on_remove):
        # Clear existing tiles properly
        while self.flow.count() > 0:
            item = self.flow.takeAt(0)
            if item is not None:
                widget = item.widget()
                if widget is not None:
                    widget.deleteLater()
        
        # Clear tiles list
        self.tiles = []
        
        # Create new tiles
        for c in cases:
            t = EnhancedTestCaseTile(
                case=c,
                on_edit=on_edit,
                on_remove=on_remove,
                on_toggle_expand=self.on_tile_toggle
            )
            self.flow.addWidget(t)
            self.tiles.append(t)
        
        # Force layout update
        self.container.updateGeometry()
        self.flow.update()
        self.container.adjustSize()
    
    def get_selected_cases(self) -> List[TestCase]:
        """Get list of selected test cases."""
        return [tile.case for tile in self.tiles if tile.checkbox.isChecked()]
    
    def select_all(self):
        """Select all test case checkboxes."""
        for tile in self.tiles:
            tile.checkbox.setChecked(True)
    
    def select_none(self):
        """Deselect all test case checkboxes."""
        for tile in self.tiles:
            tile.checkbox.setChecked(False)

class ExportPage(QWidget):
    """Export test cases with tile preview."""
    def __init__(self):
        super().__init__()
        v = QVBoxLayout(self)
        v.setSpacing(10)
        
        hdr = QLabel("Export Test Cases")
        hdr.setProperty("h2", True)
        v.addWidget(hdr)
        
        # Info banner
        info = QLabel("Select test cases to export and choose your format")
        info.setProperty("muted", True)
        v.addWidget(info)
        
        # Controls row
        controls = QHBoxLayout()
        
        # Select all/none
        self.select_all_btn = QPushButton("Select All")
        self.select_none_btn = QPushButton("Select None")
        controls.addWidget(self.select_all_btn)
        controls.addWidget(self.select_none_btn)
        
        controls.addStretch(1)
        
        # Format selector
        controls.addWidget(QLabel("Export Format:"))
        self.format_combo = QComboBox()
        self.format_combo.addItems(["TestRail CSV", "Xray CSV", "Generic CSV", "Excel (XLSX)"])
        controls.addWidget(self.format_combo)
        
        self.export_btn = QPushButton("Export Selected")
        self.export_btn.setObjectName("accent")
        controls.addWidget(self.export_btn)
        
        v.addLayout(controls)
        
        # Scroll area with tiles
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
        """)
        
        self.container = QWidget()
        self.flow = FlowLayout(self.container, hspacing=12, vspacing=12)
        self.container.setLayout(self.flow)
        self.scroll.setWidget(self.container)
        v.addWidget(self.scroll, 1)
        
        # Summary row
        summary = QHBoxLayout()
        self.summary_label = QLabel("No test cases loaded")
        self.summary_label.setProperty("muted", True)
        summary.addWidget(self.summary_label)
        summary.addStretch(1)
        v.addLayout(summary)
        
        self.tiles: List[ExportTestCaseTile] = []
    
    def set_cases(self, cases: List[TestCase]):
        """Populate export tiles."""
        # Clear existing
        while self.flow.count() > 0:
            item = self.flow.takeAt(0)
            if item and item.widget():
                item.widget().deleteLater()
        
        self.tiles = []
        
        # Create tiles
        for idx, case in enumerate(cases, start=1):
            tile = ExportTestCaseTile(case, idx)
            self.flow.addWidget(tile)
            self.tiles.append(tile)
        
        # Update summary
        self.summary_label.setText(f"{len(cases)} test case(s) available for export")
        
        # Force layout update
        self.container.updateGeometry()
        self.flow.update()
    
    def get_selected_cases(self) -> List[TestCase]:
        """Get selected test cases."""
        return [tile.case for tile in self.tiles if tile.checkbox.isChecked()]
    
    def select_all(self):
        """Select all checkboxes."""
        for tile in self.tiles:
            tile.checkbox.setChecked(True)
    
    def select_none(self):
        """Deselect all checkboxes."""
        for tile in self.tiles:
            tile.checkbox.setChecked(False)

# ---------- Worker ----------
class GenerateWorker(QThread):
    finished_ok = Signal(list, list, bool, str, object, object)
    failed = Signal(str)
    progress = Signal(int, str)
    
    def __init__(self, analyzer: Analyzer, summary: str, body: str, ac: List[str], feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None):
        super().__init__()
        self.analyzer = analyzer
        self.summary = summary
        self.body = body
        self.ac = ac
        self.feature_ctx = feature_ctx
        self.attachments = attachments
        self._is_running = True
    
    def run(self):
        try:
            cases, requirements, used_fallback, raw_ai, ai_err, critic_review = self.analyzer.generate_test_cases(
                self.summary, 
                self.body, 
                self.ac, 
                self.feature_ctx,
                self.attachments,
                max_iterations=3,
                progress_callback=self._report_progress
            )
            if self._is_running:  # Only emit if still valid
                self.finished_ok.emit(cases, requirements, used_fallback, raw_ai, ai_err, critic_review)
        except Exception as e:
            if self._is_running:  # Only emit if still valid
                self.failed.emit(str(e))
    
    def _report_progress(self, percentage: int, message: str):
        """Emit progress updates from the analyzer."""
        if self._is_running:  # Only emit if still valid
            self.progress.emit(percentage, message)
    
    def stop(self):
        """Mark the worker as stopped."""
        self._is_running = False

class OverviewWorker(QThread):
    finished_ok = Signal(str, object)
    failed = Signal(str)
    
    def __init__(self, analyzer: Analyzer, feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None):
        super().__init__()
        self.analyzer = analyzer
        self.feature_ctx = feature_ctx
        self.attachments = attachments
        self._is_running = True
    
    def run(self):
        try:
            raw_ai, ai_err = self.analyzer.generate_feature_overview(self.feature_ctx, self.attachments)
            if self._is_running:  # Only emit if still valid
                self.finished_ok.emit(raw_ai, ai_err)
        except Exception as e:
            if self._is_running:  # Only emit if still valid
                self.failed.emit(str(e))
    
    def stop(self):
        """Mark the worker as stopped."""
        self._is_running = False

class ReadinessWorker(QThread):
    finished_ok = Signal(str, object)
    failed = Signal(str)
    
    def __init__(self, analyzer: Analyzer, summary: str, body: str, ac: List[str], feature_ctx: Optional[Dict], attachments: Optional[List[Dict]] = None):
        super().__init__()
        self.analyzer = analyzer
        self.summary = summary
        self.body = body
        self.ac = ac
        self.feature_ctx = feature_ctx
        self.attachments = attachments
        self._is_running = True
    
    def run(self):
        try:
            raw_ai, ai_err = self.analyzer.assess_ticket_readiness(self.summary, self.body, self.ac, self.feature_ctx, self.attachments)
            if self._is_running:  # Only emit if still valid
                self.finished_ok.emit(raw_ai, ai_err)
        except Exception as e:
            if self._is_running:  # Only emit if still valid
                self.failed.emit(str(e))
    
    def stop(self):
        """Mark the worker as stopped."""
        self._is_running = False

class AnalyzeSplitWorker(QThread):
    """Worker for analyzing and recommending test ticket splits."""
    finished_ok = Signal(object, object)  # (split_data, error)
    failed = Signal(str)
    
    def __init__(self, analyzer: Analyzer, feature_ctx: Dict):
        super().__init__()
        self.analyzer = analyzer
        self.feature_ctx = feature_ctx
        self._is_running = True
    
    def run(self):
        try:
            split_data, ai_err = self.analyzer.analyze_and_split_for_test_tickets(self.feature_ctx)
            if self._is_running:
                self.finished_ok.emit(split_data, ai_err)
        except Exception as e:
            if self._is_running:
                self.failed.emit(str(e))
    
    def stop(self):
        self._is_running = False

class GenerateTestTicketWorker(QThread):
    """Worker for generating a single test ticket."""
    finished_ok = Signal(object, object, object)  # (ticket_data, raw_response, error)
    failed = Signal(str)
    progress = Signal(int, str)  # progress percentage and message
    
    def __init__(
        self, 
        analyzer: Analyzer, 
        epic_name: str, 
        functional_area: str, 
        child_tickets: List[Dict],
        feature_context: Dict,
        previous_attempt: Optional[str] = None,
        reviewer_feedback: Optional[Dict] = None
    ):
        super().__init__()
        self.analyzer = analyzer
        self.epic_name = epic_name
        self.functional_area = functional_area
        self.child_tickets = child_tickets
        self.feature_context = feature_context
        self.previous_attempt = previous_attempt
        self.reviewer_feedback = reviewer_feedback
        self._is_running = True
    
    def run(self):
        try:
            if self._is_running:
                self.progress.emit(30, "Creating test ticket draft...")
            
            ticket_data, raw_response = self.analyzer.generate_test_ticket(
                self.epic_name,
                self.functional_area,
                self.child_tickets,
                self.feature_context,
                self.previous_attempt,
                self.reviewer_feedback
            )
            
            if self._is_running:
                self.finished_ok.emit(ticket_data, raw_response, None)
        except Exception as e:
            if self._is_running:
                self.failed.emit(str(e))
    
    def stop(self):
        self._is_running = False

class ReviewTestTicketWorker(QThread):
    """Worker for reviewing a generated test ticket."""
    finished_ok = Signal(object, object)  # (review_data, error)
    failed = Signal(str)
    progress = Signal(int, str)
    
    def __init__(
        self,
        analyzer: Analyzer,
        ticket_content: Dict,
        epic_name: str,
        feature_context: Dict
    ):
        super().__init__()
        self.analyzer = analyzer
        self.ticket_content = ticket_content
        self.epic_name = epic_name
        self.feature_context = feature_context
        self._is_running = True
    
    def run(self):
        try:
            if self._is_running:
                self.progress.emit(60, "Reviewing ticket quality...")
            
            review_data, ai_err = self.analyzer.review_test_ticket(
                self.ticket_content,
                self.epic_name,
                self.feature_context
            )
            
            if self._is_running:
                self.finished_ok.emit(review_data, ai_err)
        except Exception as e:
            if self._is_running:
                self.failed.emit(str(e))
    
    def stop(self):
        self._is_running = False

class RegenerateTicketWorker(QThread):
    """Worker for regenerating a test ticket until it passes quality threshold."""
    finished_ok = Signal(object, object, object)  # (ticket_data, raw_response, review_data)
    failed = Signal(str)
    progress = Signal(int, str)
    
    def __init__(
        self,
        analyzer: Analyzer,
        ticket: GeneratedTestTicket,
        epic_name: str,
        functional_area: str,
        child_tickets: List[Dict],
        feature_context: Dict,
        max_attempts: int = 3
    ):
        super().__init__()
        self.analyzer = analyzer
        self.ticket = ticket
        self.epic_name = epic_name
        self.functional_area = functional_area
        self.child_tickets = child_tickets
        self.feature_context = feature_context
        self.max_attempts = max_attempts
        self._is_running = True
    
    def run(self):
        try:
            previous_attempt = None
            reviewer_feedback = None
            
            for attempt in range(self.max_attempts):
                if not self._is_running:
                    return
                
                # Generate ticket
                attempt_num = attempt + 1
                self.progress.emit(
                    int((attempt / self.max_attempts) * 50),
                    f"Regenerating ticket (attempt {attempt_num}/{self.max_attempts})..."
                )
                
                ticket_data, raw_response = self.analyzer.generate_test_ticket(
                    self.epic_name,
                    self.functional_area,
                    self.child_tickets,
                    self.feature_context,
                    previous_attempt,
                    reviewer_feedback
                )
                
                if not ticket_data:
                    if attempt == self.max_attempts - 1:
                        self.failed.emit("Failed to generate ticket after all attempts")
                    continue
                
                if not self._is_running:
                    return
                
                # Review ticket
                self.progress.emit(
                    int((attempt / self.max_attempts) * 50) + 25,
                    f"Reviewing quality (attempt {attempt_num}/{self.max_attempts})..."
                )
                
                review_data, ai_err = self.analyzer.review_test_ticket(
                    ticket_data,
                    self.epic_name,
                    self.feature_context
                )
                
                if not review_data:
                    review_data = {"approved": False, "quality_score": 70, "issues": ["Review failed"], "recommendations": []}
                
                quality_score = review_data.get('quality_score', 70)
                
                # Check if it passes
                if quality_score >= 80:
                    self.progress.emit(100, f"✅ Regeneration successful! Quality: {quality_score}%")
                    if self._is_running:
                        self.finished_ok.emit(ticket_data, raw_response, review_data)
                    return
                
                # Prepare for next attempt
                previous_attempt = raw_response or json.dumps(ticket_data)
                reviewer_feedback = review_data
                
                if attempt < self.max_attempts - 1:
                    self.progress.emit(
                        int(((attempt + 1) / self.max_attempts) * 50),
                        f"Quality {quality_score}% - retrying..."
                    )
            
            # All attempts failed
            if self._is_running:
                self.failed.emit(f"Ticket quality still below 80% after {self.max_attempts} attempts. Last score: {quality_score}%")
        
        except Exception as e:
            if self._is_running:
                self.failed.emit(str(e))
    
    def stop(self):
        self._is_running = False

# ---------- Main Window ----------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__(); self.setWindowTitle("AI Software Tester – Jira → Azure DevOps"); self.resize(1280, 820)

        self.login=LoginPage(); self.feature=FeaturePage(); self.overview=OverviewPage(); self.ticket=TicketPage(); self.create_tickets=CreateTestTicketsPage(); self.tickets_results=TestTicketsResultsPage(); self.readiness=ReadinessPage(); self.cases=CasesPage(); self.exportp=ExportPage()

        root=QWidget(); self.setCentralWidget(root)
        v=QVBoxLayout(root); v.setContentsMargins(12,12,12,12); v.setSpacing(8)
        self.topbar=TopBar("Sign In", on_back=self.back_logic); v.addWidget(self.topbar,0)
        self.tabs=TopTabs(); v.addWidget(self.tabs,0); self.tabs.hide()
        self.stack=QStackedWidget(); [self.stack.addWidget(p) for p in (self.login,self.feature,self.overview,self.ticket,self.create_tickets,self.tickets_results,self.readiness,self.cases,self.exportp)]; v.addWidget(self.stack,1)

        apply_enhanced_shadow(self.topbar, blur=40, offset_y=16, opacity=0.25)
        apply_enhanced_shadow(self.stack, blur=45, offset_y=20, opacity=0.22)

        # State
        self.jira: Optional[JiraClient]=None
        self.llm=LLM(enabled=False); self.analyzer=Analyzer(self.llm)
        self.feature_context: Optional[Dict] = None  # <-- holds Epic + children for session
        self.feature_attachments: List[Dict] = []  # Attachments from Feature/Epic/Initiative
        self.current_key=""; self.current_summary=""; self.current_body=""; self.current_ac: List[str]=[]; self.current_attachments: List[Dict] = []
        self.test_cases: List[TestCase]=[]; self.requirements: List[Dict]=[]; self.used_fallback=False; self.raw_ai_text=""; self.ai_error: Optional[str]=None
        
        # New state for test ticket creation
        self.generated_test_tickets: List[GeneratedTestTicket] = []  # Stores created test tickets
        self.split_plan: Optional[Dict] = None  # Stores recommended splits
        self.current_split_index: int = 0  # For sequential generation
        self.current_relevant_children: List[Dict] = []  # Child tickets for current generation
        self.current_generated_ticket: Optional[GeneratedTestTicket] = None  # Currently analyzed generated ticket

        # Nav wiring
        self.tabs.btn_feature.clicked.connect(lambda: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze"))
        self.tabs.btn_overview.clicked.connect(lambda: self._nav(self.overview, self.tabs.btn_overview, "AI Feature Overview"))
        self.tabs.btn_ticket.clicked.connect(lambda: self._nav(self.ticket,  self.tabs.btn_ticket,  "Ticket to analyze"))
        self.tabs.btn_readiness.clicked.connect(lambda: self._nav(self.readiness, self.tabs.btn_readiness, "Ticket Readiness"))
        self.tabs.btn_cases.clicked.connect(lambda: self._nav(self.cases,   self.tabs.btn_cases,   "Test Cases"))
        self.tabs.btn_export.clicked.connect(lambda: self._nav(self.exportp, self.tabs.btn_export, "Export"))
        
        # Wire up export page
        self.exportp.select_all_btn.clicked.connect(self.exportp.select_all)
        self.exportp.select_none_btn.clicked.connect(self.exportp.select_none)
        self.exportp.export_btn.clicked.connect(self.on_export)

        # Actions
        self.login.next_btn.clicked.connect(self.on_login)

        self.feature.fetch_btn.clicked.connect(self.on_feature_fetch)
        self.feature.load_epics_btn.clicked.connect(self.on_load_manual_epics)
        self.feature.to_overview_btn.clicked.connect(self.on_navigate_to_overview)
        
        self.overview.back_btn.clicked.connect(lambda: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze"))
        self.overview.refresh_btn.clicked.connect(self.refresh_feature_overview)
        self.overview.to_ticket_btn.clicked.connect(lambda: self._nav(self.ticket, self.tabs.btn_ticket, "Ticket to analyze"))

        # Ticket page - wire both dropdown actions and main button
        self.ticket.fetch_btn.clicked.connect(self.on_fetch_ticket)
        self.ticket.action_analyze.triggered.connect(self.on_fetch_ticket)
        self.ticket.action_create.triggered.connect(self.on_navigate_to_create_test_tickets)
        self.ticket.action_view_created.triggered.connect(self.on_view_created_tickets)
        self.ticket.assess_btn.clicked.connect(self.on_navigate_to_readiness)
        self.ticket.clear_btn.clicked.connect(self.clear_ticket)
        
        # Create test tickets page
        self.create_tickets.back_btn.clicked.connect(lambda: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze"))
        self.create_tickets.analyze_btn.clicked.connect(self.on_analyze_splits)
        self.create_tickets.generate_btn.clicked.connect(self.on_generate_test_tickets)
        
        # Test tickets results page
        self.tickets_results.back_btn.clicked.connect(lambda: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze"))
        self.tickets_results.export_all_btn.clicked.connect(self.on_export_all_test_tickets)
        
        self.readiness.back_btn.clicked.connect(lambda: self._nav(self.ticket, self.tabs.btn_ticket, "Ticket to analyze"))
        self.readiness.refresh_btn.clicked.connect(self.assess_ticket_readiness)
        self.readiness.to_cases_btn.clicked.connect(self.on_generate_clicked)

        self.cases.view_ai_btn.clicked.connect(self.on_view_raw_ai)
        self.cases.traceability_btn.clicked.connect(self.on_view_traceability)
        self.cases.edit_btn.clicked.connect(self.edit_selected_tile)
        self.cases.remove_btn.clicked.connect(self.remove_selected_tile)
        self.cases.to_export_btn.clicked.connect(self.goto_export_from_cases)
        self.cases.back_btn.clicked.connect(lambda: self._nav(self.readiness, self.tabs.btn_readiness, "Ticket Readiness"))
        self.cases.select_all_btn.clicked.connect(self.cases.select_all)
        self.cases.select_none_btn.clicked.connect(self.cases.select_none)

        # Export page navigation handled by tab buttons
        
        self._nav(self.login, None, "Sign In"); self.topbar.back.setVisible(False)
        
        # Auto-login if credentials are in .env
        if self.login.has_all_credentials():
            QTimer.singleShot(100, self.on_login)

    # Navigation
    def _nav(self, page: QWidget, btn: Optional[QPushButton], title: str):
        self.stack.setCurrentWidget(page)
        if btn: self.tabs.set_active(btn)
        self.topbar.title.setText(title)

    def back_logic(self):
        cur=self.stack.currentWidget()
        if cur is self.feature: self._nav(self.login, None, "Sign In")
        elif cur is self.overview: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze")
        elif cur is self.ticket: self._nav(self.overview, self.tabs.btn_overview, "AI Feature Overview")
        elif cur is self.create_tickets: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze")
        elif cur is self.tickets_results: self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze")
        elif cur is self.readiness: self._nav(self.ticket, self.tabs.btn_ticket, "Ticket to analyze")
        elif cur is self.cases:  self._nav(self.readiness, self.tabs.btn_readiness, "Ticket Readiness")
        elif cur is self.exportp:self._nav(self.cases,  self.tabs.btn_cases,  "Test Cases")

    # --- Sign in
    def on_login(self):
        email=self.login.email.text().strip(); token=self.login.token.text().strip()
        if not (email and token): QMessageBox.warning(self,"Missing info","Please provide Email and API Token."); return
        try:
            key=self.login.openai.text().strip(); model=self.login.model.currentText()
            if key: os.environ["OPENAI_API_KEY"]=key; os.environ["OPENAI_MODEL"]=model; self.llm=LLM(enabled=True, model=model)
            else: self.llm=LLM(enabled=False, model=model)
            if key and not self.llm.import_ok: QMessageBox.information(self,"OpenAI package missing","Run:  pip install openai")
            self.topbar.ai_status.setText(self.llm.status_label()); self.topbar.ai_status.setToolTip(self.llm.status_label())
            self.analyzer=Analyzer(self.llm)
            self.jira=JiraClient(JIRA_BASE_URL, email, token)
            self.tabs.show(); self.tabs.set_enabled_after_login(True); self.topbar.back.setVisible(True)
            # Start on Feature step
            self._nav(self.feature, self.tabs.btn_feature, "Feature to analyze")
        except Exception as e:
            QMessageBox.critical(self,"Login error",str(e))

   # --- Feature (Epic)
    def on_feature_fetch(self):
        if not self.jira: QMessageBox.warning(self,"Not signed in","Please sign in first."); return
        feature_key=self.feature.key_edit.text().strip().upper()
        if not re.match(r"^[A-Za-z][A-Za-z0-9_]+-\d+$", feature_key):
            QMessageBox.warning(self,"Invalid key","Format: <PROJECT>-<number>, e.g., UEX-100"); return
        
        try:
            # First, check what type of issue this is
            issue = self.jira.get_issue(feature_key)
            issue_type = issue.get("fields", {}).get("issuetype", {}).get("name", "").lower()
            
            if "initiative" in issue_type:
                # Handle as Initiative
                init_data = self.jira.get_initiative_details(feature_key)
                
                # Build flattened context for AI
                init = init_data["initiative"]
                all_children = []
                epic_summaries = []
                
                for epic in init_data.get("epics", []):
                    epic_summaries.append(f"{epic['key']}: {epic['summary']}")
                    all_children.extend(epic.get("children", []))
                
                # Fetch attachments from Initiative and Epics
                print(f"DEBUG: Fetching attachments for Initiative {feature_key}...")
                self.feature_attachments = []
                try:
                    # Get Initiative attachments
                    init_attachments = self.jira.get_attachments(feature_key)
                    print(f"DEBUG: Found {len(init_attachments)} attachments on Initiative")
                    
                    for attachment in init_attachments:
                        processed = self.jira.process_attachment(attachment)
                        if processed:
                            processed["source"] = f"Initiative: {feature_key}"
                            self.feature_attachments.append(processed)
                            print(f"DEBUG: Processed Initiative attachment: {processed.get('filename')}")
                    
                    # Get attachments from each Epic (limit to images)
                    for epic in init_data.get("epics", [])[:5]:  # Limit to first 5 Epics
                        epic_key = epic.get("key")
                        try:
                            epic_attachments = self.jira.get_attachments(epic_key)
                            for attachment in epic_attachments:
                                # Only process images from Epics to control token usage
                                if attachment.get("mimeType", "").startswith("image/"):
                                    processed = self.jira.process_attachment(attachment)
                                    if processed:
                                        processed["source"] = f"Epic: {epic_key}"
                                        self.feature_attachments.append(processed)
                                        print(f"DEBUG: Processed Epic attachment: {processed.get('filename')} from {epic_key}")
                        except Exception as e:
                            print(f"DEBUG: Could not fetch attachments for Epic {epic_key}: {e}")
                    
                    print(f"DEBUG: Successfully processed {len(self.feature_attachments)} total attachments for Initiative")
                except Exception as e:
                    print(f"DEBUG: Error processing Initiative attachments: {e}")
                    self.feature_attachments = []
                
                self.feature_context = {
                    "type": "initiative",
                    "initiative_key": init["key"],
                    "initiative_summary": init["summary"],
                    "initiative_desc": init["desc"],
                    "epics": epic_summaries,
                    "epic_count": len(init_data.get("epics", [])),
                    "children": all_children[:100],  # Cap at 100 for token limits
                    "total_children": len(all_children)
                }
                
                # Render HTML
                html = build_initiative_html(init_data)
                self.feature.viewer.setHtml(html)
                self.feature.to_overview_btn.setEnabled(True)
                
                # Check if no Epics were found and show manual entry option
                if len(init_data.get("epics", [])) == 0:
                    # Show manual entry field and store initiative
                    self.current_initiative = {
                        "key": init["key"],
                        "summary": init["summary"],
                        "desc": init["desc"]
                    }
                    self.feature.show_manual_epic_entry()
                    QMessageBox.information(
                        self,
                        "No Epics found automatically",
                        "Could not automatically find child Epics.\n\nPlease enter Epic keys manually in the field above (comma-separated, e.g., UEX-31, UEX-38) and click 'Load Epics'."
                    )
                else:
                    QMessageBox.information(
                        self,
                        "Initiative loaded",
                        f"Loaded Initiative {init['key']} with {len(init_data['epics'])} Epic(s) and {len(all_children)} total child issue(s)."
                    )
                
            else:
                # Handle as Epic (existing logic)
                epic_key = feature_key
                ef=issue.get("fields",{})
                epic_title=ef.get("summary", epic_key); epic_desc=ef.get("description")
                if isinstance(epic_desc, dict) and epic_desc.get("type")=="doc":
                    epic_html=adf_to_html(epic_desc); epic_plain=adf_to_plaintext(epic_desc)
                elif isinstance(epic_desc, str):
                    epic_html=plain_to_html(epic_desc); epic_plain=epic_desc
                else:
                    epic_html=plain_to_html("(no description)"); epic_plain=""
                
                # Clean epic description for LLM
                epic_plain = clean_jira_text_for_llm(epic_plain)
                
                # Get children
                issues=self.jira.get_children_of_epic(epic_key)
                children=[]
                for it in issues:
                    k=it.get("key","")
                    f=it.get("fields",{})
                    s=f.get("summary","")
                    d=f.get("description")
                    if isinstance(d, dict) and d.get("type")=="doc":
                        d_plain=adf_to_plaintext(d)
                    elif isinstance(d,str):
                        d_plain=d
                    else:
                        d_plain=""
                    # Clean child description for LLM
                    d_plain = clean_jira_text_for_llm(d_plain)
                    children.append({"key":k,"summary":s,"desc":d_plain})
                
                # Fetch attachments from Epic
                print(f"DEBUG: Fetching attachments for Epic {epic_key}...")
                self.feature_attachments = []
                try:
                    epic_attachments = self.jira.get_attachments(epic_key)
                    print(f"DEBUG: Found {len(epic_attachments)} attachments on Epic")
                    
                    for attachment in epic_attachments:
                        processed = self.jira.process_attachment(attachment)
                        if processed:
                            self.feature_attachments.append(processed)
                            print(f"DEBUG: Processed Epic attachment: {processed.get('filename')} ({processed.get('type')})")
                    
                    # Also fetch image attachments from children (to limit token usage)
                    for child in children[:10]:  # Limit to first 10 children
                        child_key = child.get("key")
                        try:
                            child_attachments = self.jira.get_attachments(child_key)
                            for attachment in child_attachments:
                                # Only process images from children
                                if attachment.get("mimeType", "").startswith("image/"):
                                    processed = self.jira.process_attachment(attachment)
                                    if processed:
                                        processed["source"] = f"Child: {child_key}"
                                        self.feature_attachments.append(processed)
                                        print(f"DEBUG: Processed child attachment: {processed.get('filename')} from {child_key}")
                        except Exception as e:
                            print(f"DEBUG: Could not fetch attachments for child {child_key}: {e}")
                    
                    print(f"DEBUG: Successfully processed {len(self.feature_attachments)} total attachments for feature")
                except Exception as e:
                    print(f"DEBUG: Error processing feature attachments: {e}")
                    self.feature_attachments = []
                
                # Store context
                self.feature_context={
                    "type": "epic",
                    "epic_key": epic_key,
                    "epic_summary": epic_title,
                    "epic_desc": epic_plain,
                    "children": children
                }
                
                # Render HTML
                html=build_feature_html(epic_key, epic_title, epic_html, children)
                self.feature.viewer.setHtml(html)
                self.feature.to_overview_btn.setEnabled(True)
                
                QMessageBox.information(self,"Epic loaded", f"Loaded Epic {epic_key} with {len(children)} child issue(s).")
                
        except Exception as e:
            QMessageBox.critical(self,"Fetch error",str(e))

    def on_load_manual_epics(self):
        """Load epics manually from comma-separated keys."""
        if not self.jira:
            QMessageBox.warning(self, "Not signed in", "Please sign in first.")
            return
        
        if not hasattr(self, 'current_initiative') or not self.current_initiative:
            QMessageBox.warning(self, "No Initiative loaded", "Please load an Initiative first.")
            return
        
        epic_keys_str = self.feature.manual_epic_keys.text().strip()
        if not epic_keys_str:
            QMessageBox.warning(self, "No Epic keys", "Please enter at least one Epic key (comma-separated).")
            return
        
        # Parse epic keys
        epic_keys = [k.strip() for k in epic_keys_str.split(',') if k.strip()]
        
        try:
            # Load epics
            self.feature.epic_status.setText(f"Loading {len(epic_keys)} Epic(s)...")
            self.feature.epic_status.setVisible(True)
            QApplication.processEvents()
            
            epics_with_children = self.jira.load_manual_epics(epic_keys)
            
            if not epics_with_children:
                QMessageBox.warning(self, "No Epics loaded", "Could not load any of the specified Epics.")
                self.feature.epic_status.setVisible(False)
                return
            
            # Update feature context
            init = self.current_initiative
            all_children = []
            epic_summaries = []
            total_children = 0
            
            for epic in epics_with_children:
                epic_summaries.append(f"{epic['key']}: {epic['summary']}")
                epic_children = epic.get("children", [])
                all_children.extend(epic_children)
                total_children += len(epic_children)
            
            self.feature_context = {
                "type": "initiative",
                "initiative_key": init["key"],
                "initiative_summary": init["summary"],
                "initiative_desc": init["desc"],
                "epics": epic_summaries,
                "epic_count": len(epics_with_children),
                "children": all_children[:100],  # Cap at 100 for token limits
                "total_children": total_children
            }
            
            # Build initiative data for HTML display
            init_data = {
                "initiative": init,
                "epics": epics_with_children
            }
            
            # Render HTML
            html = build_initiative_html(init_data)
            self.feature.viewer.setHtml(html)
            self.feature.to_overview_btn.setEnabled(True)
            
            # Update status
            status_text = f"✓ Loaded {len(epics_with_children)} Epic(s) with {total_children} total child issue(s)"
            self.feature.epic_status.setText(status_text)
            self.feature.epic_status.setStyleSheet("padding:10px; background:#0E3010; border:1px solid #22C55E; border-radius:12px; color:#86EFAC;")
            
            QMessageBox.information(
                self,
                "Epics loaded",
                f"Successfully loaded:\n" + "\n".join([f"• {epic['key']}: {epic['summary']} ({len(epic.get('children', []))} children)" for epic in epics_with_children])
            )
            
        except Exception as e:
            self.feature.epic_status.setText(f"✗ Error loading Epics: {str(e)}")
            self.feature.epic_status.setStyleSheet("padding:10px; background:#3E0A0A; border:1px solid #EF4444; border-radius:12px; color:#FCA5A5;")
            QMessageBox.critical(self, "Load error", str(e))

    def on_navigate_to_overview(self):
        """Navigate to overview page and generate AI overview."""
        if not self.feature_context:
            QMessageBox.warning(self, "No feature loaded", "Please load a feature first.")
            return
        
        # Navigate to overview page
        self._nav(self.overview, self.tabs.btn_overview, "AI Feature Overview")
        
        # Generate the overview
        self.generate_and_display_overview()

    def generate_and_display_overview(self):
        """Generate and display the AI feature overview."""
        if not self.feature_context:
            self.overview.overview_viewer.setHtml(plain_to_html("No feature loaded."))
            self.overview.feature_info.setText("No feature loaded")
            return
        
        # Update feature info banner
        epic_key = self.feature_context.get("epic_key", "")
        epic_summary = self.feature_context.get("epic_summary", "")
        num_children = len(self.feature_context.get("children", []))
        self.overview.feature_info.setText(f"Analyzing: {epic_key} - {epic_summary} ({num_children} child issues)")
        
        # Show progress bar
        self.overview.progress_bar.setVisible(True)
        self.overview.progress_bar.setValue(0)
        self.overview.progress_bar.setFormat("Initializing AI analysis... 0%")
        self.overview.overview_viewer.setVisible(False)
        
        # Animate progress
        def animate_progress():
            for val, msg in [(20, "Analyzing Epic context..."), (40, "Processing child issues..."), (60, "Generating overview with AI...")]:
                QTimer.singleShot(val * 10, lambda v=val, m=msg: (
                    self.overview.progress_bar.setValue(v),
                    self.overview.progress_bar.setFormat(f"{m} {v}%")
                ))
        animate_progress()
        
        def done(overview_raw, overview_err):
            self.overview.progress_bar.setValue(90)
            self.overview.progress_bar.setFormat("Formatting results... 90%")
            QApplication.processEvents()
            
            if overview_err:
                self.overview.overview_viewer.setHtml(plain_to_html(f"AI Overview unavailable: {overview_err}"))
            else:
                overview_data = safe_json_extract(overview_raw)
                if overview_data:
                    self.overview.overview_viewer.setHtml(build_overview_html(overview_data))
                else:
                    self.overview.overview_viewer.setHtml(plain_to_html("Could not parse AI overview."))
            
            self.overview.progress_bar.setValue(100)
            self.overview.progress_bar.setFormat("Complete! 100%")
            QTimer.singleShot(500, lambda: (
                self.overview.progress_bar.setVisible(False),
                self.overview.overview_viewer.setVisible(True)
            ))
        
        def failed(msg):
            self.overview.overview_viewer.setHtml(plain_to_html(f"Error: {msg}"))
            self.overview.progress_bar.setVisible(False)
            self.overview.overview_viewer.setVisible(True)
        
        worker = OverviewWorker(self.analyzer, self.feature_context, getattr(self, 'feature_attachments', None))
        worker.finished_ok.connect(done)
        worker.failed.connect(failed)
        worker.start()
        self.overview_worker = worker
    
    def refresh_feature_overview(self):
        """Regenerate the AI feature overview."""
        self.generate_and_display_overview()

    # --- Ticket page
    def on_fetch_ticket(self):
        if not self.jira: QMessageBox.warning(self,"Not signed in","Please sign in first."); return
        key=self.ticket.key_edit.text().strip().upper()
        if not re.match(r"^[A-Za-z][A-Za-z0-9_]+-\d+$", key):
            QMessageBox.warning(self,"Invalid key","Format: <PROJECT>-<number>, e.g., PAY-1042"); return
        try:
            issue=self.jira.get_issue(key); fields=issue.get("fields",{})
            summary=fields.get("summary", key); description=fields.get("description")
            if isinstance(description, dict) and description.get("type")=="doc":
                desc_html=adf_to_html(description); body_plain=adf_to_plaintext(description)
            elif isinstance(description, str):
                desc_html=plain_to_html(description); body_plain,_=self.analyzer.extract_sections(description)
            else:
                desc_html=plain_to_html("(no description)"); body_plain=""
            
            # Clean ticket body for LLM
            body_plain = clean_jira_text_for_llm(body_plain)
            
            ac_blocks=[]
            for fname in ["customfield_AcceptanceCriteria","customfield_12345","Acceptance Criteria","AC"]:
                val=fields.get(fname)
                if isinstance(val,str) and val.strip(): ac_blocks.append(val)
            if isinstance(description,str):
                _, ac_from_body=self.analyzer.extract_sections(description or ""); ac_blocks.extend(ac_from_body or [])
            
            # Fetch and process attachments
            print(f"DEBUG: Fetching attachments for {key}...")
            self.current_attachments = []
            try:
                attachments = self.jira.get_attachments(key)
                print(f"DEBUG: Found {len(attachments)} attachments")
                
                # Process each attachment
                for attachment in attachments:
                    processed = self.jira.process_attachment(attachment)
                    if processed:
                        self.current_attachments.append(processed)
                        print(f"DEBUG: Processed attachment: {processed.get('filename')} ({processed.get('type')})")
                
                print(f"DEBUG: Successfully processed {len(self.current_attachments)} attachments")
            except Exception as e:
                print(f"DEBUG: Error processing attachments: {e}")
                self.current_attachments = []
            
            # ... inside on_fetch_ticket method
            self.current_key=key; self.current_summary=summary; self.current_body=body_plain; self.current_ac=ac_blocks
            self.ticket.viewer.setHtml(build_ticket_html(summary, desc_html, ac_blocks))
            self.ticket.assess_btn.setEnabled(True)
            self._nav(self.ticket, self.tabs.btn_ticket, "Ticket to analyze")
        except Exception as e:
            QMessageBox.critical(self,"Fetch error",str(e))

    def on_navigate_to_readiness(self):
        """Navigate to readiness page and assess ticket."""
        # Check if we have ticket content (either Jira key or generated ticket with summary)
        if not self.current_key and not self.current_summary:
            QMessageBox.warning(self, "No ticket loaded", "Please load a ticket first.")
            return
        
        # Navigate to readiness page
        self._nav(self.readiness, self.tabs.btn_readiness, "Ticket Readiness")
        
        # Assess the ticket
        self.assess_ticket_readiness()

    def assess_ticket_readiness(self):
        """Assess the current ticket's readiness for test case generation."""
        # Check if we have ticket content (either Jira key or generated ticket with summary)
        if not self.current_key and not self.current_summary:
            self.readiness.assessment_viewer.setHtml(plain_to_html("No ticket loaded."))
            self.readiness.ticket_info.setText("No ticket loaded")
            return
        
        # Update ticket info banner (use key if available, otherwise use summary)
        ticket_identifier = self.current_key if self.current_key else "Generated Ticket"
        self.readiness.ticket_info.setText(f"Assessing: {ticket_identifier} - {self.current_summary}")
        
        # Show progress bar
        self.readiness.progress_bar.setVisible(True)
        self.readiness.progress_bar.setValue(0)
        self.readiness.progress_bar.setFormat("Initializing assessment... 0%")
        self.readiness.assessment_viewer.setVisible(False)
        self.readiness.score_value.setText("---")
        self.readiness.score_badge.setText("Assessing...")
        
        # Animate progress
        def animate_progress():
            for val, msg in [(25, "Analyzing ticket content..."), (50, "Evaluating acceptance criteria..."), (75, "Generating recommendations...")]:
                QTimer.singleShot(val * 10, lambda v=val, m=msg: (
                    self.readiness.progress_bar.setValue(v),
                    self.readiness.progress_bar.setFormat(f"{m} {v}%")
                ))
        animate_progress()
        
        def done(assessment_raw, assessment_err):
            self.readiness.progress_bar.setValue(90)
            self.readiness.progress_bar.setFormat("Finalizing assessment... 90%")
            QApplication.processEvents()
            
            if assessment_err:
                self.readiness.assessment_viewer.setHtml(plain_to_html(f"Assessment unavailable: {assessment_err}"))
                self.readiness.score_badge.setText("Error")
                self.readiness.score_value.setText("--")
            else:
                assessment_data = safe_json_extract(assessment_raw)
                if assessment_data:
                    score = assessment_data.get("score", "Unknown")
                    confidence = assessment_data.get("confidence", 0)
                    
                    # Update the score value to show actual confidence
                    self.readiness.score_value.setText(f"{confidence}%")
                    self.readiness.score_badge.setText(score)
                    
                    # Update badge color based on score
                    if score == "Excellent":
                        self.readiness.score_badge.setProperty("priority", "Low")
                        self.readiness.score_badge.setStyleSheet("""
                            background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                                stop:0 rgba(22, 163, 74, 0.35), 
                                stop:1 rgba(34, 197, 94, 0.3));
                            border: 1px solid rgba(74, 222, 128, 0.7); 
                            color: #86EFAC;
                            font-size:20px; 
                            padding:10px 20px; 
                            font-weight:800;
                            letter-spacing:0.5px;
                            border-radius: 999px;
                        """)
                    elif score == "Good":
                        self.readiness.score_badge.setProperty("priority", "Medium")
                        self.readiness.score_badge.setStyleSheet("""
                            background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                                stop:0 rgba(245, 158, 11, 0.35), 
                                stop:1 rgba(251, 191, 36, 0.3));
                            border: 1px solid rgba(252, 211, 77, 0.7); 
                            color: #FDE047;
                            font-size:20px; 
                            padding:10px 20px; 
                            font-weight:800;
                            letter-spacing:0.5px;
                            border-radius: 999px;
                        """)
                    else:  # Poor
                        self.readiness.score_badge.setProperty("priority", "High")
                        self.readiness.score_badge.setStyleSheet("""
                            background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                                stop:0 rgba(220, 38, 38, 0.35), 
                                stop:1 rgba(239, 68, 68, 0.3));
                            border: 1px solid rgba(248, 113, 113, 0.7); 
                            color: #FCA5A5;
                            font-size:20px; 
                            padding:10px 20px; 
                            font-weight:800;
                            letter-spacing:0.5px;
                            border-radius: 999px;
                        """)
                    
                    # Enable/disable next button based on score
                    self.readiness.to_cases_btn.setEnabled(score in ["Excellent", "Good"])
                    
                    if score == "Excellent":
                        self.readiness.score_badge.setProperty("priority", "Low")
                    elif score == "Good":
                        self.readiness.score_badge.setProperty("priority", "Medium")
                    else:
                        self.readiness.score_badge.setProperty("priority", "High")
                    
                    self.readiness.score_badge.style().unpolish(self.readiness.score_badge)
                    self.readiness.score_badge.style().polish(self.readiness.score_badge)
                    
                    self.readiness.to_cases_btn.setEnabled(score in ["Excellent", "Good"])
                    
                    # Build and set HTML
                    html_content = build_readiness_html(assessment_data)
                    print(f"DEBUG: Setting readiness HTML - Length: {len(html_content)} characters")
                    print(f"DEBUG: Last 200 chars of HTML: ...{html_content[-200:]}")
                    self.readiness.assessment_viewer.setHtml(html_content)
                else:
                    self.readiness.assessment_viewer.setHtml(plain_to_html("Could not parse assessment."))
            
            self.readiness.progress_bar.setValue(100)
            self.readiness.progress_bar.setFormat("Complete! 100%")
            QTimer.singleShot(500, lambda: (
                self.readiness.progress_bar.setVisible(False),
                self.readiness.assessment_viewer.setVisible(True)
            ))
        
        def failed(msg):
            self.readiness.assessment_viewer.setHtml(plain_to_html(f"Error: {msg}"))
            self.readiness.score_badge.setText("Error")
            self.readiness.progress_bar.setVisible(False)
            self.readiness.assessment_viewer.setVisible(True)
        
        worker = ReadinessWorker(self.analyzer, self.current_summary, self.current_body, self.current_ac, self.feature_context, getattr(self, 'current_attachments', None))
        worker.finished_ok.connect(done)
        worker.failed.connect(failed)
        worker.start()
        self.readiness_worker = worker

    def closeEvent(self, event):
            """Ensure worker threads are stopped before closing."""
            # Stop generate worker
            if hasattr(self, 'worker') and self.worker is not None:
                if self.worker.isRunning():
                    self.worker.stop()
                    self.worker.wait(2000)
                    if self.worker.isRunning():
                        self.worker.terminate()
            
            # Stop overview worker if exists
            if hasattr(self, 'overview_worker') and self.overview_worker is not None:
                if self.overview_worker.isRunning():
                    self.overview_worker.wait(2000)
                    if self.overview_worker.isRunning():
                        self.overview_worker.terminate()
            
            # Stop readiness worker if exists
            if hasattr(self, 'readiness_worker') and self.readiness_worker is not None:
                if self.readiness_worker.isRunning():
                    self.readiness_worker.wait(2000)
                    if self.readiness_worker.isRunning():
                        self.readiness_worker.terminate()
            
            event.accept()

    def on_generate_clicked(self):
        # Check if we have ticket content (either Jira key or generated ticket with summary)
        if not self.current_key and not self.current_summary:
            QMessageBox.information(self,"No ticket","Please load a ticket first.")
            return
    
    # Disable the assess button during generation (if we're coming from ticket page)
        if hasattr(self.ticket, 'assess_btn'):
            self.ticket.assess_btn.setEnabled(False)
    
    # Navigate to cases page and show progress
        self._nav(self.cases, self.tabs.btn_cases, "Test Cases")
        self.cases.progress_bar.setVisible(True)
        self.cases.progress_bar.setValue(0)
        self.cases.progress_bar.setFormat("Starting test case generation... 0%")
        self.cases.scroll.setVisible(False)
        QApplication.processEvents()
    
        import time
        self.cases.progress_bar.setValue(15)
        self.cases.progress_bar.setFormat("Analyzing ticket context... 15%")
        QApplication.processEvents()
        time.sleep(0.3)
    
        self.cases.progress_bar.setValue(30)
        self.cases.progress_bar.setFormat("Processing acceptance criteria... 30%")
        QApplication.processEvents()
        time.sleep(0.3)
    
        self.cases.progress_bar.setValue(50)
        self.cases.progress_bar.setFormat("Generating test cases with AI... 50%")
        QApplication.processEvents()
    
        def progress_update(percentage: int, message: str):
            """Handle progress updates from worker."""
            self.cases.progress_bar.setValue(percentage)
            self.cases.progress_bar.setFormat(f"{message} {percentage}%")
            QApplication.processEvents()
    
        def done(cases, requirements, used_fallback, raw_ai, ai_err, critic_review=None):
            print(f"DEBUG: done() called with {len(cases)} cases, used_fallback={used_fallback}")
            
            # Store critic review for potential UI display
            self.critic_review = critic_review
            if critic_review:
                print(f"DEBUG: Critic review received - Approved: {critic_review.get('approved', False)}")
            
            self.cases.progress_bar.setValue(80)
            self.cases.progress_bar.setFormat("Building test case tiles... 80%")
            QApplication.processEvents()
        
            self.test_cases = cases
            self.requirements = requirements
            self.used_fallback = used_fallback
            self.raw_ai_text = raw_ai
            self.ai_error = ai_err
            self.critic_review = critic_review  # Store critic review
            
            # If we're analyzing a generated ticket, update it with the test cases
            if hasattr(self, 'current_generated_ticket') and self.current_generated_ticket:
                self.current_generated_ticket.test_cases = cases
                self.current_generated_ticket.requirements = requirements
                self.current_generated_ticket.analyzed = True
                
                # Refresh the tickets display to show updated status
                self._refresh_ticket_display()
        
            if used_fallback:
                QMessageBox.warning(self, "AI Fallback", "OpenAI output invalid; used fallback generator.")
            
            if not cases:
                QMessageBox.warning(self, "No cases", "The AI returned no usable test cases.")
                self.cases.progress_bar.setVisible(False)
                self.cases.scroll.setVisible(True)
                return
        
            self.cases.progress_bar.setValue(100)
            self.cases.progress_bar.setFormat("Complete! 100%")
            QApplication.processEvents()
        
            # Populate tiles
            print("DEBUG: About to populate tiles...")
            self.populate_case_tiles()
            print(f"DEBUG: Tiles populated. Tile count: {len(self.cases.tiles)}")
        
            # Force visibility and layout update
            self.cases.progress_bar.setVisible(False)
            self.cases.scroll.setVisible(True)
            self.cases.container.setVisible(True)
            
            # Force complete layout refresh
            self.cases.container.updateGeometry()
            self.cases.flow.invalidate()
            self.cases.container.adjustSize()
            self.cases.scroll.viewport().update()
            QApplication.processEvents()
            
            print("DEBUG: UI should now be visible")
    
        def failed(m):
            QMessageBox.critical(self, "Generation error", m)
            self.cases.progress_bar.setVisible(False)
            self.cases.scroll.setVisible(True)
    
        def finished():
            if hasattr(self.ticket, 'assess_btn'):
                self.ticket.assess_btn.setEnabled(True)

        # Stop any existing worker first
        if hasattr(self, 'worker') and self.worker is not None:
            if self.worker.isRunning():
                self.worker.stop()
                self.worker.wait(1000)  # Wait up to 1 second
                if self.worker.isRunning():
                    self.worker.terminate()  # Force terminate if needed
        
        worker = GenerateWorker(
            self.analyzer, 
            self.current_summary, 
            self.current_body, 
            self.current_ac, 
            self.feature_context,
            getattr(self, 'current_attachments', None)
        )
        worker.finished_ok.connect(done)
        worker.failed.connect(failed)
        worker.finished.connect(finished)
        worker.progress.connect(progress_update)
        worker.start()
        self.worker = worker

    # --- Cases & Export helpers
    # populate_case_tiles - EXACTLY like old version (without the setUpdatesEnabled stuff)
    def populate_case_tiles(self):
        print(f"DEBUG: populate_case_tiles called with {len(self.test_cases)} test cases")
        
        if not self.test_cases:
            print("DEBUG: No test cases to display!")
            return
        
        sorted_cases = sorted(self.test_cases, key=lambda tc: tc.priority)
        self.test_cases = sorted_cases
        
        print(f"DEBUG: Calling set_cases with {len(self.test_cases)} cases")
        self.cases.set_cases(self.test_cases, on_edit=self.edit_tile, on_remove=self.remove_tile)
        
        print(f"DEBUG: Created {len(self.cases.tiles)} tiles")
        for idx, tile in enumerate(self.cases.tiles, start=1):
            tile.set_test_case_id(idx)
            print(f"DEBUG: Set tile {idx} ID")

    def edit_tile(self, tile: EnhancedTestCaseTile):
        dlg=StepEditor(tile.case, self)
        if dlg.exec()==QDialog.Accepted:
            tile.case=dlg.case
            tile.title_lbl.setText(dlg.case.title)
            
            # Update priority text
            priority_map = {1: "High", 2: "Medium", 3: "Low", 4: "Low"}
            priority_text = priority_map.get(dlg.case.priority, "Medium")
            tile.priority_pill.setText(priority_text)
            
            tile.refresh_steps_table()
            
            # Re-sort and re-number after edit
            self.populate_case_tiles()

    def remove_tile(self, tile: EnhancedTestCaseTile):
        try:
            self.test_cases.remove(tile.case)
        except ValueError: # The case was not in the list
            pass 
        self.populate_case_tiles()

    def edit_selected_tile(self):
        # Get first selected tile
        selected_tiles = [t for t in self.cases.tiles if t.checkbox.isChecked()]
        if not selected_tiles:
            QMessageBox.information(self,"No selection","Check a tile to edit.")
            return
        self.edit_tile(selected_tiles[0])

    def remove_selected_tile(self):
        # Get first selected tile
        selected_tiles = [t for t in self.cases.tiles if t.checkbox.isChecked()]
        if not selected_tiles:
            QMessageBox.information(self,"No selection","Check a tile to remove.")
            return
        self.remove_tile(selected_tiles[0])

    def goto_export_from_cases(self):
        selected = self.cases.get_selected_cases()
        if not selected:
            QMessageBox.information(self, "Nothing selected", "Check at least one tile to export.")
            return
        
        # Populate export page with selected test cases
        self.exportp.set_cases(selected)
        
        # Navigate to export page
        self._nav(self.exportp, self.tabs.btn_export, "Export")
    
    def on_navigate_to_export(self):
        """Called when user clicks Export tab button."""
        self.goto_export_from_cases()
    
    def on_navigate_to_export(self):
        """Called when user clicks Export tab button."""
        self.goto_export_from_cases()

    def on_view_raw_ai(self):
        RawAIDialog(self.raw_ai_text, self.ai_error, self.used_fallback, self).exec()

    def on_view_traceability(self):
        TraceabilityDialog(self.test_cases, self).exec()

    def on_export(self):
        """Export selected test cases."""
        selected = self.exportp.get_selected_cases()
        if not selected:
            QMessageBox.warning(self, "Nothing selected", "Please select at least one test case to export.")
            return
        
        fmt = self.exportp.format_combo.currentText()
        
        # Determine file extension
        if "Excel" in fmt or "XLSX" in fmt:
            ext = "xlsx"
            filter_str = "Excel Files (*.xlsx)"
        else:
            ext = "csv"
            filter_str = "CSV Files (*.csv)"
        
        path, _ = QFileDialog.getSaveFileName(
            self, 
            "Save Export", 
            f"test_cases.{ext}", 
            filter_str
        )
        
        if not path:
            return
        
        try:
            if ext == "xlsx":
                # Excel export
                self._export_xlsx(path, selected, fmt)
            else:
                # CSV export
                self._export_csv(path, selected, fmt)
            
            QMessageBox.information(
                self, 
                "Export Complete", 
                f"Exported {len(selected)} test case(s) to:\n{path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")
    
    def _export_xlsx(self, path: str, cases: List[TestCase], fmt: str):
        """Export to Excel format."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Cases"
        
        # Azure DevOps required headers (in exact order required)
        headers = ["ID", "Work Item Type", "Title", "Test Step", "Step Action", "Step Expected"]
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Add test cases with steps on separate rows (Azure DevOps format)
        for idx, case in enumerate(cases, start=1):
            # Create test case ID for title (TC-001, TC-002, etc.)
            tc_id = f"TC-{idx:03d}"
            
            # Prepend TC-ID to title if not already present
            title_with_id = case.title if case.title.startswith("TC-") else f"{tc_id}: {case.title}"
            
            # First row: Test case metadata WITHOUT step info
            ws.append([
                "",  # ID (must be empty for new test cases in Azure DevOps)
                "Test Case",  # Work Item Type
                title_with_id,  # Title with TC-ID prepended
                "",  # Test Step (empty in test case row)
                "",  # Step Action (empty in test case row)
                ""   # Step Expected (empty in test case row)
            ])
            
            # Subsequent rows: Each step
            for step_num, step in enumerate(case.steps, start=1):
                ws.append([
                    "",  # ID (empty for step rows)
                    "",  # Work Item Type (empty for step rows)
                    "",  # Title (empty for step rows)
                    step_num,  # Test Step
                    step.action,  # Step Action
                    step.expected  # Step Expected
                ])
        
        # Auto-size columns
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(path)
    
    def _export_csv(self, path: str, cases: List[TestCase], fmt: str):
        """Export to CSV format."""
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            
            if "TestRail" in fmt:
                writer.writerow(["ID", "Title", "Section", "Priority", "Type", "Steps", "Expected Result"])
                for idx, case in enumerate(cases, start=1):
                    req_id = case.requirement_id or "UNMAPPED"
                    tc_id = f"{req_id} : TC-{idx:03d}"
                    
                    steps_text = "\n".join([f"{i+1}. {s.action}" for i, s in enumerate(case.steps)])
                    expected_text = "\n".join([f"{i+1}. {s.expected}" for i, s in enumerate(case.steps)])
                    
                    priority_text = "High" if case.priority == 1 else "Medium" if case.priority == 2 else "Low"
                    
                    writer.writerow([
                        tc_id,
                        case.title,
                        case.requirement_id or "",
                        priority_text,
                        case.test_type,
                        steps_text,
                        expected_text
                    ])
            
            elif "Xray" in fmt:
                writer.writerow(["Test Type", "Test Summary", "Action", "Data", "Expected Result", "Priority", "Labels"])
                for case in cases:
                    for step in case.steps:
                        writer.writerow([
                            case.test_type,
                            case.title,
                            step.action,
                            "",
                            step.expected,
                            str(case.priority),
                            "|".join(case.tags)
                        ])
            
            else:  # Generic CSV / Azure DevOps format
                # Azure DevOps compatible format
                writer.writerow(["ID", "Work Item Type", "Title", "Test Step", "Step Action", "Step Expected"])
                
                # Add test cases with steps on separate rows (Azure DevOps format)
                for idx, case in enumerate(cases, start=1):
                    # Create test case ID for title (TC-001, TC-002, etc.)
                    tc_id = f"TC-{idx:03d}"
                    
                    # Prepend TC-ID to title if not already present
                    title_with_id = case.title if case.title.startswith("TC-") else f"{tc_id}: {case.title}"
                    
                    # First row: Test case metadata WITHOUT step info
                    writer.writerow([
                        "",  # ID (must be empty for new test cases in Azure DevOps)
                        "Test Case",  # Work Item Type
                        title_with_id,  # Title with TC-ID prepended
                        "",  # Test Step (empty in test case row)
                        "",  # Step Action (empty in test case row)
                        ""   # Step Expected (empty in test case row)
                    ])
                    
                    # Subsequent rows: Each step
                    for step_num, step in enumerate(case.steps, start=1):
                        writer.writerow([
                            "",  # ID (empty for step rows)
                            "",  # Work Item Type (empty for step rows)
                            "",  # Title (empty for step rows)
                            step_num,  # Test Step
                            step.action,  # Step Action
                            step.expected  # Step Expected
                        ])

    def on_navigate_to_create_test_tickets(self):
        """Navigate to create test tickets page."""
        if not self.feature_context:
            QMessageBox.warning(self, "No feature loaded", "Please load an Epic or Initiative first from the Feature page.")
            return
        
        if not self.llm.enabled:
            QMessageBox.warning(self, "AI not configured", "Please configure OpenAI API key in the Sign In page.")
            return
        
        # Reset state
        self.generated_test_tickets = []
        self.split_plan = None
        self.current_split_index = 0
        
        # Hide "View Created Tickets" menu option since we're starting fresh
        self.ticket.action_view_created.setVisible(False)
        
        # Update UI
        epic_name = self.feature_context.get('epic_summary', '') or self.feature_context.get('initiative_summary', '')
        child_count = len(self.feature_context.get('children', []))
        
        self.create_tickets.epic_info.setText(
            f"Epic: {epic_name}\n"
            f"Child Tickets: {child_count}"
        )
        self.create_tickets.status_label.setText("Click 'Analyze and Recommend' to begin...")
        self.create_tickets.splits_scroll.setVisible(False)
        self.create_tickets.generate_btn.setVisible(False)
        self.create_tickets.analyze_btn.setVisible(True)
        
        self._nav(self.create_tickets, None, "Create Test Ticket(s)")
    
    def on_analyze_splits(self):
        """Analyze Epic and recommend test ticket splits."""
        if not self.feature_context:
            QMessageBox.warning(self, "No feature", "Please load a feature first.")
            return
        
        # Show progress
        self.create_tickets.progress_bar.setVisible(True)
        self.create_tickets.progress_bar.setValue(0)
        self.create_tickets.progress_bar.setFormat("Analyzing Epic and children... 0%")
        self.create_tickets.analyze_btn.setEnabled(False)
        
        # Start worker
        self.split_worker = AnalyzeSplitWorker(self.analyzer, self.feature_context)
        self.split_worker.finished_ok.connect(self.on_split_analysis_complete)
        self.split_worker.failed.connect(self.on_split_analysis_failed)
        self.split_worker.start()
    
    def on_split_analysis_complete(self, split_data, error):
        """Handle split analysis completion."""
        self.create_tickets.progress_bar.setVisible(False)
        self.create_tickets.analyze_btn.setEnabled(True)
        
        if error:
            QMessageBox.critical(self, "Analysis failed", f"Failed to analyze: {error}")
            return
        
        if not split_data:
            QMessageBox.warning(self, "No data", "No split recommendations returned.")
            return
        
        self.split_plan = split_data
        
        # Display splits
        self.create_tickets.clear_splits()
        splits = split_data.get('recommended_splits', [])
        
        self.create_tickets.status_label.setText(
            f"✅ Analysis complete - {split_data.get('total_tickets', 0)} test tickets recommended\n\n"
            f"Strategy: {split_data.get('reasoning', 'N/A')}"
        )
        
        for i, split in enumerate(splits, 1):
            self.create_tickets.add_split_widget(split, i)
        
        self.create_tickets.splits_scroll.setVisible(True)
        self.create_tickets.generate_btn.setVisible(True)
        self.create_tickets.analyze_btn.setVisible(False)
    
    def on_split_analysis_failed(self, error_msg):
        """Handle split analysis failure."""
        self.create_tickets.progress_bar.setVisible(False)
        self.create_tickets.analyze_btn.setEnabled(True)
        QMessageBox.critical(self, "Analysis failed", error_msg)
    
    def on_generate_test_tickets(self):
        """Generate all test tickets sequentially."""
        if not self.split_plan:
            QMessageBox.warning(self, "No plan", "Please analyze splits first.")
            return
        
        splits = self.split_plan.get('recommended_splits', [])
        if not splits:
            QMessageBox.warning(self, "No splits", "No splits to generate.")
            return
        
        # Prepare for generation
        self.generated_test_tickets = []
        self.current_split_index = 0
        
        # Show progress and start
        self.create_tickets.progress_bar.setVisible(True)
        self.create_tickets.generate_btn.setEnabled(False)
        
        self._generate_next_ticket()
    
    def _generate_next_ticket(self):
        """Generate the next ticket in sequence."""
        splits = self.split_plan.get('recommended_splits', [])
        
        if self.current_split_index >= len(splits):
            # All done - move to results
            self._show_test_tickets_results()
            return
        
        split = splits[self.current_split_index]
        epic_name = self.feature_context.get('epic_summary', '') or self.feature_context.get('initiative_summary', '')
        functional_area = split.get('functional_area', '')
        child_ticket_keys = split.get('child_tickets', [])
        
        # Filter children to get full ticket data (key + summary)
        all_children = self.feature_context.get('children', [])
        relevant_children = [c for c in all_children if c.get('key', '') in child_ticket_keys]
        
        # Store for later use when creating the GeneratedTestTicket
        self.current_relevant_children = relevant_children
        
        # Update progress
        progress_pct = int((self.current_split_index / len(splits)) * 100)
        self.create_tickets.progress_bar.setValue(progress_pct)
        self.create_tickets.progress_bar.setFormat(
            f"Generating ticket {self.current_split_index + 1}/{len(splits)}: {functional_area} - {progress_pct}%"
        )
        
        # Start generation worker
        self.ticket_gen_worker = GenerateTestTicketWorker(
            self.analyzer,
            epic_name,
            functional_area,
            relevant_children,
            self.feature_context
        )
        self.ticket_gen_worker.finished_ok.connect(self.on_ticket_generated)
        self.ticket_gen_worker.failed.connect(self.on_ticket_generation_failed)
        self.ticket_gen_worker.start()
    
    def on_ticket_generated(self, ticket_data, raw_response, error):
        """Handle single ticket generation completion."""
        if error:
            QMessageBox.critical(self, "Generation failed", f"Failed: {error}")
            self.create_tickets.progress_bar.setVisible(False)
            self.create_tickets.generate_btn.setEnabled(True)
            return
        
        if not ticket_data:
            QMessageBox.warning(self, "No data", "No ticket data returned.")
            self.create_tickets.progress_bar.setVisible(False)
            self.create_tickets.generate_btn.setEnabled(True)
            return
        
        # Now review the ticket
        self._review_generated_ticket(ticket_data, raw_response)
    
    def _review_generated_ticket(self, ticket_data, raw_response):
        """Review a generated ticket with BA/PM persona."""
        epic_name = self.feature_context.get('epic_summary', '') or self.feature_context.get('initiative_summary', '')
        
        # Start review worker
        self.ticket_review_worker = ReviewTestTicketWorker(
            self.analyzer,
            ticket_data,
            epic_name,
            self.feature_context
        )
        self.ticket_review_worker.finished_ok.connect(
            lambda review_data, error: self.on_ticket_reviewed(ticket_data, raw_response, review_data, error)
        )
        self.ticket_review_worker.failed.connect(self.on_ticket_review_failed)
        self.ticket_review_worker.start()
    
    def _ensure_source_tickets_in_description(self, description: str, child_tickets: List[Dict]) -> str:
        """Ensure source tickets are included in the description. If AI didn't add them, append them."""
        if not child_tickets:
            return description
        
        # Check if source tickets are already mentioned
        if "Source Tickets:" in description or "source tickets:" in description.lower():
            return description
        
        # AI forgot to add them - append them now
        source_parts = []
        for child in child_tickets:
            key = child.get('key', '')
            summary = child.get('summary', '')
            if key and summary:
                source_parts.append(f"{key}: {summary}")
        
        if source_parts:
            # Add blank line if description doesn't end with one
            separator = "\n\n" if description.strip() else ""
            source_section = f"{separator}Source Tickets: {', '.join(source_parts)}"
            return description + source_section
        
        return description
    
    def on_ticket_reviewed(self, ticket_data, raw_response, review_data, error):
        """Handle ticket review completion."""
        if error or not review_data:
            QMessageBox.warning(self, "Review failed", "Could not review ticket, but will save it anyway.")
            review_data = {"approved": True, "quality_score": 70, "strengths": [], "issues": [], "recommendations": []}
        
        # Store the ticket
        splits = self.split_plan.get('recommended_splits', [])
        split = splits[self.current_split_index]
        
        # Ensure source tickets are in description
        description = self._ensure_source_tickets_in_description(
            ticket_data.get('description', ''),
            self.current_relevant_children
        )
        
        ticket = GeneratedTestTicket(
            id=len(self.generated_test_tickets) + 1,
            title=ticket_data.get('summary', ''),
            summary=ticket_data.get('summary', ''),
            description=description,
            acceptance_criteria=ticket_data.get('acceptance_criteria', []),
            quality_score=review_data.get('quality_score', 70),
            ac_count=len(ticket_data.get('acceptance_criteria', [])),
            analyzed=False,
            child_tickets=self.current_relevant_children  # Full ticket data with key + summary
        )
        
        self.generated_test_tickets.append(ticket)
        
        # Move to next
        self.current_split_index += 1
        self._generate_next_ticket()
    
    def on_ticket_generation_failed(self, error_msg):
        """Handle ticket generation failure."""
        QMessageBox.critical(self, "Generation failed", error_msg)
        self.create_tickets.progress_bar.setVisible(False)
        self.create_tickets.generate_btn.setEnabled(True)
    
    def on_ticket_review_failed(self, error_msg):
        """Handle ticket review failure."""
        QMessageBox.warning(self, "Review failed", f"Review failed: {error_msg}\nTicket will be saved anyway.")
        # Continue with default review
        self.on_ticket_reviewed({}, "", {"approved": True, "quality_score": 70, "strengths": [], "issues": [], "recommendations": []}, None)
    
    def _show_test_tickets_results(self):
        """Navigate to test tickets results page."""
        self.create_tickets.progress_bar.setVisible(False)
        self.create_tickets.generate_btn.setEnabled(True)
        
        # Make "View Created Tickets" menu option visible since we have tickets now
        self.ticket.action_view_created.setVisible(True)
        
        # Update results page
        total = len(self.generated_test_tickets)
        analyzed = sum(1 for t in self.generated_test_tickets if t.analyzed)
        
        self.tickets_results.summary_label.setText(
            f"✅ {total} test tickets created | ✅ {analyzed} analyzed | ⏳ {total - analyzed} remaining"
        )
        
        # Clear and populate ticket widgets
        self.tickets_results.clear_tickets()
        
        for ticket in self.generated_test_tickets:
            self.tickets_results.add_ticket_widget(
                ticket,
                self.on_analyze_generated_ticket,
                self.on_copy_ticket_to_clipboard,
                self.on_regenerate_ticket
            )
        
        self._nav(self.tickets_results, None, "Test Tickets Created")
    
    def on_view_created_tickets(self):
        """Navigate to view previously created test tickets."""
        if not self.generated_test_tickets:
            QMessageBox.information(self, "No Tickets", "No test tickets have been created yet.")
            return
        
        # Reuse the show results function to display the stored tickets
        self._show_test_tickets_results()
    
    def on_analyze_generated_ticket(self, ticket: GeneratedTestTicket, reanalyze: bool = False):
        """Analyze a generated test ticket."""
        if reanalyze:
            # Warn about losing existing test cases
            reply = QMessageBox.question(
                self,
                "Re-analyze Ticket?",
                f"This will replace the {len(ticket.test_cases or [])} existing test cases.\n\nContinue?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply != QMessageBox.Yes:
                return
        
        # Store reference to ticket being analyzed so we can update it after test case generation
        self.current_generated_ticket = ticket
        
        # Load ticket into current state
        self.current_key = ""  # No Jira key
        self.current_summary = ticket.summary
        self.current_body = ticket.description
        self.current_ac = ticket.acceptance_criteria
        
        # Populate ticket page
        self.ticket.viewer.setHtml(build_ticket_html(ticket.summary, plain_to_html(ticket.description), ticket.acceptance_criteria))
        self.ticket.assess_btn.setEnabled(True)
        
        # Navigate to ticket page
        self._nav(self.ticket, self.tabs.btn_ticket, "Ticket to analyze")
        
        QMessageBox.information(self, "Ticket Loaded", "Generated ticket has been loaded. Click 'Next' to assess readiness or generate test cases.")
    
    def on_copy_ticket_to_clipboard(self, ticket: GeneratedTestTicket):
        """Copy ticket to clipboard in Jira format."""
        from PySide6.QtGui import QClipboard
        
        clipboard_text = f"Summary:\n{ticket.summary}\n\n"
        clipboard_text += f"Description:\n{ticket.description}\n\n"
        clipboard_text += "Acceptance Criteria:\n"
        for i, ac in enumerate(ticket.acceptance_criteria, 1):
            clipboard_text += f"{i}. {ac}\n"
        
        QApplication.clipboard().setText(clipboard_text)
        QMessageBox.information(self, "Copied", "Ticket copied to clipboard!")
    
    def on_regenerate_ticket(self, ticket: GeneratedTestTicket):
        """Regenerate a specific ticket with in-place progress display."""
        # Find the ticket's functional area from split plan
        if not self.split_plan:
            QMessageBox.warning(self, "No context", "Cannot regenerate without split plan context.")
            return
        
        splits = self.split_plan.get('recommended_splits', [])
        ticket_index = ticket.id - 1
        
        if ticket_index >= len(splits):
            QMessageBox.warning(self, "Error", "Cannot find split information for this ticket.")
            return
        
        split = splits[ticket_index]
        functional_area = split.get('functional_area', '')
        child_ticket_keys = split.get('child_tickets', [])
        
        # Filter children
        all_children = self.feature_context.get('children', [])
        relevant_children = [c for c in all_children if c.get('key', '') in child_ticket_keys]
        
        epic_name = self.feature_context.get('epic_summary', '') or self.feature_context.get('initiative_summary', '')
        
        # Find the ticket card widgets
        card_frame = self.tickets_results.findChild(QFrame, f"ticket_card_{ticket.id}")
        progress_bar = self.tickets_results.findChild(AnimatedProgressBar, f"progress_bar_{ticket.id}")
        button_container = self.tickets_results.findChild(QWidget, f"button_container_{ticket.id}")
        
        if not all([card_frame, progress_bar, button_container]):
            QMessageBox.warning(self, "Error", "Cannot find ticket UI elements.")
            return
        
        # Show progress bar and disable buttons
        progress_bar.setVisible(True)
        progress_bar.setValue(0)
        progress_bar.setFormat("Starting regeneration... 0%")
        button_container.setEnabled(False)
        
        # Start regeneration worker
        self.regen_worker = RegenerateTicketWorker(
            self.analyzer,
            ticket,
            epic_name,
            functional_area,
            relevant_children,
            self.feature_context,
            max_attempts=3
        )
        
        # Store ticket reference for callbacks
        self.regen_worker.ticket_being_regenerated = ticket
        self.regen_worker.card_frame = card_frame
        self.regen_worker.progress_bar = progress_bar
        self.regen_worker.button_container = button_container
        
        self.regen_worker.progress.connect(
            lambda pct, msg: self._on_regen_progress(ticket.id, pct, msg)
        )
        self.regen_worker.finished_ok.connect(
            lambda ticket_data, raw, review: self._on_regen_complete(ticket, ticket_data, raw, review)
        )
        self.regen_worker.failed.connect(
            lambda err: self._on_regen_failed(ticket, err)
        )
        
        self.regen_worker.start()
    
    def _on_regen_progress(self, ticket_id: int, percentage: int, message: str):
        """Update progress bar during regeneration."""
        progress_bar = self.tickets_results.findChild(AnimatedProgressBar, f"progress_bar_{ticket_id}")
        if progress_bar:
            progress_bar.setValue(percentage)
            progress_bar.setFormat(f"{message} {percentage}%")
    
    def _on_regen_complete(self, old_ticket: GeneratedTestTicket, ticket_data: Dict, raw_response: str, review_data: Dict):
        """Handle successful regeneration - replace old ticket only if quality improved."""
        # Update the ticket in the list
        ticket_index = old_ticket.id - 1
        if ticket_index < len(self.generated_test_tickets):
            new_quality_score = review_data.get('quality_score', 80)
            old_quality_score = old_ticket.quality_score
            
            # Only replace if new score is higher
            if new_quality_score > old_quality_score:
                # Ensure source tickets are in description (in case AI forgot during regen)
                description = self._ensure_source_tickets_in_description(
                    ticket_data.get('description', ''),
                    old_ticket.child_tickets
                )
                
                # Create updated ticket with same ID
                updated_ticket = GeneratedTestTicket(
                    id=old_ticket.id,
                    title=ticket_data.get('summary', ''),
                    summary=ticket_data.get('summary', ''),
                    description=description,
                    acceptance_criteria=ticket_data.get('acceptance_criteria', []),
                    quality_score=new_quality_score,
                    ac_count=len(ticket_data.get('acceptance_criteria', [])),
                    analyzed=False,
                    child_tickets=old_ticket.child_tickets
                )
                
                self.generated_test_tickets[ticket_index] = updated_ticket
                
                # Refresh the display
                self._refresh_ticket_display()
                
                QMessageBox.information(
                    self,
                    "Regeneration Successful",
                    f"Ticket improved!\n\nOld quality score: {old_quality_score}%\nNew quality score: {new_quality_score}%"
                )
            else:
                # New score is not better - keep original ticket
                QMessageBox.information(
                    self,
                    "Regeneration Complete",
                    f"Regenerated ticket did not improve quality.\n\nOriginal score: {old_quality_score}%\nNew score: {new_quality_score}%\n\nKeeping original ticket."
                )
    
    def _on_regen_failed(self, ticket: GeneratedTestTicket, error_msg: str):
        """Handle regeneration failure."""
        # Re-enable buttons and hide progress
        progress_bar = self.tickets_results.findChild(AnimatedProgressBar, f"progress_bar_{ticket.id}")
        button_container = self.tickets_results.findChild(QWidget, f"button_container_{ticket.id}")
        
        if progress_bar:
            progress_bar.setVisible(False)
        if button_container:
            button_container.setEnabled(True)
        
        QMessageBox.warning(
            self,
            "Regeneration Failed",
            f"Could not improve ticket quality:\n\n{error_msg}"
        )
    
    def _refresh_ticket_display(self):
        """Refresh the ticket display after regeneration."""
        # Clear and repopulate
        self.tickets_results.clear_tickets()
        
        for ticket in self.generated_test_tickets:
            self.tickets_results.add_ticket_widget(
                ticket,
                self.on_analyze_generated_ticket,
                self.on_copy_ticket_to_clipboard,
                self.on_regenerate_ticket
            )
        
        # Update summary
        total = len(self.generated_test_tickets)
        analyzed = sum(1 for t in self.generated_test_tickets if t.analyzed)
        self.tickets_results.summary_label.setText(
            f"✅ {total} test tickets created | ✅ {analyzed} analyzed | ⏳ {total - analyzed} remaining"
        )
    
    def on_export_all_test_tickets(self):
        """Export all generated test tickets to Jira format."""
        if not self.generated_test_tickets:
            QMessageBox.warning(self, "No tickets", "No test tickets to export.")
            return
        
        # Ask for directory
        dir_path = QFileDialog.getExistingDirectory(self, "Select Export Directory")
        if not dir_path:
            return
        
        try:
            export_dir = Path(dir_path)
            
            for ticket in self.generated_test_tickets:
                filename = f"{slugify(ticket.title)}.txt"
                filepath = export_dir / filename
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(f"Summary:\n{ticket.summary}\n\n")
                    f.write(f"Description:\n{ticket.description}\n\n")
                    f.write("Acceptance Criteria:\n")
                    for i, ac in enumerate(ticket.acceptance_criteria, 1):
                        f.write(f"{i}. {ac}\n")
            
            QMessageBox.information(self, "Export Complete", f"Exported {len(self.generated_test_tickets)} tickets to:\n{dir_path}")
        
        except Exception as e:
            QMessageBox.critical(self, "Export failed", f"Failed to export: {e}")

    def clear_ticket(self):
        self.current_key=self.current_summary=self.current_body=""; self.current_ac=[]; self.current_attachments=[]; self.feature_attachments=[]
        self.ticket.key_edit.clear(); self.ticket.viewer.setHtml(plain_to_html("(cleared)")); self.ticket.assess_btn.setEnabled(False)
        QMessageBox.information(self,"Cleared","Ticket cleared. Enter a new Jira key to continue.")

# ---------- Main ----------
def main():
    app=QApplication(sys.argv)
    load_brand_fonts(); app.setStyleSheet(NEBULA_QSS)
    w=MainWindow(); w.show(); sys.exit(app.exec())

if __name__ == "__main__":
    main()
